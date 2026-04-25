#!/usr/bin/env python3
"""
ksd_ai_translator.py

AI Karminder Singh Dhillon — переводчик СГГС через ChatGPT (Playwright).

Берёт ang_json → обрабатывает каждый шабд через ChatGPT в браузере
с опорой на ksd_knowledge.db → выдаёт docx.

Принципы вывода:
  - Оригинальный порядок строк и шабдов НЕ меняется.
  - На каждую строку: Гурмукхи → романизация → KSD-перевод (смысл выше художественности).
  - Синий блок (контекст/переход) — только там, где это важно.
  - Художественный вариант — отдельная строка, курсив, только если добавляет ценность.
  - Confidence: HIGH/MEDIUM/LOW — хранится в JSON, в docx — только резюме.
  - Свои JSON в ksd_ang_json/ — исходные ang_json НЕ ТРОГАЕМ.

Использование:
  python3 ksd_ai_translator.py --ang 1          # один анг
  python3 ksd_ai_translator.py --ang 1-8        # диапазон (Джап Джи)
  python3 ksd_ai_translator.py --resume         # продолжить с места остановки
  python3 ksd_ai_translator.py --rebuild-docx   # пересобрать docx из готовых JSON
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sqlite3
import sys
import time
import unicodedata
import shutil
from itertools import groupby
from pathlib import Path
from typing import Any

from gurbani_romanization import roman_display_from_gurmukhi as shared_roman_display_from_gurmukhi

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Inches = None
    Pt = None

    def RGBColor(r, g, b):  # type: ignore[no-redef]
        return (r, g, b)

try:
    from playwright.sync_api import Page, TimeoutError as PWTimeout, sync_playwright
    from playwright_stealth import Stealth
except ImportError:
    Page = Any  # type: ignore[assignment,misc]
    PWTimeout = TimeoutError  # type: ignore[assignment]
    sync_playwright = None
    Stealth = None

# ─── пути ────────────────────────────────────────────────────────────────────

SCRIPT_DIR   = Path(__file__).parent
BASE_DIR     = SCRIPT_DIR.parent
ANG_JSON_DIR = BASE_DIR / "custom_khoj_sahib_singh" / "ang_json"
DB_PATH      = SCRIPT_DIR / "ksd_knowledge.db"

KSD_JSON_DIR = SCRIPT_DIR / "ksd_ang_json"   # наши JSON — исходные ang_json НЕ трогаем
KSD_BACKUP_DIR = SCRIPT_DIR / "ksd_ang_json_backups"
OUT_DIR      = SCRIPT_DIR / "output"
LOG_DIR      = SCRIPT_DIR / "logs"

for d in [KSD_JSON_DIR, KSD_BACKUP_DIR, OUT_DIR, LOG_DIR]:
    d.mkdir(exist_ok=True)

BOT_PROFILE    = BASE_DIR / "custom_khoj_sahib_singh" / "bot_profile"
PROGRESS_FILE  = SCRIPT_DIR / "ksd_translator.progress"
LAST_ATTEMPT_FILE = SCRIPT_DIR / "ksd_last_attempt.json"
OUTPUT_DOCX    = OUT_DIR / "SGGS_KSD_Russian.docx"
CHAT_URL       = "https://chatgpt.com/"

# ─── Playwright stealth ───────────────────────────────────────────────────────

_stealth = Stealth() if Stealth is not None else None


JAPJI_CONTEXT_RULES = """\
━━━ КОНТЕКСТ ДЖАП ДЖИ: ЗАПРЕТ ПЛОСКОЙ КОСМОЛОГИИ ━━━
Для ангов 1-8, особенно Джап Джи, держи высший приоритет контекста произведения:
Гурбани прежде всего о духовной реализации внутри сознания, а не о физическом
описании мироздания, создании миров, подсчёте миров или космологии.

Если строка выглядит как «создание мира», «сотворение миров», «тысячи миров»,
«реки», «небеса/земли/планы», сначала проверь:
1. Рахао/тему паури и ближайшие предыдущие строки.
2. Связь с Хукамом, кудратом, внутренней реализацией, духовными состояниями.
3. Пользовательский/художественный AST-reference, если он дан.

Гурбани себе не противоречит. Если перевод создаёт конфликт с базовыми KSD-
принципами (Творец внутри, Гурбани о сейчас, не внешняя космология, не
ритуализм), пересмотри рамку перевода до того, как писать JSON.

Веды, Пураны, грантхи и другие тексты в Гурбани обычно упоминаются не ради их
«возвышенного смысла» и не как авторитет выше Гурбани. Читай такие места как
вопрос об открытии или не-открытии их внутренней духовной реальности.

Если в строке явно не указано, что речь о внешних людях, сначала проверь чтение
как внутреннего процесса: состояния разума, совести, самости, пороков,
добродетелей, духовного поиска и реализации внутри сознания.

Пиши живым человеческим русским. Не используй канцелярские слова вроде
«интернализация», если можно сказать «внутреннее принятие», «осознание внутри»,
«впитать», «сделать частью себя», «прожить внутри» или другой контекстный
вариант.

Sahib Singh используй только как низкоприоритетную лексико-грамматическую
подсказку. Он НЕ должен задавать смысловую рамку, если тянет к буквальной
космологии или внешнему мироописанию.

Правило для ਅਸੰਖ / asaṅkh:
«Неисчислимость» в Джап Джи продолжает контекст кудрата и бесчисленных
проявлений/реализаций Хукама. Не превращай это в инвентаризацию внешних миров.
Чтения грантхов, вед, духовных практик и голосов - это множество проявлений
духовного поиска внутри кудрата Творца.

Правило для ਅਖਰੀ / ਅਖਰ / akhrī / akhar:
Не своди это к механическим «буквам». Учитывай 52 A-Khar: обычные буквы и язык
являются временным берегом, через который человек выражает Наам, хвалу, знание,
письмо, речь и связь обстоятельств; но Тот, кто дал буквам силу описывать, Сам
не заключён в запись. Движение смысла: от ਅਖਰ как буквального языка к ਅ+ਖਰ как
неразрушимому духовному сообщению.

Правило для ਮਾਈ / māī / mayi:
Не переводи автоматически как «мама/мать». В Гурбани у этого слова несколько
контекстных значений:
1. Майя в обычном смысле.
2. Непросветлённость, невежество, майическое влияние на жизнь.
3. Состояние ума/разума; иногда обращение к собственному разуму.
4. Мать/мама — только если контекст явно родительский: отец/мать, семья,
рождение, физическая мать.

Пример: «ਕਹੁ ਨਾਨਕ ਜਿਨਿ ਧੂਰਿ ਸੰਤ ਪਾਈ ॥ ਤਾ ਕੈ ਨਿਕਟਿ ਨ ਆਵੈ ਮਾਈ» — здесь māī
это не мать, а Майя/непросветлённость, которая не приближается к тому, кто
осознал Шабд. Пример: «ਕਵਨ ਗੁਨ ਪ੍ਰਾਨਪਤਿ ਮਿਲਉ ਮੇਰੀ ਮਾਈ» — здесь «моя māī»
может быть разумом: какое качество должен реализовать мой разум, чтобы
соединиться с Прана-мастером. Пример: «ਸੋ ਸੇਵਹੁ ਜਿਸੁ ਮਾਈ ਨ ਬਾਪੁ» — здесь
родительский контекст, поэтому māī = мать.

Правило для ਮੋਹ / moh / moha:
Не оставляй «мох» как непонятное русское слово и не своди его к простой
эмоциональной привязанности. В Гурбани это майическая привязанность,
очарованность преходящим, захваченность майей и помрачение различения, из-за
которых разум забывает Единого и теряет духовное движение. Если рядом есть
майя, тришна, переправа через поток майической запутанности, тонущий разум,
болото, грязь, трясина, сеть, оковы или забывание Творца, читай moh как
внутреннюю связанность майей: разум застревает в очаровании
майического и принимает преходящее за опору. Художественный образ болота,
грязи или трясины допустим, но обязательно раскрывай, что это трясина
майической привязанности/очарования майей, а не физическая грязь и не просто
бытовая привязанность к людям.

Правило для ਭਵਜਲ / bhavjal:
Не пиши в русских переводческих полях «бхаваджал»: это непрозрачный термин для
читателя. В анализе можно хранить исходный roman, но в переводе раскрывай
смысл: поток/океан майической запутанности, внутреннее блуждание, круг
захваченности майей и повторяющимися состояниями сознания. Если строка говорит
о taran/переправе, переводи как «переправиться через майическую запутанность»,
«выйти из потока майического блуждания» или близко по контексту.

Правило для ਰਜਾਇ / razā / razāi:
В русских переводческих полях не оставляй «раза» как основной термин. Передавай
смысл по-русски: Воля, Высшая Воля, Воля Хукама, согласие с Волей Творца.
Правильные падежи: «в Воле», «с Волей», «к Воле», «жизнь в Воле», «согласно
Воле». Если важно сохранить связь с Хукамом, пиши «Воля Хукама» или «движение
в Воле Хукама».

Правило для ਕੁਦਰਤ / kudrat:
Не отделяй Кудрат от Хукама. Кудрат — это проявленная природа Воли Творца,
действие Хукама, проявление этой Воли внутри и вокруг. В русских полях лучше
раскрывать: «проявление Воли/Хукама», «поле действия Хукама», «проявленная
Кудрат Творца», чем оставлять слово как самостоятельную “природу”. Если
оставляешь «Кудрат», рядом уточняй её как проявление Воли Хукама.

Правило для ਸੋਹਿਲਾ / sohilā / sohilai:
Не оставляй «сохила» как голый термин, если строка не является только
заголовком произведения. Раскрывай смысл как песнь Мира, песнь внутреннего
соединения, поток хвалы/Хукама, через который сознание приходит к постоянному
духовному счастью. В Рахао «ਹਉ ਵਾਰੀ ਜਿਤੁ ਸੋਹਿਲੈ ਸਦਾ ਸੁਖੁ ਹੋਇ» не пиши «отдаю себя той
сохиле» как вещному объекту; лучше: «склоняюсь перед той песней Мира / тем
потоком Хукама, через который раскрывается постоянный внутренний покой».

Правило для ਸਤਿਗੁਰ / ਸਤਿਗੁਰਿ / satigur:
В русском переводе пиши «Сатгуру», а не «Сатигур/Сатигуру». Сихари в таких
формах часто грамматическая и не должна становиться русской гласной. При этом
можно оставлять имена Бога и обращения как имена: Сахиб, Хар, Раам, Гобинд,
Прабху и другие, если это не скрывает смысл строки.

Правило для имен Бога, особенно ਸਾਹਿਬ / ਹਰਿ:
Если в строке стоит «Сахиб», его можно оставлять как имя Бога. Не обязательно
заменять на «Творец». Если в gurmukhi/roman стоит ਹਰਿ / har, не растворяй его
автоматически в «Творец»: чаще пиши «Хар», «Вездесущий Хар», «Хар-Наам»,
«качества Хара» по падежу и контексту. «Творец» допустим как пояснение, но не
должен стирать имя Хар там, где оно прямо дано в строке.

Правило для ਸੁਖ / sukh:
Не пиши в русских переводческих полях «сух». Читатель воспринимает это как
русское прилагательное, а не как термин Гурбани. Переводи по контексту:
духовное счастье, внутреннее счастье, устойчивый покой, блаженное состояние,
счастье сахаджа. Если рядом sahaj, Naam, Har или seva, чаще выбирай «духовное
счастье» или «счастье сахаджа», а не сухой термин.

Правило для ਲਿਵ / liv:
Не оставляй «лив» как непонятный термин в русском переводе. Передавай по
контексту как устойчивый внутренний фокус, любовное внимание, сосредоточенное
погружение сознания, направленность ума к Хару/Нааму/Шабду. Если строка
говорит «ਹਰਿ ਲਿਵ», пиши «устойчивый фокус на Харе», «любовное внимание к
Хару» или «сознание сосредоточено на Харе», а не «лив с Творцом».

Правило для ਸੰਜੋਗ / sanjog:
Не оставляй «санжог/санджог» как непонятное слово. По контексту это соединение,
связь, сложившееся условие, а иногда прежний духовный посев, который созревает
в Хукаме. Не превращай это в механическую судьбу. В строках о «прошлом
написании/прежних санжогах» раскрывай как прежний посев/связь условий в
Хукаме, через которые сознание приходит к Гуру-Шабду.

Правило для ਸਾਕਤ / sākat:
Не оставляй «сакат» как непонятный термин и не путай с «шакта» как внешней
религиозной группой. В KSD сначала читай как внутреннее манмукх-состояние:
сознание, отвернувшееся от Хара, Гуру-Шабда и Наама, захваченное майей и
Хоумэ. В переводе раскрывай: «оторванное от Хара сознание», «манмукх-самость»,
«состояние, отвернувшееся от Наама».

Правило для ਹਉਮੈ / haumai / Хоумэ:
Не своди Хоумэ к простой гордыне, самолюбию или психологическому «эго».
Хоумэ — это подверженность влиянию «я» или «мы», эгоистическая
самососредоточенность, чувство «я существую отдельно», а также тот внутренний
механизм, через который Майя сокращает и искажает восприятие. Майя на уровне
восприятия действует через Хоумэ: ум присваивает внешние образы, реакции,
истории, групповые «мы» и личные «я», а затем принимает это за реальность.

Хоумэ не называй абсолютным злом. Это условие опыта отдельности и часть
проявления в Хукаме; но когда оно не осознано, человек попадает под влияние
своей истории, генетики, кармы, группы, чужой оценки, социальной картинки,
страха, желания, гнева, привязанности и гордости. Когда хватка Хоумэ
ослабевает через Гуру-Шабд, Наам, совесть и жизнь в Воле Хукама, ослабевает и
влияние Майи: восприятие становится шире, точнее, сострадательнее и меньше
зависит от внешнего одобрения или коллективного «мы».

Современный прикладной слой Хоумэ:
Не переводить Хоумэ как банальное «эгоизм плох». В обычной жизни здоровый
интерес к себе, забота о себе, достоинство, гордость за реальное достижение и
умение не растворяться в других могут быть добродетельными. Проблемным
становится то Хоумэ, которое делает разум слугой уже имеющихся эмоций,
убеждений и групповой идентичности. Его признаки: избирательное восприятие,
мотивированное мышление, поиск резонанса вместо реальности, поиск подтверждения
у похожих людей, превращение собственных омрачений в добродетели, отказ от
самостоятельного мышления под видом духовности, зависимость от социального
одобрения.

Главный порок Хоумэ в современной формулировке: игнорирование обратной связи
от Реальности — тела, совести, здравого смысла, логики, окружающего мира,
плодов действий и Хукама. Разум здесь является «сикхом», учеником: он может
быть обращён к самому себе и утопать в собственных отражениях, либо быть
обращён к Гуру — к живой обратной связи от Реальности, Гуру-Шабду, совести и
глубокому освоению Хукама. В этой рамке manmukh — обращённость к уму и его
отражениям; gurmukh — обращённость к Гуру и Реальности.

Если строка говорит о рае/аде, нараке/сураге, рождении/смерти в контексте
Хоумэ, не делай это внешней географией после смерти. Читай как состояния
восприятия внутри «я»-концепции: жар Хоумэ, наслаждение и страдание,
ограниченность отдельного существования. Освобождение от Хоумэ — не уничтожение
функционального человека, а растворение ложной отдельности в Хукаме, когда
капля начинает видеть ближе к Океану Ооангкара.

Для образов шипа/колючки пиши «колючка Хоумэ», «шип эгоистической
самососредоточенности», «боль отдельного “я”». Для «ਹਉਮੈ ਦੀਰਘ ਰੋਗੁ ਹੈ ਦਾਰੂ ਭੀ
ਇਸੁ ਮਾਹਿ» держи рамку: Хоумэ — глубокая/длительная болезнь, но лекарство
находится внутри неё, потому что осознавание Хоумэ открывает выход за её
пределы.

Правило для ਸਤਿ / ਸਚੁ / ਸਤਸੰਗਤਿ / ਸਾਧਸੰਗਤਿ:
Sat/Sach не тяни автоматически в абстрактную «истину». В KSD это обычно то,
что относится к Творцу, Божественному, подлинно пребывающему, а не философская
категория «Истина». Переводи по контексту: Божественный, относящийся к Творцу,
подлинно пребывающий, Наам-Творца, Хукам-Творца.

Satsangat/Sadhsangat не обязательно означает внешнее общество людей. Сначала
проверь чтение как соединение с Божественным, с Гуру-Шабдом, с Наамом, с
пространством божественного осознания. Внешняя община допустима только если
контекст явно говорит о людях/собрании/совместном слушании.

Правило для ਘਰੁ / ਘਰਿ / ghar / дом:
Не переводи автоматически как физический дом. В Гурбани ghar часто может быть
внутренним пространством соединения с Создателем: совесть, духовная мотивация,
место внутри, где звучит Гуру-Шабд, происходит вичар и удерживается Наам.
Если ghar связан с сатсангат/садхсангат, хвалой, киртаном, вичаром,
духовным счастьем/сахаджем или «береги этот дом», сначала читай это как пространство
совести и соединения с Божественным, которое нужно беречь. Физическое место
или внешнее собрание допустимы только при явных внешних маркерах.

Правило для ਦਰੁ / ਦੁਆਰੁ / ਮਹਲੁ / dar / duar / mahal / обитель / врата:
Не превращай строки о «вратах», «двери», «доме», «махале» или «обители»
Творца в внешний адрес, куда нужно прийти, или в вопрос «где живёт Творец».
Творец всепронизывает: Он не ограничен местом, Он узнаётся внутри разума,
совести, Наама и Хукама. Сначала проверяй чтение как внутреннее пространство
осознания: врата восприятия, дом совести, махал присутствия, место внутри, где
сознание признаёт заботу и власть Творца. Для «ਸੋ ਦਰੁ ... ਸੋ ਘਰੁ ... ਜਿਤੁ ਬਹਿ
ਸਰਬ ਸਮਾਲੇ» и похожих строк смысл не в географии небесного двора, а в том,
каким становится внутреннее пространство, где распознаётся всепронизывающий
Творец, поддерживающий всё.

Правило для художественного варианта:
Художественный вариант не удаляй и не делай сухим пересказом. Он должен
сохранять тот же KSD-смысл, что и основной перевод, но звучать живо, образно и
человечески. Не добавляй красивость, которая меняет смысл, тянет к внешней
космологии, ритуализму, физическим людям/местам без контекста или противоречит
Гурбани. Если основной перевод исправляет буквальную ошибку, художественный
вариант тоже обязан быть исправлен в той же рамке.

Правило для 21-й паури Джап Джи: ਵੇਲਾ / ਵਖਤੁ / ਬਾਣੀ ਬਰਮਾਉ
Эта паури НЕ о времени физического создания мира. Читай ਵੇਲਾ и ਵਖਤੁ через
призму Амрит-велы как состояния реализации: это вопрос о той Веле/Вакхте
внутреннего раскрытия, которую не знают внешние религиозные системы, календарь
и книжные записи. Пандиты не нашли эту Велу в Пуранах; кадии не нашли этот
Вакхт в записи Корана; йогические системы тоже не дают такой меры. Это не спор
о космологии, а утверждение предела внешнего знания перед внутренней
реализацией Хукама.

В строке «ਸੁਅਸਤਿ ਆਥਿ ਬਾਣੀ ਬਰਮਾਉ» не тянуть ਬਰਮਾਉ автоматически к Брахме.
В этом месте проверяй персидско-лексическое чтение: бармао/бармау как щит,
защита, покров, опора. В связке с ਬਾਣੀ это может указывать на Бани как
защитный щит/опору духовной реализации, а не на пуранического Брахму.

Правило для богов, девте, Аллаха, Ишара/Шивы и других религиозных образов:
Гуру Нанак часто переопределяет уже существующие представления, зная, что
слушатель уже вкладывает в них свой смысл. Если в строке сказано, что боги,
девте, Ишар, Брахма, Индра, Аллах, пандиты, кадии или другие религиозные
фигуры «воспевают» Творца, не читай это как утверждение их внешнего статуса.
Проверь внутреннее чтение: через слушание Хукама разные внутренние образы
божественного, уже живущие в сознании человека, устремляются к Наивысшему.
Человек мог верить в Аллаха, Ишара/Шиву или любое божество; Гурбани обращается
к этому внутреннему образу и переводит его к Единому, Нааму, Хукаму, Творцу
внутри.

Это правило связано с ੴ / Ik Oankaar: смысл символа в том, что любые разные
внутренние представления о божественном, известные человеку из его религиозной
культуры, не остаются самостоятельными высшими реальностями. Они получают
правильное место только через Единого Ооангкара — единый Хукам, единый Наам,
единый Творец внутри.
"""

# ─── цвета docx ──────────────────────────────────────────────────────────────

COLOR_GURMUKHI    = RGBColor(0x55, 0x00, 0x00)   # тёмно-красный
COLOR_ROMAN       = RGBColor(0x55, 0x55, 0x55)   # серый
COLOR_TRANSLATION = RGBColor(0x00, 0x55, 0x00)   # тёмно-зелёный
COLOR_CONTEXT     = RGBColor(0x00, 0x55, 0xAA)   # синий — "crossing over"
COLOR_COMMENT     = RGBColor(0x88, 0x88, 0x88)   # серый мелкий
COLOR_ARTISTIC    = RGBColor(0x33, 0x33, 0x66)   # тёмно-синий — художественный
COLOR_RAHAO       = RGBColor(0x88, 0x00, 0x44)   # пурпурный — Рахао
COLOR_HEADING     = RGBColor(0x22, 0x22, 0x22)

# ─── system-блок (начало каждого нового чата) ────────────────────────────────
# ChatGPT не имеет отдельного system-поля — вставляем в первое сообщение.

def build_system_block(db_conn: sqlite3.Connection) -> str:
    cur = db_conn.cursor()

    # Принципы KSD
    cur.execute("SELECT num, title, description FROM ksd_principles ORDER BY num")
    principles_text = "\n".join(
        f"Принцип {p[0]} ({p[1]}): {p[2][:280]}"
        for p in cur.fetchall()
    )

    # Концепты Nanak Canvas
    cur.execute(
        "SELECT concept, ksd_meaning FROM canvas_concepts "
        "WHERE source='canvas_ksd' LIMIT 14"
    )
    canvas_text = "\n".join(
        f"  • {c[0].upper()}: {c[1][:180]}"
        for c in cur.fetchall()
    )

    # Грамматические правила
    cur.execute("SELECT pattern, meaning FROM grammar_rules LIMIT 18")
    grammar_text = "\n".join(f"  {g[0]} → {g[1]}" for g in cur.fetchall())

    # Few-shot примеры (паури 1 из Джап Бани)
    cur.execute(
        "SELECT pauri_num, word_analysis, ksd_translation "
        "FROM ksd_examples WHERE source='jbani_v2' LIMIT 2"
    )
    few_shot_parts = []
    for row in cur.fetchall():
        wa = json.loads(row[1])
        wa_text = "\n".join(
            f"    {w['roman']}: [{w['literal']}] → [{w.get('ksd_meta', '')}]"
            for w in wa[:6]
        )
        few_shot_parts.append(
            f"  [Паури {row[0]}]\n{wa_text}\n  Translation: {row[2][:200]}"
        )
    few_shot = "\n\n".join(few_shot_parts)

    return f"""Ты — AI-версия Карминдера Сингха Диллона (KSD), интерпретатора Сири Гуру Грантх Сахиб.

━━━ МЕТОДОЛОГИЯ KSD (8 принципов) ━━━
{principles_text}

━━━ ХОЛСТ НАНАКА — концептуальный фрейм ━━━
Следующие понятия в Гурбани НЕ буквальны — KSD их переопределяет:
{canvas_text}

━━━ ГРАММАТИКА ГУРБАНИ ━━━
Окончания слов меняют смысл:
{grammar_text}

━━━ ПРИМЕРЫ KSD-РАЗБОРА (few-shot) ━━━
{few_shot}

━━━ НЕИЗМЕННЫЕ ПРАВИЛА ━━━
1. Порядок строк в шабде НЕ меняется никогда.
2. Смысл > художественность. Художественный вариант — только если добавляет ценность.
3. Confidence 0.0–1.0 на каждую строку.
4. Рахао (ਰਹਾਉ) — тема/суть шабда. Начинать анализ с неё.
5. Нет внешних философий (индуизм, ислам, йога).
6. Гурбани — ДЛЯ МЕНЯ (читателя), не про кого-то снаружи.
7. Творец ВНУТРИ, не снаружи.
8. Гурбани о СЕЙЧАС, не о загробной жизни.
9. Не объяснять очевидное. Не быть снисходительным.
10. Все русские поля вывода должны быть написаны по-русски, без латиницы.
11. В `ksd_translation`, `artistic_ru`, `context_note`, `shabad_summary`, `rahao_theme`, `confidence_reason` латинские буквы запрещены.
12. Если термин Гурбани нельзя перевести буквально, передавай его кириллицей:
    maya -> майя, mayic -> майический, seva -> сева, Thakur -> Тхакур,
    bhagat -> бхагат, bhagti -> бхагти, raza/razaa/razā -> Воля,
    dhyan/dhyan -> дхьян, jaap/jap -> джап, tapas -> тапас.
13. Латиница допустима только в поле `roman`.

{JAPJI_CONTEXT_RULES}

━━━ СИНИЕ ВСТАВКИ [[...]] — КЛЮЧЕВОЙ ПРИЁМ ━━━
Если для понимания строки необходимо добавить слова, которых нет в ней буквально,
но которые подразумеваются из предыдущих строк, Рахао или контекста шабда —
заключи эти слова в двойные квадратные скобки [[вот так]] прямо внутри ksd_translation.

В итоговом документе они будут выделены СИНИМ цветом в квадратных скобках [вот так].
Это показывает читателю: «это слово добавлено из контекста, не буквально».

Примеры:
  «[[Хукам Творца —]] сила, движущая нашим существованием»
  «Множество [[духовных искателей]] воспринимают [[это как]] дар»
  «погрузиться внутрь [[в своё Сознание]]»

Когда добавлять:
  ✓ Пропущенный подлежащий, восстанавливаемый из предыдущей строки
  ✓ Тема шабда (из Рахао), без которой строка непонятна
  ✓ Понятие, определённое несколькими строками раньше
Когда НЕ добавлять:
  ✗ Слова «для красоты» или удлинения фразы
  ✗ Интерпретации, не следующие из текста однозначно
  ✗ Если строка понятна без добавлений

━━━ ФОРМАТ ОТВЕТА ━━━
Верни строго JSON между BEGIN_KSD_JSON и END_KSD_JSON.
Никакого markdown, никаких code fences снаружи маркеров.
Никогда не используй ASCII-кавычки " " внутри строк JSON — ломает парсинг. Используй «ёлочки».

{{
  "ang": <int>,
  "shabad_id": <int>,
  "rahao_verse_id": <int или null>,
  "rahao_theme": "<тема шабда — одна фраза>",
  "lines": [
    {{
      "verse_id": <int>,
      "is_rahao": <true/false>,
      "word_analysis": [
        {{
          "roman": "<слово>",
          "literal_ru": "<буквальное>",
          "ksd_meta_ru": "<KSD-метафора>",
          "confidence": <0.0–1.0>,
          "grammar_note": "<опционально>"
        }}
      ],
      "ksd_translation": "<KSD-перевод строки>",
      "artistic_ru": "<художественный вариант или пустая строка>",
      "context_note": "<синий блок: пустая строка если не нужен>",
      "confidence": <0.0–1.0>,
      "confidence_reason": "<почему такой уровень>"
    }}
  ],
  "shabad_summary": "<резюме смысла шабда — 1-2 предложения>"
}}"""


# ─── ang_json загрузка ────────────────────────────────────────────────────────

def load_ang(ang_num: int) -> dict | None:
    path = ANG_JSON_DIR / f"ang_{ang_num:04d}.json"
    if not path.exists():
        return None
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def group_by_shabad(ang_data: dict) -> list[list[dict]]:
    shabads = []
    for _, group in groupby(ang_data["lines"], key=lambda l: l["shabad_id"]):
        shabads.append(list(group))
    return shabads


# Паттерн конца строфы: ॥੧॥ ॥੨॥ ... ॥੧੦॥ — цифры гурмукхи после двойной черты
_STANZA_END_RE = re.compile(r"॥[੦-੯]+॥")

def find_rahao_vids(lines: list[dict]) -> list[int]:
    """
    Возвращает verse_id всех строк Рахао-куплета.

    Правило: находим строку с маркером ਰਹਾਉ — это последняя строка куплета.
    Затем идём назад и добавляем предыдущие строки, у которых НЕТ маркера
    конца строфы (॥੧॥ ॥੨॥ и т.д.) — они продолжают тот же куплет.

    Пример:
      v557  ਕੈਸੀ ਆਰਤੀ ਹੋਇ ॥ ਭਵ ਖੰਡਨਾ ਤੇਰੀ ਆਰਤੀ ॥        ← входит в куплет
      v558  ਅਨਹਤਾ ਸਬਦ ਵਾਜੰਤ ਭੇਰੀ ॥੧॥ ਰਹਾਉ ॥              ← Рахао (последняя)
    """
    # Ищем индекс строки с ਰਹਾਉ
    rahao_idx = None
    for i, line in enumerate(lines):
        if "ਰਹਾਉ" in line.get("gurmukhi", "") or "ਰਹਾਓ" in line.get("gurmukhi", ""):
            rahao_idx = i
            break
    if rahao_idx is None:
        return []

    rahao_vids = [lines[rahao_idx]["verse_id"]]

    # Идём назад, собираем строки куплета
    for i in range(rahao_idx - 1, -1, -1):
        gm = lines[i].get("gurmukhi", "")
        # Если у предыдущей строки есть маркер конца строфы — она не часть Рахао
        if _STANZA_END_RE.search(gm):
            break
        rahao_vids.insert(0, lines[i]["verse_id"])

    return rahao_vids


# Обратная совместимость — одна строка
def find_rahao_vid(lines: list[dict]) -> int | None:
    vids = find_rahao_vids(lines)
    return vids[-1] if vids else None


# ─── фикс битой романизации ───────────────────────────────────────────────────

# Символы не из латиницы/диакритики — признак битой кодировки в roman
_NON_LATIN_RE = re.compile(r"[\u0600-\u06FF\u0900-\u097F\u0A00-\u0A7F]")
_ROMAN_TOKEN_RE = re.compile(r"[A-Za-zāīūṭḍṇṛñśṃṅ]+")
_GURMUKHI_TOKEN_RE = re.compile(r"[\u0A00-\u0A7F]+")
MAYI_ROMAN_TOKENS = {"mai", "mayi", "maii", "maee", "māī"}
MAYI_GURMUKHI_FORMS = ("ਮਾਈ", "ਮਾਇ", "ਮਾਈਂ")
MOH_ROMAN_TOKENS = {"moh", "moha", "mohi", "mohī", "mōh", "mōhi"}
MOH_GURMUKHI_FORMS = ("ਮੋਹ", "ਮੋਹੁ", "ਮੋਹਿ", "ਮੋਹੀ")
BHAVJAL_ROMAN_TOKENS = {"bhavjal", "bhavajal", "bhaujal"}
BHAVJAL_GURMUKHI_FORMS = ("ਭਵਜਲ", "ਭਉਜਲ", "ਭਵ ਜਲ", "ਭਉ ਜਲ")
SOHILA_ROMAN_TOKENS = {"sohila", "sohilai", "sohilā", "sōhilā", "sōhilai"}
SOHILA_GURMUKHI_FORMS = ("ਸੋਹਿਲਾ", "ਸੋਹਿਲੈ")
COMMON_ROMAN_TOKENS = {
    "har", "gur", "jan", "jin", "tin", "kau", "mo", "mai", "main", "hau",
    "tu", "toon", "mera", "mere", "meri", "ham", "te", "so", "jo", "sab",
    "sabh", "nahi", "nahi", "sat", "rahau", "mahala", "mehala", "rag", "kar",
}
IMPORTANT_ROMAN_TOKENS = {
    "nam", "naam", "jam", "sangat", "satsangat", "simran", "hukam", "haumai",
    "gurmukh", "manmukh", "satgur", "satigur", "pargas", "paragasa", "ras",
    "seva", "sevi", "sahaj", "kirat", "kiratam", "pias",
    "isar", "brahma", "barma", "ind", "indra", "dev", "devta", "avtar",
}

CONCEPT_ALIASES: dict[str, tuple[str, ...]] = {
    "наам": ("naam", "nam", "nām", "ਨਾਮ", "ਨਾਮੁ"),
    "Naam / Наам (ਨਾਮ)": ("naam", "nam", "nām", "ਨਾਮ", "ਨਾਮੁ"),
    "Sat Naam / Сат Наам (ਸਤਿ ਨਾਮੁ)": ("sat naam", "sat nam", "ਸਤਿ ਨਾਮੁ"),
    "хукам": ("hukam", "ਹੁਕਮ"),
    "симран": ("simran", "ਸਿਮਰਨ"),
    "Simran / Симран (ਸਿਮਰਨ)": ("simran", "ਸਿਮਰਨ", "ਚੇਤਾ", "ਯਾਦ"),
    "Sikhi / Сикхи (ਸਿਖੀ)": ("sikhi", "ਸਿਖੀ", "ਸਿਖਿਆ"),
    "Sikhi-Guru-Vichar / Ученичество-Гуру-Вичар": ("sikhi", "guru", "vichar", "ਸਿਖੀ", "ਗੁਰੂ", "ਵੀਚਾਰ"),
    "Guru / Гуру (ਗੁਰੂ)": ("guru", "gur", "ਗੁਰੂ", "ਗੁਰ"),
    "Vichar / Вичар (ਵੀਚਾਰ)": ("vichar", "ਵੀਚਾਰ", "vīchār"),
    "Ooangkar / Ооангкар (ਓਅੰਕਾਰ)": ("ooangkar", "oangkar", "ik ooangkar", "ik oankaar", "ੴ", "ਓਅੰਕਾਰ"),
    "Devte / Vedantic Gods / Боги в Гурбани": (
        "isar", "ishar", "eesar", "brahma", "barma", "ind", "indra",
        "dev", "devta", "devi", "avtar", "avatar", "ramchandra", "krishna",
        "durga", "chandi", "yamraj", "dharam raj", "chitragupt",
        "ਈਸਰੁ", "ਬਰਮਾ", "ਇੰਦੁ", "ਦੇਵ", "ਅਵਤਾਰੀ",
    ),
    "мукти": ("mukti", "mukat", "mukat(i)", "ਮੁਕਤਿ", "ਮੁਕਤ"),
    "джам": ("jam", "ਜਮ"),
    "дхарм радж": ("dharam", "dharam raj", "ਧਰਮ", "ਰਾਜ"),
    "смерть": ("marn", "maran", "ਮਰਣ"),
    "реинкарнация": ("janam", "birth", "ਜਨਮ"),
    "Sant (ਸੰਤ/ਸੰਤੁ) — Sant, святой, сант": ("sant", "ਸੰਤ", "ਸੰਤੁ"),
    "гурбани": ("gurbani", "ਗੁਰਬਾਣੀ", "ਗੁਰਬਾਨੀ"),
    "дарга": ("dargah", "ਦਰਗਹ", "ਦਰਗਾਹ"),
    "нарак": ("narak", "ਨਰਕ"),
    "кал": ("kal", "ਕਾਲ", "ਕਲਿ"),
    "мох": ("moh", "moha", "ਮੋਹ", "ਮੋਹੁ", "ਮੋਹਿ"),
    "bhavjal": ("bhavjal", "bhavajal", "bhaujal", "ਭਵਜਲ", "ਭਉਜਲ"),
    "sohila": ("sohila", "sohilai", "ਸੋਹਿਲਾ", "ਸੋਹਿਲੈ"),
}
_LATIN_LEAK_RE = re.compile(r"[A-Za-zāīūṭḍṇṛñśṃṅĀĪŪṬḌṆṚÑŚṂṄ]")


MAYI_CONTEXT_HINT = """\
━━━ ЛОКАЛЬНАЯ ПОДСКАЗКА: ਮਾਈ / māī / mayi ━━━
В этом шабде встречается māī/mayi. Не переводи автоматически как «мать».
Сначала выбери значение по контексту:
1. Майя.
2. Непросветлённость/невежество/майическое влияние.
3. Состояние ума или обращение к собственному разуму.
4. Физическая мать — только при явном родительском контексте: отец/мать,
семья, рождение, физическая мать.
Если строка обращается «ਮੇਰੀ ਮਾਈ / моя māī» без явного семейного контекста,
проверь чтение как «мой разум».
"""


SATSANGAT_CONTEXT_HINT = """\
━━━ ЛОКАЛЬНАЯ ПОДСКАЗКА: SAT/SACH/SATSANGAT ━━━
Не переводи sat/sach автоматически как абстрактную «истину». Обычно это
относится к Творцу/Божественному/подлинно пребывающему. Satsangat/Sadhsangat
сначала проверяй как соединение с Божественным, Гуру-Шабдом и Наамом, а не как
внешнее общество людей.
"""


GHAR_CONTEXT_HINT = """\
━━━ ЛОКАЛЬНАЯ ПОДСКАЗКА: GHAR / ДОМ ━━━
В этом шабде встречается ghar/дом. Не переводи автоматически как физический
дом. Если рядом есть сатсангат, хвала, киртан, вичар, сукх, сахадж, Наам,
Гуру-Шабд или призыв «береги этот дом», сначала читай ghar как внутреннее
пространство совести, духовной мотивации и соединения с Создателем.
"""


MOH_CONTEXT_HINT = """\
━━━ ЛОКАЛЬНАЯ ПОДСКАЗКА: ਮੋਹ / MOH / MOHA ━━━
В этом шабде встречается moh/moha. Не оставляй в русском тексте голое слово
«мох» без раскрытия. Сначала читай его как майическую привязанность:
очарование майей, захваченность преходящим, помрачение различения и внутреннюю
связанность, из-за которой разум забывает Единого. Если есть образ болота,
грязи, трясины, сети, оков, тришны или тонущего сознания, раскрой это как
трясину майической привязанности/очарования майей.
"""


BHAVJAL_CONTEXT_HINT = """\
━━━ ЛОКАЛЬНАЯ ПОДСКАЗКА: ਭਵਜਲ / BHAVJAL ━━━
В этом шабде встречается bhavjal/bhavajal. Не пиши в русских полях
«бхаваджал»: читатель не обязан знать этот термин. Раскрывай смысл по
контексту: океан/поток майической запутанности, круг внутреннего блуждания,
состояние захваченности майей и повторяющимися состояниями сознания. Если
строка говорит о taran/переправе, переводи как переправу через майическую
запутанность или выход из потока майического блуждания, а не как неизвестное
слово.
"""


SOHILA_CONTEXT_HINT = """\
━━━ ЛОКАЛЬНАЯ ПОДСКАЗКА: ਸੋਹਿਲਾ / SOHILA ━━━
В этом шабде встречается sohilā/sohilai. В русских переводческих полях не
оставляй «сохила» как непонятный термин, кроме чистого заголовка. Раскрывай
смысл как песнь Мира, песнь внутреннего соединения, поток хвалы/Хукама, через
который приходит постоянное духовное счастье. Для строки «ਹਉ ਵਾਰੀ ਜਿਤੁ ਸੋਹਿਲੈ
ਸਦਾ ਸੁਖੁ ਹੋਇ» не пиши «отдаю себя той сохиле»; передай: склоняюсь перед той
песней Мира / тем потоком Хукама, где раскрывается постоянное духовное счастье.
"""

def fix_roman(roman: str) -> str:
    """
    Если поле roman содержит арабские/деванагари/гурмукхи символы —
    возвращаем пустую строку: ChatGPT сгенерирует романизацию заново.
    """
    if _NON_LATIN_RE.search(roman):
        return ""
    return roman


def best_roman_line(line: dict[str, Any]) -> str:
    """Берёт лучшую доступную roman-строку: сначала очищенный roman, затем site_roman."""
    primary = fix_roman(line.get("roman", ""))
    if primary:
        return primary
    fallback = fix_roman(line.get("site_roman", ""))
    if fallback:
        return fallback
    return ""


def roman_display_from_gurmukhi(gurmukhi: str, roman: str) -> str:
    """
    Backward-compatible wrapper around the shared project romanization rules.
    Keep lookup/DB roman unchanged; normalize only display output.
    """
    return shared_roman_display_from_gurmukhi(gurmukhi, roman)


def normalize_roman_token(token: str) -> str:
    """Сводит roman к plain-ascii для lookup по базе."""
    token = unicodedata.normalize("NFKD", token.lower())
    token = "".join(ch for ch in token if not unicodedata.combining(ch))
    token = re.sub(r"[^a-z]", "", token)
    return token


def roman_tokens(text: str) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for token in _ROMAN_TOKEN_RE.findall(text or ""):
        clean = normalize_roman_token(token)
        if len(clean) < 2 or clean in seen:
            continue
        seen.add(clean)
        out.append(clean)
    return out


def roman_ngrams(text: str, min_len: int = 2, max_len: int = 4) -> list[tuple[str, ...]]:
    tokens = roman_tokens(text)
    out: list[tuple[str, ...]] = []
    seen: set[tuple[str, ...]] = set()
    for size in range(min_len, max_len + 1):
        for i in range(0, max(0, len(tokens) - size + 1)):
            gram = tuple(tokens[i:i + size])
            if all(tok in COMMON_ROMAN_TOKENS for tok in gram):
                continue
            if gram not in seen:
                seen.add(gram)
                out.append(gram)
    return out


def gurmukhi_tokens(text: str) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for token in _GURMUKHI_TOKEN_RE.findall(text or ""):
        clean = token.strip("।॥")
        if len(clean) < 2 or clean in seen:
            continue
        seen.add(clean)
        out.append(clean)
    return out


def normalize_gurmukhi_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def shabad_has_mayi(shabad_lines: list[dict]) -> bool:
    """Находит māī/mayi как отдельное слово в roman или gurmukhi."""
    for line in shabad_lines:
        roman_set = set(roman_tokens(best_roman_line(line)))
        if roman_set & MAYI_ROMAN_TOKENS:
            return True
        gurmukhi = line.get("gurmukhi", "")
        if any(form in gurmukhi for form in MAYI_GURMUKHI_FORMS):
            return True
    return False


def shabad_has_sat_sangat(shabad_lines: list[dict]) -> bool:
    markers = {
        "sat", "sach", "sachi", "sacha", "sāchā", "sāchai", "satsangat",
        "sadhsangat", "sādhsaṅgat", "sādh", "sangat", "saṅgat",
    }
    gurmukhi_markers = ("ਸਤਿ", "ਸਚੁ", "ਸਾਚ", "ਸਤਸੰਗ", "ਸਾਧਸੰਗ", "ਸੰਗਤ")
    for line in shabad_lines:
        roman_set = set(roman_tokens(best_roman_line(line)))
        if roman_set & markers:
            return True
        gurmukhi = line.get("gurmukhi", "")
        if any(marker in gurmukhi for marker in gurmukhi_markers):
            return True
    return False


def shabad_has_ghar(shabad_lines: list[dict]) -> bool:
    markers = {"ghar", "ghari", "gharī", "nijghar", "nij", "mahal"}
    gurmukhi_markers = ("ਘਰ", "ਘਰਿ", "ਮਹਲੁ")
    for line in shabad_lines:
        roman_set = set(roman_tokens(best_roman_line(line)))
        if roman_set & markers:
            return True
        gurmukhi = line.get("gurmukhi", "")
        if any(marker in gurmukhi for marker in gurmukhi_markers):
            return True
    return False


def shabad_has_moh(shabad_lines: list[dict]) -> bool:
    """Находит moh/moha как маркер майической привязанности."""
    for line in shabad_lines:
        roman_set = set(roman_tokens(best_roman_line(line)))
        if roman_set & MOH_ROMAN_TOKENS:
            return True
        gurmukhi = line.get("gurmukhi", "")
        if any(form in gurmukhi for form in MOH_GURMUKHI_FORMS):
            return True
    return False


def shabad_has_bhavjal(shabad_lines: list[dict]) -> bool:
    """Находит bhavjal/bhavajal, чтобы раскрывать термин по-русски."""
    for line in shabad_lines:
        roman_set = set(roman_tokens(best_roman_line(line)))
        if roman_set & BHAVJAL_ROMAN_TOKENS:
            return True
        gurmukhi = line.get("gurmukhi", "")
        if any(form in gurmukhi for form in BHAVJAL_GURMUKHI_FORMS):
            return True
    return False


def shabad_has_sohila(shabad_lines: list[dict]) -> bool:
    """Находит sohilā/sohilai, чтобы раскрывать термин по-русски."""
    for line in shabad_lines:
        roman_set = set(roman_tokens(best_roman_line(line)))
        if roman_set & SOHILA_ROMAN_TOKENS:
            return True
        gurmukhi = line.get("gurmukhi", "")
        if any(form in gurmukhi for form in SOHILA_GURMUKHI_FORMS):
            return True
    return False


# ─── knowledge DB lookups ─────────────────────────────────────────────────────

def lookup_word_hints(db_conn: sqlite3.Connection, roman_line: str) -> list[str]:
    """Возвращает подсказки из DB для слов в строке."""
    rows = lookup_word_rows(
        db_conn,
        roman_terms=roman_tokens(roman_line),
        gurmukhi_terms=[],
        roman_ngram_terms=roman_ngrams(roman_line),
        limit=5,
    )
    hints = [
        f"{row['roman']}: [{row['literal_ru']}] → [{row['ksd_meta_ru'][:60]}]"
        for row in rows
        if row.get("ksd_meta_ru")
    ]
    return hints[:5]


def gather_shabad_tokens(shabad_lines: list[dict]) -> tuple[list[str], list[str]]:
    roman_seen: set[str] = set()
    gurmukhi_seen: set[str] = set()
    roman_out: list[str] = []
    gurmukhi_out: list[str] = []

    for line in shabad_lines:
        roman_src = best_roman_line(line)
        for token in roman_tokens(roman_src):
            if token not in roman_seen:
                roman_seen.add(token)
                roman_out.append(token)
        for token in gurmukhi_tokens(line.get("gurmukhi", "")):
            if token not in gurmukhi_seen:
                gurmukhi_seen.add(token)
                gurmukhi_out.append(token)

    return roman_out, gurmukhi_out


def gather_shabad_ngrams(shabad_lines: list[dict]) -> list[tuple[str, ...]]:
    seen: set[tuple[str, ...]] = set()
    out: list[tuple[str, ...]] = []
    for line in shabad_lines:
        for gram in roman_ngrams(best_roman_line(line)):
            if gram not in seen:
                seen.add(gram)
                out.append(gram)
    return out


def word_candidate_score(
    candidate_roman: str,
    *,
    roman_term_set: set[str],
    roman_ngram_set: set[tuple[str, ...]],
    gurmukhi_term_set: set[str],
    row_gurmukhi: str,
    confidence: float,
    has_meta: bool,
    has_grammar: bool,
) -> float:
    tokens = roman_tokens(candidate_roman)
    if not tokens:
        return -1.0

    token_tuple = tuple(tokens)
    overlap = roman_term_set & set(tokens)
    non_common_overlap = {tok for tok in overlap if tok not in COMMON_ROMAN_TOKENS}
    phrase_match = token_tuple in roman_ngram_set
    gurmukhi_match = bool(row_gurmukhi and row_gurmukhi in gurmukhi_term_set)

    if not phrase_match and not gurmukhi_match:
        if len(non_common_overlap) == 0:
            return -1.0
        if len(tokens) == 1 and next(iter(overlap), "") in COMMON_ROMAN_TOKENS:
            return -1.0
        if len(non_common_overlap) < 2:
            important_hit = any(tok in IMPORTANT_ROMAN_TOKENS for tok in non_common_overlap)
            exact_single = len(tokens) == 1 and tokens[0] in IMPORTANT_ROMAN_TOKENS and tokens[0] in overlap
            if not important_hit and not exact_single:
                return -1.0

    score = confidence * 100.0
    score += len(non_common_overlap) * 25.0
    score += len(overlap) * 8.0
    if phrase_match:
        score += 70.0 + len(tokens) * 6.0
    if gurmukhi_match:
        score += 90.0
    if has_meta:
        score += 12.0
    if has_grammar:
        score += 8.0
    if len(tokens) == 1 and tokens[0] in COMMON_ROMAN_TOKENS and not gurmukhi_match:
        score -= 25.0
    return score


def lookup_word_rows(
    db_conn: sqlite3.Connection,
    *,
    roman_terms: list[str],
    gurmukhi_terms: list[str],
    roman_ngram_terms: list[tuple[str, ...]] | None = None,
    limit: int = 14,
) -> list[dict[str, Any]]:
    """Подтягивает релевантные слова/значения из БД по roman и gurmukhi токенам."""
    cur = db_conn.cursor()
    cur.execute(
        "SELECT gurmukhi, roman, literal_ru, ksd_meta_ru, grammar_note, confidence "
        "FROM words"
    )
    candidates = cur.fetchall()

    roman_term_set = set(roman_terms)
    gurmukhi_term_set = set(gurmukhi_terms)
    roman_ngram_set = set(roman_ngram_terms or [])

    scored: list[tuple[float, dict[str, Any]]] = []
    for row in candidates:
        row_gurmukhi = row[0] or ""
        row_roman = row[1] or ""
        row_literal = row[2] or ""
        row_meta = row[3] or ""
        row_grammar = row[4] or ""
        row_conf = row[5] or 0.0

        score = word_candidate_score(
            row_roman,
            roman_term_set=roman_term_set,
            roman_ngram_set=roman_ngram_set,
            gurmukhi_term_set=gurmukhi_term_set,
            row_gurmukhi=row_gurmukhi,
            confidence=row_conf,
            has_meta=bool(row_meta),
            has_grammar=bool(row_grammar),
        )
        if score < 0:
            continue
        if not row_meta and not row_grammar and tuple(roman_tokens(row_roman)) not in roman_ngram_set and row_gurmukhi not in gurmukhi_term_set:
            continue
        scored.append((score, {
            "gurmukhi": row_gurmukhi,
            "roman": row_roman,
            "literal_ru": row_literal,
            "ksd_meta_ru": row_meta,
            "grammar_note": row_grammar,
            "confidence": row_conf,
            "_score": score,
        }))

    scored.sort(key=lambda item: (-item[0], -(item[1]["confidence"] or 0.0), item[1]["roman"]))

    rows: list[dict[str, Any]] = []
    seen_keys: set[tuple[str, str]] = set()
    seen_roman: set[str] = set()
    for _, row in scored:
        key = ("gm", row["gurmukhi"]) if row["gurmukhi"] else ("rm", row["roman"].lower())
        roman_key = " ".join(roman_tokens(row["roman"]))
        if key in seen_keys:
            continue
        if roman_key and roman_key in seen_roman:
            continue
        seen_keys.add(key)
        if roman_key:
            seen_roman.add(roman_key)
        row.pop("_score", None)
        rows.append(row)
        if len(rows) >= limit:
            break

    return rows[:limit]


def find_relevant_concepts(
    db_conn: sqlite3.Connection,
    *,
    roman_terms: list[str],
    gurmukhi_terms: list[str],
    shabad_lines: list[dict],
    limit: int = 6,
) -> list[tuple[str, str]]:
    """Ищет компактный набор KSD-концептов, реально связанных с шабдом."""
    cur = db_conn.cursor()
    cur.execute("SELECT concept, ksd_meaning, source FROM canvas_concepts ORDER BY concept")
    concepts = cur.fetchall()

    line_text = " ".join(
        " ".join(
            [
                best_roman_line(line),
                line.get("gurmukhi", ""),
                line.get("sahib_singh_pa", ""),
            ]
        )
        for line in shabad_lines
    ).lower()
    alias_pool = set(roman_terms) | set(gurmukhi_terms)

    source_rank = {"canvas_ksd": 0, "concepts_sikhi": 1, "canvas_combined_ru": 2}
    hits: list[tuple[int, str, str]] = []
    seen: set[str] = set()
    for concept, meaning, source in concepts:
        aliases = {concept.lower()}
        aliases.update(alias.lower() for alias in CONCEPT_ALIASES.get(concept, ()))
        matched = False
        for alias in aliases:
            alias_norm = normalize_roman_token(alias)
            if alias_norm and alias_norm in alias_pool:
                matched = True
                break
            if alias and alias in gurmukhi_terms:
                matched = True
                break
        if not matched or concept in seen:
            continue
        seen.add(concept)
        rank = source_rank.get(source or "", 99)
        hits.append((rank, concept, (meaning or "").replace("\n", " ")[:220]))

    hits.sort(key=lambda item: (item[0], item[1]))
    return [(concept, meaning) for _, concept, meaning in hits[:limit]]


def find_similar_examples(
    db_conn: sqlite3.Connection,
    *,
    roman_terms: list[str],
    include_user_japji: bool = False,
    limit: int = 3,
) -> list[dict[str, Any]]:
    """Выбирает few-shot примеры с пересечением по ключевым roman словам."""
    if not roman_terms:
        return []

    cur = db_conn.cursor()
    cur.execute(
        "SELECT ang, pauri_num, roman, word_analysis, ksd_translation, source "
        "FROM ksd_examples"
    )

    scored: list[tuple[int, dict[str, Any]]] = []
    roman_set = set(roman_terms)
    for ang, pauri_num, example_roman, word_analysis, ksd_translation, source in cur.fetchall():
        source = source or ""
        if source == "ast_translation_line":
            continue
        if source.startswith("ast_") and not include_user_japji:
            continue
        try:
            wa = json.loads(word_analysis or "[]")
        except Exception:
            wa = []
        example_terms: set[str] = set()
        example_terms.update(roman_tokens(example_roman or ""))
        for item in wa:
            example_terms.update(roman_tokens(item.get("roman", "")))
        matches = roman_set & example_terms
        meaningful_matches = {tok for tok in matches if tok not in COMMON_ROMAN_TOKENS}
        score = len(meaningful_matches) * 2 + len(matches)
        if score <= 0:
            continue
        if not meaningful_matches:
            continue
        scored.append((score, {
            "ang": ang,
            "pauri_num": pauri_num,
            "source": source,
            "ksd_translation": (ksd_translation or "").replace("\n", " ")[:220],
            "matches": sorted(meaningful_matches)[:6],
        }))

    scored.sort(key=lambda item: (-item[0], item[1].get("pauri_num") or 0))
    return [item[1] for item in scored[:limit]]


def find_ast_line_references(
    db_conn: sqlite3.Connection,
    shabad_lines: list[dict],
) -> list[dict[str, str]]:
    """Возвращает готовые AST-переводы для строк Джап Джи, если они есть."""
    cur = db_conn.cursor()
    cur.execute(
        "SELECT gurmukhi, roman, ksd_translation FROM ksd_examples "
        "WHERE source = 'ast_translation_line'"
    )
    rows = cur.fetchall()
    if not rows:
        return []

    by_gurmukhi = {
        normalize_gurmukhi_text(row_gurmukhi): (row_gurmukhi, row_roman, row_translation)
        for row_gurmukhi, row_roman, row_translation in rows
        if normalize_gurmukhi_text(row_gurmukhi)
    }

    refs: list[dict[str, str]] = []
    for line in shabad_lines:
        key = normalize_gurmukhi_text(line.get("gurmukhi", ""))
        if not key or key not in by_gurmukhi:
            continue
        best_row = by_gurmukhi[key]
        refs.append({
            "verse_id": str(line.get("verse_id", "")),
            "gurmukhi": best_row[0] or "",
            "roman": best_row[1] or "",
            "translation": (best_row[2] or "").replace("\n", " "),
        })

    return refs


def build_retrieval_block(
    db_conn: sqlite3.Connection,
    shabad_lines: list[dict],
    *,
    include_user_japji: bool = False,
    ignore_sahib_singh: bool = False,
) -> str:
    """Компактный релевантный контекст вместо попытки засунуть весь корпус в промпт."""
    roman_terms, gurmukhi_terms = gather_shabad_tokens(shabad_lines)
    roman_ngram_terms = gather_shabad_ngrams(shabad_lines)
    word_rows = lookup_word_rows(
        db_conn,
        roman_terms=roman_terms,
        gurmukhi_terms=gurmukhi_terms,
        roman_ngram_terms=roman_ngram_terms,
        limit=16,
    )

    sections: list[str] = []

    ast_refs = find_ast_line_references(db_conn, shabad_lines) if include_user_japji else []
    if ast_refs:
        refs_text = "\n".join(
            f"  • verse_id={ref['verse_id']}\n"
            f"    roman: {ref['roman']}\n"
            f"    готовый русский AST-reference: {ref['translation'][:360]}"
            for ref in ast_refs
        )
        sections.append(
            "━━━ ГОТОВЫЙ ПОСТРОЧНЫЙ ПЕРЕВОД ДЖАП ДЖИ ИЗ AST ━━━\n"
            "Используй это как главный смысловой reference пользователя. Не переводи эти строки с нуля; "
            "сохраняй художественный/духовный вектор AST, нормализуя по KSD, словарю, контексту паури "
            "и запрету латиницы в русских полях.\n"
            + refs_text
        )

    if word_rows:
        words_text = "\n".join(
            f"  • {(row['gurmukhi'] + ' / ') if row['gurmukhi'] else ''}{row['roman']}: "
            f"[{row['literal_ru'][:80]}] → [{row['ksd_meta_ru'][:120]}]"
            for row in word_rows[:14]
        )
        sections.append("━━━ РЕЛЕВАНТНЫЕ СЛОВА ИЗ БАЗЫ ━━━\n" + words_text)

    grammar_notes = []
    seen_grammar: set[str] = set()
    for row in word_rows:
        note = (row.get("grammar_note") or "").strip()
        roman = row.get("roman", "")
        if note and note not in seen_grammar:
            seen_grammar.add(note)
            grammar_notes.append(f"  • {roman}: {note[:180]}")
    if grammar_notes:
        sections.append("━━━ РЕЛЕВАНТНАЯ ГРАММАТИКА ━━━\n" + "\n".join(grammar_notes[:8]))

    concept_rows = find_relevant_concepts(
        db_conn,
        roman_terms=roman_terms,
        gurmukhi_terms=gurmukhi_terms,
        shabad_lines=[] if ignore_sahib_singh else shabad_lines,
        limit=6,
    )
    if concept_rows:
        concepts_text = "\n".join(
            f"  • {concept}: {meaning}" for concept, meaning in concept_rows
        )
        sections.append("━━━ РЕЛЕВАНТНЫЕ KSD-КОНЦЕПТЫ ━━━\n" + concepts_text)

    examples = find_similar_examples(
        db_conn,
        roman_terms=roman_terms,
        include_user_japji=include_user_japji,
        limit=3,
    )
    if examples:
        examples_text = "\n".join(
            f"  • source={ex['source']}, pauri={ex['pauri_num']}, "
            f"match={', '.join(ex['matches'])}: {ex['ksd_translation']}"
            for ex in examples
        )
        sections.append("━━━ ПОХОЖИЕ KSD-ПРИМЕРЫ ━━━\n" + examples_text)

    return ("\n\n".join(sections) + "\n\n") if sections else ""


# ─── промпт для одного шабда ──────────────────────────────────────────────────

def build_shabad_prompt(
    ang_data: dict,
    shabad_lines: list[dict],
    db_conn: sqlite3.Connection,
    *,
    ignore_sahib_singh: bool = False,
) -> str:
    ang         = ang_data["ang"]
    shabad_id   = shabad_lines[0]["shabad_id"]
    rahao_vids  = find_rahao_vids(shabad_lines)
    n           = len(shabad_lines)
    include_user_japji = 1 <= int(ang) <= 8
    retrieval   = build_retrieval_block(
        db_conn,
        shabad_lines,
        include_user_japji=include_user_japji,
        ignore_sahib_singh=ignore_sahib_singh,
    )

    lines_text = ""
    for line in shabad_lines:
        vid      = line["verse_id"]
        if vid in rahao_vids:
            if vid == rahao_vids[-1]:
                rahao_mark = "← РАХАО (последняя строка куплета, с маркером ਰਹਾਉ)"
            else:
                rahao_mark = "← РАХАО-КУПЛЕТ (часть куплета Рахао)"
        else:
            rahao_mark = ""
        roman    = best_roman_line(line)
        roman_note = " [roman битый, сгенерируй заново]" if not roman and line.get("roman") else ""
        ss_pa    = "" if ignore_sahib_singh else line.get("sahib_singh_pa", "")
        hints    = lookup_word_hints(db_conn, roman)
        hint_str = ("  [DB-подсказки: " + " | ".join(hints) + "]") if hints else ""
        ss_line = (
            ""
            if ignore_sahib_singh
            else f"  sahib_singh_low_priority_lexical_hint: {ss_pa[:180] if ss_pa else '—'}\n"
        )
        lines_text += (
            f"verse_id={vid} {rahao_mark}\n"
            f"  gurmukhi:    {line['gurmukhi']}\n"
            f"  roman:       {roman}{roman_note}\n"
            f"{ss_line}"
            f"{hint_str}\n\n"
        )

    japji_rules = f"{JAPJI_CONTEXT_RULES}\n\n" if include_user_japji else ""
    local_hint_parts = []
    if shabad_has_mayi(shabad_lines):
        local_hint_parts.append(MAYI_CONTEXT_HINT)
    if shabad_has_sat_sangat(shabad_lines):
        local_hint_parts.append(SATSANGAT_CONTEXT_HINT)
    if shabad_has_ghar(shabad_lines):
        local_hint_parts.append(GHAR_CONTEXT_HINT)
    if shabad_has_moh(shabad_lines):
        local_hint_parts.append(MOH_CONTEXT_HINT)
    if shabad_has_bhavjal(shabad_lines):
        local_hint_parts.append(BHAVJAL_CONTEXT_HINT)
    if shabad_has_sohila(shabad_lines):
        local_hint_parts.append(SOHILA_CONTEXT_HINT)
    local_hints = ("\n\n".join(local_hint_parts) + "\n\n") if local_hint_parts else ""

    return (
        f"Игнорируй весь предыдущий контекст этого чата.\n\n"
        f"АНГ {ang}, ШАБД {shabad_id} — {n} строк. Рахао verse_ids: {rahao_vids or 'не найдено'}.\n\n"
        f"{'━━━ РЕЖИМ БЕЗ SAHIB SINGH ━━━\nSahib Singh полностью исключён из prompt. Переводи по gurmukhi, roman, грамматике, KSD-концептам, AST-reference и контексту произведения.\n\n' if ignore_sahib_singh else ''}"
        f"{japji_rules}"
        f"{local_hints}"
        f"{retrieval}"
        f"━━━ ИСХОДНЫЕ СТРОКИ ШАБДА ━━━\n\n"
        f"{lines_text}"
        f"Верни KSD-анализ и перевод ВСЕХ {n} строк в указанном порядке.\n"
        f"Порядок строк НЕ менять. Все {n} строк в массиве lines.\n"
        f"Ответ — строго между BEGIN_KSD_JSON и END_KSD_JSON."
    )


REPAIR_PROMPT = """\
Игнорируй весь предыдущий контекст этого чата, кроме своего последнего ответа.
Перепиши свой последний ответ в правильном JSON-формате между BEGIN_KSD_JSON и END_KSD_JSON.
Требования:
- ровно {n} элементов в массиве lines
- порядок verse_id сохранить
- никаких ASCII-кавычек " " внутри строк — только «ёлочки» или одинарные
- никаких markdown code fences
- никаких латинских букв в `ksd_translation`, `artistic_ru`, `context_note`, `shabad_summary`, `rahao_theme`, `confidence_reason`
- если нужен термин Гурбани, передавай его кириллицей, а не латиницей
Формат тот же что был запрошен изначально.
"""

MAX_SHABAD_ATTEMPTS = 3


def collect_latin_leaks(result: dict[str, Any]) -> list[str]:
    """Находит латиницу в русских полях, где она недопустима."""
    leaks: list[str] = []

    for field in ["rahao_theme", "shabad_summary"]:
        text = result.get(field, "")
        if isinstance(text, str) and _LATIN_LEAK_RE.search(text):
            leaks.append(f"{field}: {text[:120]}")

    for line in result.get("lines", []):
        vid = line.get("verse_id", "?")
        for field in ["ksd_translation", "artistic_ru", "context_note", "confidence_reason"]:
            text = line.get(field, "")
            if isinstance(text, str) and _LATIN_LEAK_RE.search(text):
                leaks.append(f"verse_id={vid} {field}: {text[:120]}")

    return leaks[:12]


def build_latin_repair_prompt(n: int, leaks: list[str]) -> str:
    leak_text = "\n".join(f"- {item}" for item in leaks[:12])
    return (
        REPAIR_PROMPT.format(n=n)
        + "\nИсправь конкретно эти места с латиницей:\n"
        + leak_text
        + "\nСохрани смысл, но передай все такие слова по-русски или кириллицей."
    )


def validate_shabad_result(result: dict[str, Any] | None, expected_n: int) -> str | None:
    if result is None:
        return "нет JSON-результата"
    lines = result.get("lines")
    if not isinstance(lines, list):
        return "нет массива lines"
    if len(lines) != expected_n:
        return f"ожидалось {expected_n} строк, получено {len(lines)}"
    return None

# ─── JSON парсинг ─────────────────────────────────────────────────────────────

def repair_json_quotes(text: str) -> str:
    """Экранирует незакрытые ASCII-кавычки внутри строк JSON."""
    result = []
    in_string = False
    escape_next = False
    for c in text:
        if escape_next:
            result.append(c)
            escape_next = False
            continue
        if c == "\\":
            result.append(c)
            escape_next = True
            continue
        if c == '"':
            if not in_string:
                in_string = True
                result.append(c)
            else:
                # peek ahead
                tail = "".join(result[result.index('"'):] if '"' in result else []) + c
                # упрощённо: проверяем что следует за кавычкой в оригинале
                result.append(c)
                in_string = False
            continue
        result.append(c)
    return "".join(result)


def extract_ksd_json(text: str) -> dict | None:
    m = re.search(r"BEGIN_KSD_JSON\s*(.*?)\s*END_KSD_JSON", text, re.DOTALL)
    if not m:
        return None
    raw = m.group(1).strip()
    raw = re.sub(r"^```json\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"^```\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    # Заменяем типографские кавычки на обычные в ключах
    raw = raw.replace('\u201c', '\u00ab').replace('\u201d', '\u00bb')
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        fixed = repair_json_quotes(raw)
        try:
            return json.loads(fixed)
        except Exception:
            return None


# ─── Playwright interactions ──────────────────────────────────────────────────

def open_chat_tab(context, page_timeout_ms: int = 30_000) -> Page:
    page = context.new_page()
    if _stealth is not None:
        _stealth.apply_stealth_sync(page)
    page.goto(CHAT_URL, wait_until="domcontentloaded", timeout=page_timeout_ms)
    page.wait_for_timeout(2000)
    return page


def assistant_msg_count(page: Page) -> int:
    return page.locator('[data-message-author-role="assistant"]').count()


def insert_text(page: Page, text: str) -> None:
    page.evaluate(
        """(t) => {
            const el = document.getElementById('prompt-textarea');
            if (!el) return;
            el.focus();
            document.execCommand('selectAll', false, null);
            document.execCommand('insertText', false, t);
        }""",
        text,
    )
    page.wait_for_timeout(500)


def click_send(page: Page) -> None:
    btn = page.locator('button[data-testid="send-button"]')
    try:
        btn.wait_for(state="visible", timeout=5_000)
        btn.click()
    except PWTimeout:
        page.locator("#prompt-textarea").press("Enter")


def drain_generation(page: Page, timeout_ms: int = 900_000) -> None:
    stop = page.locator('button[data-testid="stop-button"]')
    try:
        stop.wait_for(state="visible", timeout=20_000)
    except PWTimeout:
        pass
    try:
        stop.wait_for(state="hidden", timeout=timeout_ms)
    except PWTimeout:
        print("    ⚠ Таймаут ожидания ответа — беру что есть")
    page.wait_for_timeout(1200)
    # Continue generating если надо
    for _ in range(4):
        cont = page.locator("button").filter(
            has_text=re.compile(r"Continue generating|Продолжить", re.I)
        )
        if cont.count() > 0 and cont.first.is_visible():
            cont.first.click()
            page.wait_for_timeout(600)
            drain_generation(page, timeout_ms)
            break


def get_last_assistant_msg(page: Page, before_count: int, wait_ms: int = 30_000) -> str | None:
    try:
        page.wait_for_function(
            "(b) => document.querySelectorAll('[data-message-author-role=\"assistant\"]').length > b",
            arg=before_count,
            timeout=wait_ms,
        )
    except PWTimeout:
        pass
    msgs = page.locator('[data-message-author-role="assistant"]')
    n = msgs.count()
    if n == 0 or n <= before_count:
        return None
    text = msgs.nth(n - 1).inner_text()
    return text.strip() or None


def get_last_assistant_any(page: Page) -> str | None:
    msgs = page.locator('[data-message-author-role="assistant"]')
    n = msgs.count()
    if n == 0:
        return None
    try:
        text = msgs.nth(n - 1).inner_text(timeout=10_000)
    except Exception:
        return None
    return text.strip() or None


def write_chatgpt_debug_dump(
    page: Page,
    *,
    prompt: str,
    raw_result: str | None,
    label: str,
) -> None:
    stamp = time.strftime("%Y%m%d_%H%M%S")
    safe_label = re.sub(r"[^A-Za-z0-9_.-]+", "_", label).strip("_")
    base = LOG_DIR / f"chatgpt_{safe_label}_{stamp}"
    try:
        (base.with_suffix(".prompt.txt")).write_text(prompt, encoding="utf-8")
    except Exception:
        pass
    try:
        last_any = get_last_assistant_any(page) or ""
        (base.with_suffix(".last_assistant.txt")).write_text(last_any, encoding="utf-8")
    except Exception:
        pass
    try:
        (base.with_suffix(".result.txt")).write_text(raw_result or "", encoding="utf-8")
    except Exception:
        pass
    try:
        page.screenshot(path=str(base.with_suffix(".png")), full_page=True)
    except Exception:
        pass
    print(f"    [DEBUG] dump: {base}.*")


def send_and_get(
    page: Page,
    prompt: str,
    response_ms: int = 900_000,
    debug_label: str | None = None,
) -> str | None:
    ta = page.locator("#prompt-textarea")
    try:
        ta.wait_for(state="visible", timeout=15_000)
    except PWTimeout:
        print("    ✗ Поле ввода не найдено")
        return None
    ta.click()
    page.keyboard.press("Control+a")
    page.keyboard.press("Delete")
    page.wait_for_timeout(250)
    insert_text(page, prompt)
    before = assistant_msg_count(page)
    click_send(page)
    print("    → Ждём ответа ChatGPT…", end=" ", flush=True)
    drain_generation(page, response_ms)
    result = get_last_assistant_msg(page, before)
    if result is None:
        fallback = get_last_assistant_any(page)
        if fallback:
            print("найден последний assistant msg без счётчика")
            result = fallback
        elif debug_label:
            print("нет ответа")
            write_chatgpt_debug_dump(
                page,
                prompt=prompt,
                raw_result=result,
                label=debug_label,
            )
            return None
    else:
        print("получен")
        return result
    print("получен" if result else "нет ответа")
    return result


# ─── сохранение KSD-JSON ──────────────────────────────────────────────────────

def save_ksd_ang_json(ang: int, shabad_result: dict) -> None:
    """Дополняет/создаёт ksd_ang_XXXX.json. Исходный ang_json НЕ трогается."""
    path = KSD_JSON_DIR / f"ksd_ang_{ang:04d}.json"
    if path.exists():
        stamp = time.strftime("%Y%m%d_%H%M%S")
        backup_path = KSD_BACKUP_DIR / f"ksd_ang_{ang:04d}.{stamp}.bak.json"
        shutil.copy2(path, backup_path)
        with open(path, encoding="utf-8") as f:
            existing = json.load(f)
    else:
        existing = {"ang": ang, "shabads": []}
    # Убираем старый вариант шабда если есть
    shabad_id = shabad_result.get("shabad_id")
    existing["shabads"] = [
        s for s in existing.get("shabads", [])
        if s.get("shabad_id") != shabad_id
    ]
    existing["shabads"].append(shabad_result)
    existing["shabads"].sort(
        key=lambda s: (
            s.get("lines", [{}])[0].get("verse_id", 10**9)
            if s.get("lines")
            else 10**9,
            s.get("shabad_id", 10**9),
        )
    )
    with open(path, "w", encoding="utf-8") as f:
        json.dump(existing, f, ensure_ascii=False, indent=2)


# ─── docx builder ─────────────────────────────────────────────────────────────

def add_run(para, text: str, *, color=None, bold=False, italic=False, size_pt=None):
    if Pt is None:
        raise RuntimeError("Для сборки docx нужен пакет python-docx")
    run = para.add_run(text)
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def add_translation_with_blue(
    para,
    text: str,
    *,
    base_color: RGBColor,
    bold: bool = False,
    italic: bool = False,
    size_pt: int = 11,
) -> None:
    """
    Рендерит строку перевода с поддержкой [[синих вставок]].
    Текст внутри [[...]] выводится синим цветом в квадратных скобках [].
    Весь остальной текст — цветом base_color.

    Пример входа:  "[[Хукам]] — сила, движущая [[нашим существованием]]"
    Результат:     "[Хукам]" синим + " — сила, движущая " base_color + "[нашим существованием]" синим
    """
    # Разбиваем по маркерам [[...]]
    parts = re.split(r"(\[\[.*?\]\])", text)
    for part in parts:
        if not part:
            continue
        m = re.match(r"\[\[(.*?)\]\]", part)
        if m:
            # Синяя вставка: показываем в [скобках] синим
            add_run(
                para,
                f"[{m.group(1)}]",
                color=COLOR_CONTEXT,
                bold=bold,
                italic=italic,
                size_pt=size_pt,
            )
        else:
            add_run(
                para,
                part,
                color=base_color,
                bold=bold,
                italic=italic,
                size_pt=size_pt,
            )


def ensure_docx(path: Path) -> None:
    if Document is None or Inches is None or WD_ALIGN_PARAGRAPH is None or Pt is None:
        raise RuntimeError("Для сборки docx нужен пакет python-docx")
    if path.exists():
        return
    doc = Document()
    doc.core_properties.title = "СГГС — KSD-перевод"
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
    title = doc.add_heading("Шри Гуру Грантх Сахиб — KSD-перевод", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sub, "Методология: Karminder Singh Dhillon",
            color=COLOR_COMMENT, italic=True, size_pt=9)
    doc.add_paragraph()
    doc.save(str(path))


def append_shabad_to_docx(path: Path, result: dict) -> None:
    ensure_docx(path)
    doc = Document(str(path))

    ang       = result.get("ang", "?")
    shabad_id = result.get("shabad_id", "?")
    theme     = result.get("rahao_theme", "")

    # Заголовок анга/шабда
    p_head = doc.add_paragraph()
    p_head.paragraph_format.space_before = Pt(14)
    add_run(p_head, f"Анг {ang}", bold=True, color=COLOR_HEADING, size_pt=11)

    if theme:
        p_theme = doc.add_paragraph()
        add_run(p_theme, f"↳ {theme}", italic=True, color=COLOR_RAHAO, size_pt=9)

    for line in result.get("lines", []):
        is_rahao = line.get("is_rahao", False)

        # Гурмукхи
        p_gm = doc.add_paragraph()
        p_gm.paragraph_format.space_before = Pt(7)
        p_gm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p_gm, line.get("gurmukhi", ""),
                color=COLOR_GURMUKHI, bold=is_rahao, size_pt=14)

        # Романизация
        roman = line.get("roman", "")
        roman_display = line.get("roman_display") or roman_display_from_gurmukhi(
            line.get("gurmukhi", ""),
            roman,
        )
        if roman_display:
            p_rm = doc.add_paragraph()
            p_rm.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p_rm, roman_display, color=COLOR_ROMAN, italic=True, size_pt=9)

        # KSD-перевод (с поддержкой [[синих вставок]])
        ksd = line.get("ksd_translation", "")
        if ksd:
            p_tr = doc.add_paragraph()
            p_tr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_translation_with_blue(
                p_tr, ksd,
                base_color=COLOR_TRANSLATION,
                bold=is_rahao,
                size_pt=11,
            )

        # Художественный вариант (только если отличается и не пустой)
        artistic = line.get("artistic_ru", "")
        ksd_clean = re.sub(r"\[\[.*?\]\]", "", ksd).strip()
        if artistic and artistic.strip() and artistic.strip() != ksd_clean:
            p_art = doc.add_paragraph()
            p_art.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_translation_with_blue(
                p_art, f"〜 {artistic}",
                base_color=COLOR_ARTISTIC,
                italic=True,
                size_pt=10,
            )

        # Синий блок — context_note
        ctx = line.get("context_note", "")
        if ctx and ctx.strip():
            p_ctx = doc.add_paragraph()
            add_run(p_ctx, ctx, color=COLOR_CONTEXT, size_pt=9)

    # Резюме шабда
    summary = result.get("shabad_summary", "")
    if summary:
        p_sum = doc.add_paragraph()
        p_sum.paragraph_format.space_before = Pt(4)
        add_run(p_sum, f"[ {summary} ]", color=COLOR_COMMENT, italic=True, size_pt=9)

    doc.add_paragraph("─" * 45)
    doc.save(str(path))


def rebuild_docx_from_ksd_json(ang_start: int, ang_end: int) -> None:
    """Пересобирает docx из готовых ksd_ang_JSON файлов."""
    if OUTPUT_DOCX.exists():
        OUTPUT_DOCX.unlink()
    ensure_docx(OUTPUT_DOCX)
    rebuilt = 0
    for ang in range(ang_start, ang_end + 1):
        path = KSD_JSON_DIR / f"ksd_ang_{ang:04d}.json"
        if not path.exists():
            continue
        ang_data = load_ang(ang)
        orig_by_vid = {
            line["verse_id"]: line
            for line in (ang_data or {}).get("lines", [])
        }
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        shabads = sorted(
            data.get("shabads", []),
            key=lambda s: (
                s.get("lines", [{}])[0].get("verse_id", 10**9)
                if s.get("lines")
                else 10**9,
                s.get("shabad_id", 10**9),
            ),
        )
        for shabad in shabads:
            # Перезапускаем детекцию Рахао-куплета из гурмукхи-поля
            lines = shabad.get("lines", [])
            rahao_vids = find_rahao_vids(lines)
            for ln in lines:
                vid = ln.get("verse_id")
                orig = orig_by_vid.get(vid)
                if orig:
                    if not ln.get("gurmukhi"):
                        ln["gurmukhi"] = orig.get("gurmukhi", "")
                    if not ln.get("roman"):
                        ln["roman"] = best_roman_line(orig)
                ln["roman_display"] = roman_display_from_gurmukhi(
                    ln.get("gurmukhi", ""),
                    ln.get("roman", ""),
                )
                ln["is_rahao"] = ln.get("verse_id") in rahao_vids
            append_shabad_to_docx(OUTPUT_DOCX, shabad)
            rebuilt += 1
    print(f"Пересобрано: {rebuilt} шабдов → {OUTPUT_DOCX}")


# ─── progress ─────────────────────────────────────────────────────────────────

def load_progress() -> set[str]:
    if not PROGRESS_FILE.exists():
        return set()
    return set(PROGRESS_FILE.read_text(encoding="utf-8").splitlines())


def load_saved_json_keys() -> set[str]:
    """Считает готовыми шабды, уже лежащие в ksd_ang_json."""
    done: set[str] = set()
    for path in KSD_JSON_DIR.glob("ksd_ang_*.json"):
        m = re.search(r"ksd_ang_(\d+)\.json$", path.name)
        if not m:
            continue
        ang = int(m.group(1))
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        for shabad in data.get("shabads", []):
            shabad_id = shabad.get("shabad_id")
            if shabad_id is not None:
                done.add(f"{ang}:{shabad_id}")
    return done


def load_resume_keys() -> set[str]:
    return load_progress() | load_saved_json_keys()


def save_progress(done: set[str]) -> None:
    PROGRESS_FILE.write_text("\n".join(sorted(done)), encoding="utf-8")


def source_shabad_ids_for_ang(ang: int) -> list[int]:
    ang_data = load_ang(ang)
    if not ang_data:
        return []
    return [group[0]["shabad_id"] for group in group_by_shabad(ang_data) if group]


def available_source_angs() -> list[int]:
    out = []
    for path in ANG_JSON_DIR.glob("ang_*.json"):
        m = re.search(r"ang_(\d+)\.json$", path.name)
        if m:
            out.append(int(m.group(1)))
    return sorted(out)


def next_incomplete_position(done_keys: set[str]) -> tuple[int, int | None] | None:
    for ang in available_source_angs():
        ids = source_shabad_ids_for_ang(ang)
        if not ids:
            continue
        for shabad_id in ids:
            if f"{ang}:{shabad_id}" not in done_keys:
                return ang, shabad_id
    return None


def count_completed_for_ang(ang: int, done_keys: set[str]) -> tuple[int, int]:
    ids = source_shabad_ids_for_ang(ang)
    return sum(1 for shabad_id in ids if f"{ang}:{shabad_id}" in done_keys), len(ids)


def parse_ang_range(text: str) -> list[int]:
    text = text.strip()
    if "-" in text:
        a, b = text.split("-", 1)
        return list(range(int(a), int(b) + 1))
    return [int(text)]


def save_last_attempt(
    ang_list: list[int],
    *,
    resume: bool,
    test: bool,
    shabad_id_filter: int | None,
) -> None:
    if test or not ang_list:
        return
    payload = {
        "ang_start": min(ang_list),
        "ang_end": max(ang_list),
        "angs": ang_list,
        "resume": resume,
        "shabad_id_filter": shabad_id_filter,
        "started_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    LAST_ATTEMPT_FILE.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def load_last_attempt() -> dict[str, Any] | None:
    if not LAST_ATTEMPT_FILE.exists():
        return None
    try:
        return json.loads(LAST_ATTEMPT_FILE.read_text(encoding="utf-8"))
    except Exception:
        return None


def run_interactive_menu() -> None:
    progress_keys = load_progress()
    saved_keys = load_saved_json_keys()
    done_keys = progress_keys | saved_keys
    source_angs = available_source_angs()
    existing_angs = sorted(
        int(m.group(1))
        for p in KSD_JSON_DIR.glob("ksd_ang_*.json")
        if (m := re.search(r"ksd_ang_(\d+)\.json$", p.name))
    )
    next_pos = next_incomplete_position(done_keys)
    next_ang = next_pos[0] if next_pos else None
    next_shabad = next_pos[1] if next_pos else None
    last_attempt = load_last_attempt()
    last_start = int(last_attempt.get("ang_start", 0)) if last_attempt else 0
    last_end = int(last_attempt.get("ang_end", 0)) if last_attempt else 0
    if next_ang and last_start <= next_ang <= last_end:
        default_end = last_end
    else:
        default_end = min((next_ang or 1) + 11, max(source_angs or [1]))

    print("\n══ KSD_AI Translator ══\n")
    print(f"  Progress file:          {PROGRESS_FILE}")
    print(f"  Шабдов в progress:      {len(progress_keys)}")
    print(f"  Шабдов в JSON:          {len(saved_keys)}")
    print(f"  Готовых ключей всего:   {len(done_keys)}")
    print(f"  KSD JSON ангов:         {len(existing_angs)}" + (f" ({existing_angs[0]}..{existing_angs[-1]})" if existing_angs else ""))
    if last_attempt:
        print(f"  Последняя попытка:      анги {last_start}..{last_end} ({last_attempt.get('started_at', '?')})")
    if next_pos:
        done, total = count_completed_for_ang(next_ang, done_keys)
        print(f"  Следующий незавершённый: анг {next_ang}, шабд {next_shabad} ({done}/{total} готово)")
    else:
        print("  Прогресс: всё доступное выглядит готовым")

    print("\n  Что делаем?")
    if next_pos:
        print(
            f"  1) Продолжить предыдущую попытку → с анг {next_ang}, "
            f"шабд {next_shabad}; до анга {default_end} (--resume)"
        )
    print("  2) Свой диапазон ангов")
    print("  3) Один шабд")
    print("  4) Пересобрать DOCX из JSON")
    print("  5) Показать prompt")
    print("  6) Показать прогресс по диапазону")
    print("  0) Выход\n")

    choice = input("  Выбор: ").strip()
    print()

    if choice == "0":
        print("Выход.")
        return

    if choice == "1" and next_pos:
        print(
            f"Запускаю с первого незавершённого: анг {next_ang}, "
            f"шабд {next_shabad}; диапазон {next_ang}..{default_end} с --resume\n"
        )
        process_ang_list(list(range(next_ang, default_end + 1)), resume=True)
        return

    if choice == "2":
        rng = input("  Диапазон, например 2-13 или 5: ").strip()
        try:
            ang_list = parse_ang_range(rng)
        except ValueError:
            print("  ✗ Некорректный диапазон")
            return
        use_resume = input("  Учитывать progress/JSON и пропускать готовые? [Y/n] ").strip().lower() != "n"
        process_ang_list(ang_list, resume=use_resume)
        return

    if choice == "3":
        try:
            ang = int(input("  Анг: ").strip())
            shabad_id = int(input("  Shabad ID: ").strip())
        except ValueError:
            print("  ✗ Некорректный ввод")
            return
        process_ang_list([ang], resume=False, shabad_id_filter=shabad_id)
        return

    if choice == "4":
        rng = input("  Диапазон rebuild [1-13]: ").strip() or "1-13"
        try:
            ang_list = parse_ang_range(rng)
        except ValueError:
            print("  ✗ Некорректный диапазон")
            return
        rebuild_docx_from_ksd_json(min(ang_list), max(ang_list))
        return

    if choice == "5":
        rng = input("  Диапазон prompt [1]: ").strip() or "1"
        shabad_s = input("  Shabad ID [пусто = все/первые по режиму]: ").strip()
        try:
            ang_list = parse_ang_range(rng)
            shabad_id = int(shabad_s) if shabad_s else None
        except ValueError:
            print("  ✗ Некорректный ввод")
            return
        print_prompt_bundle(ang_list, shabad_id=shabad_id)
        return

    if choice == "6":
        rng = input("  Диапазон проверки [1-13]: ").strip() or "1-13"
        try:
            ang_list = parse_ang_range(rng)
        except ValueError:
            print("  ✗ Некорректный диапазон")
            return
        for ang in ang_list:
            done, total = count_completed_for_ang(ang, done_keys)
            status = "✓" if total and done == total else ("-" if total else "?")
            print(f"  {status} Анг {ang}: {done}/{total}")
        return

    print("  ✗ Неизвестный выбор")


def print_prompt_bundle(
    ang_list: list[int],
    shabad_id: int | None = None,
    *,
    ignore_sahib_singh: bool = False,
) -> None:
    """Печатает system-блок и prompt для локальной проверки без Playwright."""
    db_conn = sqlite3.connect(str(DB_PATH))
    system_block = build_system_block(db_conn)
    print("===== SYSTEM BLOCK =====\n")
    print(system_block)

    for ang_num in ang_list:
        ang_data = load_ang(ang_num)
        if not ang_data:
            print(f"\n[SKIP] ang {ang_num} — файл не найден")
            continue
        for shabad_lines in group_by_shabad(ang_data):
            current_shabad_id = shabad_lines[0]["shabad_id"]
            if shabad_id is not None and current_shabad_id != shabad_id:
                continue
            print(f"\n===== PROMPT ang={ang_num} shabad={current_shabad_id} =====\n")
            print(build_shabad_prompt(
                ang_data,
                shabad_lines,
                db_conn,
                ignore_sahib_singh=ignore_sahib_singh,
            ))

    db_conn.close()


# ─── основной цикл ────────────────────────────────────────────────────────────

def process_ang_list(
    ang_list: list[int],
    resume: bool = False,
    test: bool = False,
    shabad_id_filter: int | None = None,
    ignore_sahib_singh: bool = False,
) -> None:
    if sync_playwright is None:
        raise RuntimeError("Для запуска перевода нужен playwright")
    save_last_attempt(
        ang_list,
        resume=resume,
        test=test,
        shabad_id_filter=shabad_id_filter,
    )
    db_conn = sqlite3.connect(str(DB_PATH))
    system_block = build_system_block(db_conn)

    done_keys = load_resume_keys() if resume else set()

    # В тестовом режиме — отдельный docx, прогресс не сохраняется
    out_docx = OUT_DIR / "SGGS_KSD_TEST.docx" if test else OUTPUT_DOCX
    if test and out_docx.exists():
        out_docx.unlink()
    ensure_docx(out_docx)

    with sync_playwright() as pw:
        context = pw.chromium.launch_persistent_context(
            str(BOT_PROFILE),
            headless=False,
            args=["--disable-blink-features=AutomationControlled"],
        )
        page = open_chat_tab(context)

        # Отправляем system-блок один раз в начале чата
        print("Инициализация чата (system-блок)…")
        init_resp = send_and_get(page, system_block + "\n\nПодтверди: ты готов переводить шабды СГГС по методологии KSD.")
        if init_resp:
            print(f"  ChatGPT: {init_resp[:120]}…")
        else:
            print("  [WARN] нет подтверждения, продолжаем всё равно")

        for ang_num in ang_list:
            ang_data = load_ang(ang_num)
            if not ang_data:
                print(f"  [SKIP] ang {ang_num} — файл не найден")
                continue

            shabads = group_by_shabad(ang_data)
            if shabad_id_filter is not None:
                shabads = [s for s in shabads if s and s[0]["shabad_id"] == shabad_id_filter]
            if test:
                shabads = shabads[:1]  # только первый шабд
                print(f"\nАнг {ang_num}: тест — первый шабд из {len(group_by_shabad(ang_data))}")
            else:
                print(f"\nАнг {ang_num}: {len(shabads)} шабдов")

            for shabad_lines in shabads:
                shabad_id = shabad_lines[0]["shabad_id"]
                key = f"{ang_num}:{shabad_id}"

                if key in done_keys:
                    print(f"  → Шабд {shabad_id} [пропуск — уже готов]")
                    continue

                print(f"  → Шабд {shabad_id} ({len(shabad_lines)} строк)…")

                result = None
                last_error = "нет попыток"
                for attempt in range(1, MAX_SHABAD_ATTEMPTS + 1):
                    if attempt > 1:
                        print(f"    [RETRY] попытка {attempt}/{MAX_SHABAD_ATTEMPTS}…")
                    prompt  = build_shabad_prompt(
                        ang_data,
                        shabad_lines,
                        db_conn,
                        ignore_sahib_singh=ignore_sahib_singh,
                    )
                    debug_label = f"ang{ang_num:04d}_shabad{shabad_id}_attempt{attempt}"
                    raw_ans = send_and_get(page, prompt, debug_label=debug_label)

                    if not raw_ans:
                        last_error = "ChatGPT не вернул ответ"
                        continue

                    result = extract_ksd_json(raw_ans)
                    validation_error = validate_shabad_result(result, len(shabad_lines))
                    if validation_error:
                        last_error = validation_error
                        print(f"    [REPAIR] {validation_error}, пробую repair…")
                        repair_ans = send_and_get(
                            page,
                            REPAIR_PROMPT.format(n=len(shabad_lines)),
                            debug_label=f"{debug_label}_repair_json",
                        )
                        if repair_ans:
                            result = extract_ksd_json(repair_ans)
                            validation_error = validate_shabad_result(result, len(shabad_lines))
                        if validation_error:
                            last_error = validation_error
                            result = None
                            continue

                    leaks = collect_latin_leaks(result)
                    if leaks:
                        print("    [REPAIR] найдена латиница в русских полях, пробую repair…")
                        repair_ans = send_and_get(
                            page,
                            build_latin_repair_prompt(len(shabad_lines), leaks),
                            debug_label=f"{debug_label}_repair_latin",
                        )
                        if repair_ans:
                            repaired = extract_ksd_json(repair_ans)
                            validation_error = validate_shabad_result(repaired, len(shabad_lines))
                            if not validation_error:
                                repaired_leaks = collect_latin_leaks(repaired)
                                if not repaired_leaks:
                                    result = repaired
                                else:
                                    last_error = "после repair латиница всё ещё осталась"
                                    print("    [WARN] после repair латиница всё ещё осталась")
                                    result = None
                                    continue
                            else:
                                last_error = validation_error
                                result = None
                                continue
                        else:
                            last_error = "repair латиницы не вернул ответ"
                            result = None
                            continue

                    break

                if result:
                    # Обогащаем из оригинала — оригинальный ang_json не трогаем
                    rahao_vids = find_rahao_vids(shabad_lines)
                    orig_by_vid = {l["verse_id"]: l for l in shabad_lines}
                    for ln in result.get("lines", []):
                        vid = ln.get("verse_id")
                        if vid and vid in orig_by_vid:
                            orig = orig_by_vid[vid]
                            if not ln.get("gurmukhi"):
                                ln["gurmukhi"] = orig.get("gurmukhi", "")
                            # roman: берём из результата если чистый, иначе из оригинала
                            if not ln.get("roman"):
                                ln["roman"] = best_roman_line(orig)
                            ln["sahib_singh_pa"] = orig.get("sahib_singh_pa", "")
                        ln["roman_display"] = roman_display_from_gurmukhi(
                            ln.get("gurmukhi", ""),
                            ln.get("roman", ""),
                        )
                        # Выставляем is_rahao для всего куплета
                        if vid in rahao_vids:
                            ln["is_rahao"] = True
                    result["ang"]      = ang_num
                    result["shabad_id"] = shabad_id

                    # Считаем confidence
                    lines_conf = result.get("lines", [])
                    avg_conf = (
                        sum(l.get("confidence", 0.8) for l in lines_conf) / len(lines_conf)
                        if lines_conf else 0.8
                    )
                    badge = "✓✓ HIGH" if avg_conf >= 0.85 else ("✓ MED" if avg_conf >= 0.65 else "? LOW")
                    print(f"    OK [{badge} {avg_conf:.2f}]")

                    # Сохраняем в ksd_ang_json (наши данные, не оригинал)
                    if not test:
                        save_ksd_ang_json(ang_num, result)
                        done_keys.add(key)
                        save_progress(done_keys)
                    append_shabad_to_docx(out_docx, result)
                else:
                    raise RuntimeError(
                        f"Шабд {shabad_id} ang {ang_num} не переведён после "
                        f"{MAX_SHABAD_ATTEMPTS} попыток: {last_error}. "
                        "Прогон остановлен, чтобы не получить дырявый результат."
                    )

                time.sleep(1)

        context.close()

    db_conn.close()
    print(f"\nГотово. DOCX: {OUTPUT_DOCX}")


# ─── CLI ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="AI KSD Translator (ChatGPT + Playwright)")
    parser.add_argument("--ang",          help="Анг или диапазон: 1 или 1-8")
    parser.add_argument("--all",          action="store_true", help="Все доступные анги")
    parser.add_argument("--resume",       action="store_true", help="Продолжить с остановки")
    parser.add_argument("--menu",         action="store_true", help="Интерактивное меню")
    parser.add_argument("--test",         action="store_true", help="Тест: только первый шабд анга, без сохранения прогресса")
    parser.add_argument("--shabad-id",    type=int, help="Ограничить обработку одним shabad_id")
    parser.add_argument("--print-prompt", action="store_true", help="Напечатать system-блок и prompt без запуска браузера")
    parser.add_argument("--ignore-sahib-singh", action="store_true", help="Не включать объяснения Sahib Singh в prompt")
    parser.add_argument("--rebuild-docx", action="store_true", help="Пересобрать docx из JSON")
    parser.add_argument("--rebuild-range", default="1-1430",   help="Диапазон для rebuild-docx")
    args = parser.parse_args()

    if args.menu or len(sys.argv) == 1:
        run_interactive_menu()
        return

    if args.rebuild_docx:
        a, b = args.rebuild_range.split("-")
        rebuild_docx_from_ksd_json(int(a), int(b))
        return

    # Определяем список ангов
    if args.ang:
        if "-" in str(args.ang):
            a, b = str(args.ang).split("-")
            ang_list = list(range(int(a), int(b) + 1))
        else:
            ang_list = [int(args.ang)]
    elif args.all or args.resume:
        ang_list = sorted(
            int(p.stem.replace("ang_", ""))
            for p in ANG_JSON_DIR.glob("ang_*.json")
        )
    else:
        parser.print_help()
        sys.exit(1)

    if args.print_prompt:
        print_prompt_bundle(
            ang_list,
            shabad_id=args.shabad_id,
            ignore_sahib_singh=args.ignore_sahib_singh,
        )
        return

    process_ang_list(
        ang_list,
        resume=args.resume,
        test=args.test,
        shabad_id_filter=args.shabad_id,
        ignore_sahib_singh=args.ignore_sahib_singh,
    )


if __name__ == "__main__":
    main()
