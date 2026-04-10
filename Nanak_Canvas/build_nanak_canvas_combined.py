#!/usr/bin/env python3
"""
Combine all 12 parts of Nanak Canvas into a single DOCX.

Features:
- Table of Contents with internal hyperlinks (bookmarks)
- Proper image merging (copies binary image parts, remaps relationship IDs)
- Post-pass: normalise Gurmukhi, transliteration and body-text formatting
- Page break between every part
- Nothing is dropped: every paragraph and image from each source file is included
"""

import copy, hashlib, os, re, sys
from pathlib import Path

from docx import Document
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Part metadata ─────────────────────────────────────────────────────────────

PARTS = [
    (1,  'Холст Гуру Нанака',                    'Nanak_Canvas_Part1_RU.docx'),
    (2,  'Смерть',                                'Nanak_Canvas_Part2_RU.docx'),
    (3,  'Жизнь После Смерти',                   'Nanak_Canvas_Part3_RU.docx'),
    (4,  'Чаураси Лакх — 8,4 Миллиона Форм Жизни', 'Nanak_Canvas_Part4_RU.docx'),
    (5,  'Реинкарнация — Ава Гаун',              'Nanak_Canvas_Part5_RU.docx'),
    (6,  'Рай и Ад',                              'Nanak_Canvas_Part6_RU.docx'),
    (7,  'Спасение — Мукти — Джеван Мукт',       'Nanak_Canvas_Part7_RU.docx'),
    (8,  'Суд — Дарга',                           'Nanak_Canvas_Part8_RU.docx'),
    (9,  'Дхарм Радж',                            'Nanak_Canvas_Part9_RU.docx'),
    (10, 'Яма Дуты и Читра Гупт',                'Nanak_Canvas_Part10_RU.docx'),
    (11, 'Предки — Питар',                        'Nanak_Canvas_Part11_RU.docx'),
    (12, 'Применение Холста',                     'Nanak_Canvas_Part12_RU.docx'),
]

BASE_DIR = Path(__file__).parent
OUTPUT   = BASE_DIR / 'Nanak_Canvas_COMBINED.docx'

# ── Formatting constants ──────────────────────────────────────────────────────

GURMUKHI_COLOR  = RGBColor(0x55, 0x00, 0x00)   # dark red
TRANSLIT_COLOR  = RGBColor(0x40, 0x40, 0x40)   # dark gray
BODY_COLOR      = RGBColor(0x00, 0x00, 0x00)
PART_HDR_COLOR  = RGBColor(0x1A, 0x1A, 0x5A)   # dark navy for part headers

FONT_NAME = 'Arial'


def is_gurmukhi_text(text: str) -> bool:
    """True if the paragraph contains Gurmukhi Unicode characters (U+0A00–U+0A7F)."""
    return any('\u0A00' <= ch <= '\u0A7F' for ch in text)


# ── XML / DOCX helpers ────────────────────────────────────────────────────────

_bm_counter = [0]   # mutable counter for bookmark IDs

def _next_bm_id() -> int:
    _bm_counter[0] += 1
    return _bm_counter[0]


def _set_run_font(run, size_pt, bold=False, italic=False, color=BODY_COLOR):
    run.font.name  = FONT_NAME
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_para(doc, text, size=11, bold=False, italic=False,
             color=BODY_COLOR, indent_cm=0, center=False, space_before=0):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    _set_run_font(run, size, bold, italic, color)
    return p


def add_page_break(doc):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    doc.element.body.append(p)


def add_bookmark_to_para(p_elem, bm_name: str) -> int:
    """
    Insert bookmarkStart + bookmarkEnd around the first run of a paragraph element.
    Returns the bookmark ID used.
    """
    bm_id = _next_bm_id()
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bm_id))
    bm_start.set(qn('w:name'), bm_name)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bm_id))
    # Insert at beginning of paragraph
    p_elem.insert(0, bm_start)
    p_elem.append(bm_end)
    return bm_id


def add_hyperlink_para(doc, text: str, anchor: str, size=12,
                       color=(0x1a, 0x1a, 0xaa)):
    """Add a paragraph containing an internal hyperlink (anchor → bookmark)."""
    p = doc.add_paragraph()
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Blue underline style for the link
    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), '%02X%02X%02X' % (color[0], color[1], color[2]))
    rPr.append(color_elem)
    u_elem = OxmlElement('w:u')
    u_elem.set(qn('w:val'), 'single')
    rPr.append(u_elem)
    sz_elem = OxmlElement('w:sz')
    sz_elem.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz_elem)
    szCs_elem = OxmlElement('w:szCs')
    szCs_elem.set(qn('w:val'), str(int(size * 2)))
    rPr.append(szCs_elem)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rPr.append(rFonts)

    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)
    return p


# ── Image merging ──────────────────────────────────────────────────────────────

# Global cache so identical images (same hash) are only added once per master
_img_cache: dict[str, str] = {}   # content_hash → new rId in master

IMAGE_RELTYPE = ('http://schemas.openxmlformats.org/officeDocument/'
                 '2006/relationships/image')


def _copy_images_from(master: Document, source: Document) -> dict[str, str]:
    """
    Copy all image parts from source into master's package.
    Returns a mapping {source_rId → master_rId}.
    """
    rel_map: dict[str, str] = {}
    for rId, rel in source.part.rels.items():
        if 'image' not in rel.reltype:
            continue
        img_part = rel.target_part
        blob     = img_part.blob
        ct       = img_part.content_type

        # Determine extension
        ext = ct.split('/')[-1]
        if ext in ('jpeg', 'jpg'):
            ext = 'jpg'
        elif ext == 'svg+xml':
            ext = 'svg'

        content_hash = hashlib.md5(blob).hexdigest()

        if content_hash in _img_cache:
            rel_map[rId] = _img_cache[content_hash]
            continue

        # Build a unique part name
        partname = PackURI(f'/word/media/c_{content_hash[:12]}.{ext}')
        new_part = Part(partname, ct, blob)
        new_rId  = master.part.relate_to(new_part, IMAGE_RELTYPE)

        _img_cache[content_hash] = new_rId
        rel_map[rId] = new_rId

    return rel_map


def _remap_blips(elem, rel_map: dict[str, str]):
    """Update all a:blip r:embed attributes in elem using rel_map."""
    for blip in elem.findall('.//' + qn('a:blip')):
        old = blip.get(qn('r:embed'))
        if old in rel_map:
            blip.set(qn('r:embed'), rel_map[old])


# ── Merge one document into master ────────────────────────────────────────────

def merge_into(master: Document, source: Document,
               bookmark_name: str | None = None):
    """
    Append every body element of source into master.
    - Remaps image relationship IDs
    - Inserts bookmark into the first paragraph if bookmark_name given
    """
    rel_map     = _copy_images_from(master, source)
    bookmark_done = False

    for elem in source.element.body:
        # Skip the final section properties (page size / margins of source)
        if elem.tag == qn('w:sectPr'):
            continue

        new_elem = copy.deepcopy(elem)
        _remap_blips(new_elem, rel_map)

        # Attach bookmark to very first paragraph of this part
        if not bookmark_done and bookmark_name and new_elem.tag == qn('w:p'):
            add_bookmark_to_para(new_elem, bookmark_name)
            bookmark_done = True

        master.element.body.append(new_elem)


# ── Post-pass: normalise formatting ──────────────────────────────────────────

def _para_runs_text(para) -> str:
    return ''.join(r.text or '' for r in para.runs)


def normalise_formatting(doc: Document):
    """
    Walk every paragraph in the combined document and:
    1. Gurmukhi paragraphs → dark-red, bold, 13pt Arial
    2. Italic+indented paragraphs (transliterations) → gray, italic, 10pt Arial, 1.2cm
    3. All other non-empty paragraphs → ensure Arial font is set (size/style unchanged)
    """
    for para in doc.paragraphs:
        text = _para_runs_text(para)
        if not text.strip():
            continue

        if is_gurmukhi_text(text):
            # Normalise Gurmukhi style
            para.alignment = None          # no forced alignment
            para.paragraph_format.left_indent = None
            for run in para.runs:
                run.font.name   = FONT_NAME
                run.font.size   = Pt(13)
                run.font.bold   = True
                run.font.italic = False
                run.font.color.rgb = GURMUKHI_COLOR
            continue

        # Detect transliteration: italic runs AND indented paragraph
        has_italic = any(r.font.italic for r in para.runs if r.text.strip())
        is_indented = (para.paragraph_format.left_indent is not None and
                       para.paragraph_format.left_indent > 0)
        # Also catch: starts with ASCII letter, small font
        starts_ascii = bool(re.match(r'^[A-Za-z]', text))
        small_font   = any(r.font.size and r.font.size <= Pt(10.5)
                           for r in para.runs if r.text.strip())

        if (has_italic or is_indented) and starts_ascii:
            para.paragraph_format.left_indent = Cm(1.2)
            for run in para.runs:
                run.font.name   = FONT_NAME
                run.font.size   = Pt(10)
                run.font.bold   = False
                run.font.italic = True
                run.font.color.rgb = TRANSLIT_COLOR
            continue

        # Regular body text — just ensure font name is Arial
        for run in para.runs:
            if run.font.name is None:
                run.font.name = FONT_NAME


# ── Title page ────────────────────────────────────────────────────────────────

def build_title_page(doc: Document):
    add_para(doc, '', size=11)  # top spacer

    add_para(doc, 'Karminder Singh Dhillon',
             size=13, color=RGBColor(0x40, 0x40, 0x40), center=True)

    add_para(doc, '',  size=8)

    add_para(doc, 'ХОЛСТ ГУРУ НАНАКА',
             size=26, bold=True, color=PART_HDR_COLOR, center=True, space_before=6)

    add_para(doc, 'Концепции Сикхи',
             size=16, color=PART_HDR_COLOR, center=True)

    add_para(doc, 'Серия из 12 частей',
             size=13, italic=True, color=RGBColor(0x40, 0x40, 0x40), center=True)

    for _ in range(4):
        add_para(doc, '', size=11)

    add_para(doc, 'Перевод: русский язык',
             size=11, italic=True, color=RGBColor(0x50, 0x50, 0x50), center=True)

    add_para(doc, 'Asia Samachar, 2021',
             size=11, color=RGBColor(0x50, 0x50, 0x50), center=True)

    add_page_break(doc)


# ── Table of contents ─────────────────────────────────────────────────────────

def build_toc(doc: Document, parts):
    add_para(doc, 'СОДЕРЖАНИЕ',
             size=16, bold=True, color=PART_HDR_COLOR, space_before=4)
    add_para(doc, '', size=6)

    for num, title, _, bm_name in parts:
        label = f'Часть {num}.  {title}'
        add_hyperlink_para(doc, label, bm_name, size=12)

    add_page_break(doc)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    # Build the bookmark name list alongside PARTS
    parts_with_bm = [(num, title, fname, f'part{num}')
                     for num, title, fname in PARTS]

    # ── Create master document ────────────────────────────────────────────────
    master = Document()

    # Set margins
    for sec in master.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── Title page ────────────────────────────────────────────────────────────
    build_title_page(master)

    # ── Table of contents ─────────────────────────────────────────────────────
    build_toc(master, parts_with_bm)

    # ── Merge each part ───────────────────────────────────────────────────────
    total_paras = 0
    total_imgs  = 0

    for num, title, fname, bm_name in parts_with_bm:
        path = BASE_DIR / fname
        if not path.exists():
            print(f'  WARNING: {path} not found — skipping', file=sys.stderr)
            continue

        source = Document(str(path))
        src_imgs  = sum(1 for r in source.part.rels.values() if 'image' in r.reltype)
        src_paras = len(source.paragraphs)

        print(f'  Part {num:2d}: {src_paras} paragraphs, {src_imgs} images — {fname}')

        # Part heading paragraph (before the source content)
        hdg = add_para(
            master,
            f'Часть {num}: {title}',
            size=15, bold=True, color=PART_HDR_COLOR, space_before=6,
        )
        # Add bookmark to THIS heading paragraph
        add_bookmark_to_para(hdg._p, bm_name)
        # Update the flag so merge_into won't try to add another bookmark
        merge_into(master, source, bookmark_name=None)

        total_paras += src_paras
        total_imgs  += src_imgs

        # Page break after each part except the last
        if num < 12:
            add_page_break(master)

    # ── Post-pass: normalise Gurmukhi + transliteration formatting ────────────
    print('\nNormalising Gurmukhi / transliteration formatting...')
    normalise_formatting(master)

    # ── Save ─────────────────────────────────────────────────────────────────
    master.save(str(OUTPUT))
    print(f'\nSaved: {OUTPUT}')
    print(f'Total source paragraphs merged: {total_paras}')
    print(f'Total source images merged:     {total_imgs}')


if __name__ == '__main__':
    main()
