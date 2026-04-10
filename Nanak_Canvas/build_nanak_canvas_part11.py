#!/usr/bin/env python3
"""
Build Russian docx for Guru Nanak's Canvas — Part 11: Ancestors (Piter).
"""

import re, os, html as html_module, urllib.request
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

URL     = "https://asiasamachar.com/2021/03/26/35742/"
OUTPUT  = "/home/royal/Work/Spiritual/Nanak_Canvas/Nanak_Canvas_Part11_RU.docx"
IMG_DIR = "/tmp/nanak_canvas_part11_imgs"
os.makedirs(IMG_DIR, exist_ok=True)

RU = {
    0:   'Карминдер Сингх | Концепции Сикхи | Часть 11 из 12 |',

    1:   'Прежде всего нам нужно рассмотреть '
         'до-1468-е верования, касающиеся Предков.',

    2:   'По учению старого холста, '
         'наши Питар (предки) пребывают в месте под названием Питар Лок. '
         'Там они существуют и нуждаются в пище, одежде и предметах обихода.',

    3:   'Ритуал, связанный с предками, называется Шраадх. '
         'Его цель — обеспечить усопших едой и предметами быта '
         'через духовенство, выступающее посредником.',

    4:   'Семья, совершающая Шраадх, приглашает духовенство. '
         'Его кормят, ему делают подношения — '
         'считается, что всё это достигает усопших.',

    5:   'Еду предлагают также душам ушедших. '
         'Подношения делаются трём поколениям умерших предков.',

    6:   'В сикхской общине наше духовенство '
         'распространило принцип поклонения предкам '
         'и ритуалов, с ним связанных.',

    7:   'Подношения — постели, еда, предметы обихода, '
         'посуда — делаются либо в гурдваре, '
         'либо у рек и водоёмов.',

    8:   'Большинство наших грантхи, рагиев и киртаниясов '
         'отстаивали такой ритуал или практику — '
         'прямо или косвенно, опираясь на стихи из СГГС.',

    9:   'Нельзя отрицать, что Гурбани употребляет '
         'слово «Питар» в Шри Гуру Грантх Сахибе. '
         'Стихов немало.',

    10:  'Рассмотрим эти стихи: '
         'пропагандирует ли Гурбани ритуалы, '
         'связанные с Питар?',

    11:  'Стих 1: Бхагат Кабир (СГГС, страница 332).',

    # 012: H4 Gurmukhi
    # 013: transliteration

    14:  'Смысл стиха: предков не чтили, пока они были живы. '
         'После смерти устраивают Шраадх. '
         'Но бедные предки — где они возьмут пищу? '
         'Ею питаются только ворона и собака.',

    15:  'И подношения еды в честь предков '
         'достаются в конечном счёте лишь воронам и собакам — '
         'тем, кто пирует на месте жертвоприношений.',

    16:  'Послание ясно. '
         'Эти стихи — критика ритуала Питар. '
         'Они говорят: ритуал бессмысленен.',

    17:  'А в стихе-Рахао того же шабда '
         'Бхагат Кабир говорит следующее.',

    18:  'Стих 2: Бхагат Кабир (СГГС, страница 332).',

    # 019: H4 Gurmukhi
    # 020: transliteration

    21:  'Кабир спрашивает: скажите мне, '
         'какое благо приносит этот ритуал? '
         'Весь мир совершает его и гибнет. '
         'Как тут обрести благо?',

    22:  'Послание снова ясно: '
         'от этого ритуала ничего не придёт к нам — '
         'потому что усопшие предки '
         'не могут ничего нам дать.',

    23:  'Стих 3: Гуру Нанак (СГГС, страница 472).',

    # 024: H4 Gurmukhi
    # 025: transliteration

    26:  'Смысл: если вор ограбил чей-то дом '
         'и сделал подношение своим умершим предкам из краденого — '
         'в загробном мире его узнают как вора, '
         'и его предки тоже станут ворами. '
         'Писец Дхарм Раджа всё запишет.',

    27:  'Аргумент Гуру Нанака проясняет '
         'для Сикха три вещи. '
         'Первое: любой ритуал подношения предкам отвергается, '
         'потому что они не могут его получить.',

    28:  'Второе: вся практика отправки вещей предкам '
         'нелогична и потому отвергается.',

    29:  'Третье: роль посредника-духовенства — '
         'это мошенничество. '
         'Он зарабатывает на живых, '
         'делая вид, что обеспечивает умерших.',

    30:  'Завершение стиха вносит ещё один вклад '
         'в нашу духовность. '
         'Используя такие слова, как «вор» и «краденое», '
         'Гуру Нанак указывает: '
         'подношения из краденого не достигают предков.',

    31:  'Нетрудно понять: '
         'даже если слово «Питар» встречается в стихах '
         'в контексте ритуалов — '
         'это не значит, что Гурбани их одобряет.',

    32:  'Послания Гурбани призывают нас '
         'держаться подальше от ритуалов '
         'и от духовенства, живущего за счёт этих ритуалов.',

    33:  'Но послания Гурбани не остановили наше духовенство — '
         'грантхи, рагиев, киртаниясов — '
         'от пропаганды этих ритуалов.',

    34:  'В поддержку своей позиции '
         'наше духовенство опирается на стих Гуру Арджуна, '
         'в котором встречается слово «Питри».',

    35:  'Стих 4: Гуру Арджун (СГГС, страница 496).',

    # 036: H4 Gurmukhi
    # 037: transliteration

    38:  'Первая часть стиха '
         'ਜਿਸੁ ਸਿਮਰਤ ਸਭਿ ਕਿਲਵਿਖ ਨਾਸਹਿ '
         'переводится как: '
         '«Памятование о Ком устраняет все грехи».',

    39:  'Наше духовенство переводит '
         'ਪਿਤਰੀ ਹੋਇ ਉਧਾਰੋ как '
         '«и спасёт ваших предков».',

    40:  'Утверждение духовенства таково: '
         'этот стих однозначно поддерживает концепцию '
         'Шраадха и поклонения предкам.',

    41:  'Послание духовенства Сикхам: '
         'Гуру Арджун прямо говорит — '
         'совершай Симран, это спасёт тебя '
         'и обеспечит спасение твоих предков.',

    42:  'Однако в этом переводе духовенства '
         'три ошибки. '
         'Первая: принцип Гурмат состоит в том, '
         'что духовные плоды — '
         'это результат собственных усилий.',

    43:  'Этот стих Гуру Нанака '
         'на странице 474 СГГС '
         'излагает этот основной принцип.',

    44:  'Стих 5: Гуру Нанак (СГГС, страница 474).',

    # 045: H4 Gurmukhi
    # 046: transliteration

    47:  'Смысл: духовные плоды — '
         'результат наших собственных действий, '
         'совершённых нами самими.',

    48:  'Нам также известен этот стих Гуру Нанака.',

    49:  'Стих 7: Гуру Нанак (СГГС, страница 4).',

    # 050: H4 Gurmukhi

    51:  'Смысл: сам посеял — сам и пожнёшь.',

    52:  'Этот стих Гуру Арджуна '
         'на странице 134 СГГС подтверждает.',

    53:  'Стих 8: Гуру Арджун (СГГС, страница 134).',

    # 054: H4 Gurmukhi
    # 055: transliteration

    56:  'Смысл: духовность — это поле, '
         'где сеются семена поступков '
         'и вызревают плоды этих поступков.',

    57:  'В совокупности принцип Гурмат прост: '
         'Сикхи — это духовность индивидуального усилия. '
         'Никто не может совершить духовный путь за другого.',

    58:  'Вторая ошибка: слово ਹੋਇ '
         'в выражении ਪਿਤਰੀ ਹੋਇ ਉਧਾਰੋ — '
         'стоит в прошедшем времени. '
         'Значит, правильный перевод духовенства '
         'должен был бы быть '
         '«как спаслись предки» — '
         'а не «предки будут спасены».',

    59:  'Третья ошибка: это стих Гуру Арджуна. '
         'Поскольку Гурбани написана от первого лица, '
         'говорящий в стихе — сам Гуру Арджун.',

    60:  'Нужно ли предкам Гуру Арджуна спасение? '
         'Вопрос риторический.',

    61:  'Собирая все три возражения, '
         'перевод духовенства '
         'ਪਿਤਰੀ ਹੋਇ ਉਧਾਰੋ как «спасёт ваших предков» '
         'несостоятелен.',

    62:  'Правильный перевод стиха таков: '
         'Памятование о Котором устраняет все грехи — '
         'и точно так же спасло моих предков в настоящем.',

    63:  'ਪਿਤਰੀ ਹੋਇ ਉਧਾਰੋ означает: '
         'так же как памятование спасло наших предков '
         'здесь и сейчас — '
         'так и ты практикуй это памятование.',

    64:  'Следующий стих того же двустишия '
         'подтверждает это.',

    65:  'Стих 9: Гуру Арджун (СГГС, страница 496).',

    # 066: H4 Gurmukhi
    # 067: transliteration

    68:  'Смысл: постигай того Непостижимого '
         'Вездесущего Творца, '
         'постижение Которого спасало наших предков — '
         'Которому нет ни конца, ни края.',

    69:  'Таким образом ясно: '
         'в этом шабде Гуру Арджуна '
         'выражение ਪਿਤਰੀ ਹੋਇ ਉਧਾਰੋ '
         'не имеет отношения к ритуалам.',

    70:  'Как сказано выше, '
         'слово «Питар», «Питра(н)» и «Питри(н)» '
         'встречается в СГГС. '
         'Во ВСЕХ этих случаях — '
         'либо для критики ритуала, '
         'либо в переопределённом смысле.',

    71:  'Единственное послание Гурбани '
         'о наших Предках заключено '
         'в следующем стихе.',

    72:  'Стих 10: Бхагат Кабир (СГГС, страница 332).',

    # 073: H4 Gurmukhi
    # 074: transliteration

    75:  'Смысл: предков не чтили при жизни — '
         'после смерти хотим отправить им пищу. '
         'Бедные предки — где им взять эту пищу? '
         'Ею питается лишь ворона.',

    76:  'Послание — УСТРАНИТЬ это противоречие. '
         'ОТВЕРГНУТЬ это лицемерие. '
         'И путь к этому — '
         'чтить предков, пока они живы, '
         'а не после смерти.',

    77:  'По существу, Гурбани освободила нас '
         'от всего нарратива '
         'о культе Предков из четырёхтысячелетнего холста.',

    78:  'Когда ты перестаёшь хотеть '
         'поклоняться предкам после их смерти — '
         'ты начинаешь уважать их, '
         'пока они живы.',

    79:  'Когда ты перестаёшь гнаться '
         'за духовностью загробной жизни — '
         'ты начинаешь жить '
         'духовностью настоящей жизни.',

    80:  'Когда ты отвергаешь нарративы, '
         'которые сама Гурбани отвергла, — '
         'ты начинаешь жить по Гурбани.',

    81:  'Когда ты ПЕРЕСТАЁШЬ гнаться '
         'за неправильным — '
         'у правильного появляется шанс '
         'настичь тебя.',
}

SKIP = set(range(82, 107))

# ── helpers ──────────────────────────────────────────────────────────────────

def best_img_src(tag_html):
    m = re.search(r'\bsrc="([^"]+)"', tag_html)
    if not m:
        return None
    src = m.group(1)
    if re.search(r'-\d+x\d+\.(jpg|png)', src):
        return None
    return src


def download_img(src):
    fname = os.path.basename(src.split('?')[0])
    dest  = os.path.join(IMG_DIR, fname)
    if not os.path.exists(dest):
        try:
            urllib.request.urlretrieve(src, dest)
            print(f'  Downloaded: {fname}')
        except Exception as e:
            print(f'  FAILED {src}: {e}')
            return None
    return dest


def emit_img_node(tag_html):
    src = best_img_src(tag_html)
    if not src:
        return None
    if 'uploads/2021' not in src and 'uploads/2020' not in src:
        return None
    path = download_img(src)
    if not path:
        return None
    return {'type': 'img', 'path': path}


def split_by_imgs(raw_html):
    img_re = re.compile(r'<img\b[^>]*/?>|<a[^>]*>\s*<img\b[^>]*/?>.*?</a>',
                        re.IGNORECASE | re.DOTALL)
    nodes = []
    last  = 0
    for m in img_re.finditer(raw_html):
        if m.start() > last:
            nodes.append({'type': 'text', 'html': raw_html[last:m.start()]})
        img_node = emit_img_node(m.group(0))
        if img_node:
            nodes.append(img_node)
        last = m.end()
    if last < len(raw_html):
        nodes.append({'type': 'text', 'html': raw_html[last:]})
    return nodes


# ── fetch & parse ─────────────────────────────────────────────────────────────

print('Fetching article...')
req = urllib.request.Request(URL, headers={'User-Agent': 'Mozilla/5.0'})
with urllib.request.urlopen(req) as r:
    raw = r.read().decode('utf-8', errors='replace')

parts = list(re.finditer(r'<div[^>]+class="td-post-content"', raw))
start = parts[0].start()
end_m = re.search(r'<div[^>]+class="td-post-source-tags"', raw[start:])
content = raw[start: start + end_m.start()] if end_m else raw[start: start+60000]

block_re = re.compile(r'<(h[1-6]|p|hr)\b([^>]*)>(.*?)</\1>',
                      re.IGNORECASE | re.DOTALL)
raw_nodes = list(block_re.finditer(content))
print(f'Parsed {len(raw_nodes)} nodes')

# ── build docx ────────────────────────────────────────────────────────────────

HEADING_STYLES = {
    4: dict(size=13, bold=True,  color=(0x55, 0x00, 0x00)),
    5: dict(size=12, bold=True,  color=(0x20, 0x20, 0x60)),
    6: dict(size=11, bold=True,  color=(0x20, 0x20, 0x60)),
}

doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)


def add_para(text, size=11, bold=False, italic=False,
             color=(0, 0, 0), indent=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.left_indent = Cm(1.2)
    run = p.add_run(text)
    run.font.name   = 'Arial'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color)
    return p


def add_image(path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(5.5))


images_inserted = 0

for i, node in enumerate(raw_nodes):
    if i in SKIP:
        continue

    tag    = node.group(1).lower()
    inner  = node.group(3)
    hlevel = int(tag[1]) if tag.startswith('h') else 0

    plain_text = html_module.unescape(re.sub(r'<[^>]+>', '', inner).strip())

    # H4 (Gurmukhi) — keep as-is in dark red
    if hlevel == 4:
        style = HEADING_STYLES[4]
        sub_nodes = split_by_imgs(inner)
        for sn in sub_nodes:
            if sn['type'] == 'img':
                add_image(sn['path'])
                images_inserted += 1
            else:
                t = html_module.unescape(re.sub(r'<[^>]+>', '', sn['html']).strip())
                if t:
                    add_para(t, size=style['size'], bold=style['bold'],
                             color=style['color'])
        continue

    # Transliteration: only the node immediately following an H4
    prev_tag = raw_nodes[i-1].group(1).lower() if i > 0 else ''
    is_transliteration = (
        hlevel == 0 and
        prev_tag == 'h4' and
        '<strong' not in inner.lower() and
        re.match(r'^[A-Za-z]', plain_text)
    )

    if is_transliteration:
        sub_nodes = split_by_imgs(inner)
        for sn in sub_nodes:
            if sn['type'] == 'img':
                add_image(sn['path'])
                images_inserted += 1
            else:
                t = html_module.unescape(re.sub(r'<[^>]+>', '', sn['html']).strip())
                if t:
                    add_para(t, size=10, italic=True,
                             color=(0x40, 0x40, 0x40), indent=True)
        continue

    display = RU.get(i, None)

    sub_nodes = split_by_imgs(inner)
    for sn in sub_nodes:
        if sn['type'] == 'img':
            add_image(sn['path'])
            images_inserted += 1
        else:
            if display is None:
                continue
            t = display
            if not t or not t.strip():
                continue
            if hlevel in HEADING_STYLES:
                style = HEADING_STYLES[hlevel]
                add_para(t, size=style['size'], bold=style['bold'],
                         color=style['color'])
            else:
                add_para(t)

doc.save(OUTPUT)
print(f'\nSaved: {OUTPUT}')
print(f'Images inserted: {images_inserted}')
