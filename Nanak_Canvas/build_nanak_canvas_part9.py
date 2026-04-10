#!/usr/bin/env python3
"""
Build Russian docx for Guru Nanak's Canvas — Part 9: Dhrm Raj.
"""

import re, os, html as html_module, urllib.request
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

URL     = "https://asiasamachar.com/2021/02/27/36844/"
OUTPUT  = "/home/royal/Work/Spiritual/Nanak_Canvas/Nanak_Canvas_Part9_RU.docx"
IMG_DIR = "/tmp/nanak_canvas_part9_imgs"
os.makedirs(IMG_DIR, exist_ok=True)

RU = {
    0:   'Карминдер Сингх | Концепции Сикхи | Часть 9 из 12 |',

    1:   'Каков нарратив о Дхарм Радже в Гурбани? '
         'Основные принципы можно выстроить, '
         'обратившись к нескольким ключевым стихам.',

    2:   'Гурбани делает две вещи. '
         'Первое: Гурбани РАЗРУШАЕТ миф о Дхарм Радже. '
         'Второе: Гурбани ПЕРЕОПРЕДЕЛЯЕТ концепцию Дхарм Раджа.',

    # 003: H5 "TO WATCH THE VIDEO" — skip

    4:   'Стих 1: Гуру Арджун на странице 614 СГГС.',

    # 005: H4 Gurmukhi

    6:   'Смысл: Мой Гуру наделил меня богатством просветления. '
         'В результате обман, живший в моём уме, рассеян. '
         'Дхарм Радж — кто этот Дхарм Радж? — '
         'записи моих дел порваны.',

    7:   'Миф разрушен, и страх исчезает. '
         'Я сам веду свои счета, сам оцениваю свои дела '
         'здесь и сейчас — и записи этих дел ведёт '
         'Голос Совести внутри.',

    8:   'Стих 2: Гуру Арджун на странице 1348 СГГС.',

    # 009: H4 Gurmukhi

    10:  'Смысл: Я постиг Божественные добродетели '
         'и отверг человеческие пороки. '
         'Тем самым я уничтожил записи Дхарм Рая — '
         'ни одной строки не осталось.',

    11:  'И снова: миф разрушен, страх перед Дхарм Раджем '
         'в загробной жизни уходит. '
         'Стих Гуру Арджуна ясно говорит: '
         'Дхарм Радж не ведёт записей. '
         'Я сам учитываю добро и зло '
         'через Голос Совести внутри.',

    12:  'Стих 3: Бхагат Кабир на странице 793 СГГС.',

    # 013: H4 Gurmukhi

    14:  'Смысл: Единый Творец пребывает в полноте моего ума. '
         'Его Божественные добродетели укоренились во мне. '
         'Я обыскал контору Дхарм Рая — '
         'и не нашёл ни одного долга.',

    15:  'Послание ясно. '
         'Обыскать контору — это идиоматическое выражение: '
         'для меня мифа о Дхарм Радже больше не существует. '
         'Долгов нет. '
         'Кабир, кожевник из низшей касты, '
         'говорит это языком бесстрашия и свободы.',

    16:  'Стих 4: Гуру Рамдас на странице 698 СГГС.',

    # 017: H4 Gurmukhi

    18:  'Смысл: Творец даровал мне постижение внутри. '
         'Нанак, записи Дхарм Рая разорваны — '
         'счёт мой завершён.',

    19:  'Послание снова ясно. '
         'Для Гурбани и Гуру Рамдас Джи '
         'реальность такова: '
         'Нанак, твои записи уничтожены, '
         'твой счёт завершён.',

    20:  'Стих 5: Гуру Рамдас на странице 1326 СГГС.',

    # 021: H4 Gurmukhi

    22:  'Смысл: Я отверг эго и всё стремление '
         'к мирской славе и вошёл в общество Садху. '
         'Это позволило мне избавиться '
         'от страха перед Дхарм Раем — '
         'и перебросить мне спасательную верёвку.',

    23:  'Подводя итог: позиция Гурбани '
         'в отношении выдуманного мифа о Дхарм Радже '
         'как судье в загробной жизни '
         'такова: это вымысел, миф, обман — '
         'и Гурбани его разрушает.',

    24:  'Начало Сикхи в том, что Творец — внутри. '
         'Если Творец внутри — Его Суд внутри. '
         'Если Суд внутри — Судья внутри. '
         'Судья — Голос Совести внутри. '
         'И всё это — в НАСТОЯЩЕМ.',

    25:  'Гурбани также ПЕРЕОПРЕДЕЛИЛА концепцию Дхарм Раджа. '
         'Радж означает Царство, Дхарм означает духовную жизнь. '
         'Дхарм Радж, таким образом, — '
         'это Царство Дхармы, Царство Духовной Жизни.',

    26:  'Итак, не только Творец, Его Суд, '
         'Его Приговор и Судья — всё это внутри. '
         'Царство Дхармы — это Царство жизни по Дхарму, '
         'жизни праведной, жизни по Воле Творца. '
         'И достигается оно здесь и сейчас, '
         'в нынешней жизни.',

    27:  'Стих 6: Гуру Арджун на странице 406 СГГС.',

    # 028: H4 Gurmukhi
    # 029: H4 Gurmukhi (second stanza)

    30:  'Смысл: Мой Гуру благословил моё бытие. '
         'Все мои чувства, прежде уходившие на дурное, '
         'теперь служат Тебе. '
         'Ты Сам — Творец, Хранитель всей вселенной, '
         'пребывающий во всём. '
         'Дхарм Радж изумлён и поклоняется Тебе.',

    31:  'Очевидно, что Дхарм Радж в Гурбани — '
         'это совершенно иная концепция, '
         'никак не связанная с Дхарм Раджем '
         'четырёхтысячелетнего старого холста.',

    32:  'Стих 7: Бхагат Кабир делает свой выбор '
         'и сообщает нам о нём '
         'на странице 1158 СГГС.',

    # 033: H4 Gurmukhi

    34:  'У меня нет никакого конфликта в уме. '
         'Никакой внутренней борьбы больше нет. '
         '(Я достиг этого,) оставив позади '
         'и пандита, и муллу.',

    # 035: H4 Gurmukhi

    36:  'Всё, что написали мне пандит и мулла, '
         'я отбросил, идя своим путём, '
         'и не взял с собой ничего.',

    37:  'Такой пересмотр нашей системы убеждений '
         'и нашей психики необходим, '
         'чтобы истины Гурбани '
         'могли войти в нас и обосноваться в нас.',

    # 038: H4 Gurmukhi

    39:  'Смысл: (Плод отбрасывания таков:) '
         'я могу наполнить очищенный ум чистотой. '
         'Я исследовал себя — и встретил Кабира.',

    40:  'Следует помнить: '
         'всякий раз, когда в Гурбани встречается '
         'слово «Дхарм Радж» или «Дхарм Рай», '
         'оно употребляется в одном из двух значений. '
         'Первое — это отрицание: разрушение мифа '
         'о Дхарм Радже как судье в загробной жизни.',

    41:  'Второе значение — переопределённое. '
         'Дхарм Радж в контексте Гурбани — '
         'это Царство Дхармы, '
         'которое созидается внутри, '
         'здесь и сейчас.',

    42:  'Наша задача — как Сикхов Гурбани, '
         'Шри Гуру Грантх Сахиба — '
         'распознавать, в каком из двух контекстов '
         'употребляется слово «Дхарм Радж» в Гурбани.',

    43:  'По существу, Гурбани '
         'освободила нас от всего нарратива '
         'о Дхарм Радже из четырёхтысячелетнего холста, '
         'описанного как судья в загробной жизни.',

    44:  'Когда ты перестаёшь бояться Дхарм Раджа '
         'загробного мира — '
         'ты начинаешь созидать Радж Дхармы '
         'здесь и сейчас, '
         'в своей нынешней жизни.',
}

SKIP = set(range(45, 72)) | {3}

# ── helpers ──────────────────────────────────────────────────────────────────

def best_img_src(tag_html):
    m = re.search(r'\bsrc="([^"]+)"', tag_html)
    if not m:
        return None
    src = m.group(1)
    if re.search(r'-\d+x\d+\.(jpg|png)', src):
        return None
    return src


def download_img(src):
    fname = os.path.basename(src.split('?')[0])
    dest  = os.path.join(IMG_DIR, fname)
    if not os.path.exists(dest):
        try:
            urllib.request.urlretrieve(src, dest)
            print(f'  Downloaded: {fname}')
        except Exception as e:
            print(f'  FAILED {src}: {e}')
            return None
    return dest


def emit_img_node(tag_html):
    src = best_img_src(tag_html)
    if not src:
        return None
    if 'uploads/2021' not in src and 'uploads/2020' not in src:
        return None
    path = download_img(src)
    if not path:
        return None
    return {'type': 'img', 'path': path}


def split_by_imgs(raw_html):
    img_re = re.compile(r'<img\b[^>]*/?>|<a[^>]*>\s*<img\b[^>]*/?>.*?</a>',
                        re.IGNORECASE | re.DOTALL)
    nodes = []
    last  = 0
    for m in img_re.finditer(raw_html):
        if m.start() > last:
            nodes.append({'type': 'text', 'html': raw_html[last:m.start()]})
        img_node = emit_img_node(m.group(0))
        if img_node:
            nodes.append(img_node)
        last = m.end()
    if last < len(raw_html):
        nodes.append({'type': 'text', 'html': raw_html[last:]})
    return nodes


# ── fetch & parse ─────────────────────────────────────────────────────────────

print('Fetching article...')
req = urllib.request.Request(URL, headers={'User-Agent': 'Mozilla/5.0'})
with urllib.request.urlopen(req) as r:
    raw = r.read().decode('utf-8', errors='replace')

parts = list(re.finditer(r'<div[^>]+class="td-post-content"', raw))
start = parts[0].start()
end_m = re.search(r'<div[^>]+class="td-post-source-tags"', raw[start:])
content = raw[start: start + end_m.start()] if end_m else raw[start: start+60000]

block_re = re.compile(r'<(h[1-6]|p|hr)\b([^>]*)>(.*?)</\1>',
                      re.IGNORECASE | re.DOTALL)
raw_nodes = list(block_re.finditer(content))
print(f'Parsed {len(raw_nodes)} nodes')

# ── build docx ────────────────────────────────────────────────────────────────

HEADING_STYLES = {
    4: dict(size=13, bold=True,  color=(0x55, 0x00, 0x00)),
    5: dict(size=12, bold=True,  color=(0x20, 0x20, 0x60)),
    6: dict(size=11, bold=True,  color=(0x20, 0x20, 0x60)),
}

doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)


def add_para(text, size=11, bold=False, italic=False,
             color=(0, 0, 0), indent=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.left_indent = Cm(1.2)
    run = p.add_run(text)
    run.font.name   = 'Arial'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color)
    return p


def add_image(path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(5.5))


images_inserted = 0

for i, node in enumerate(raw_nodes):
    if i in SKIP:
        continue

    tag   = node.group(1).lower()
    inner = node.group(3)
    hlevel = int(tag[1]) if tag.startswith('h') else 0

    plain_text = html_module.unescape(re.sub(r'<[^>]+>', '', inner).strip())

    # H4 (Gurmukhi) — keep as-is
    if hlevel == 4:
        style = HEADING_STYLES[4]
        sub_nodes = split_by_imgs(inner)
        for sn in sub_nodes:
            if sn['type'] == 'img':
                add_image(sn['path'])
                images_inserted += 1
            else:
                t = html_module.unescape(re.sub(r'<[^>]+>', '', sn['html']).strip())
                if t:
                    add_para(t, size=style['size'], bold=style['bold'],
                             color=style['color'])
        continue

    display = RU.get(i, None)

    sub_nodes = split_by_imgs(inner)
    for sn in sub_nodes:
        if sn['type'] == 'img':
            add_image(sn['path'])
            images_inserted += 1
        else:
            if display is None:
                continue
            t = display
            if not t or not t.strip():
                continue
            if hlevel in HEADING_STYLES:
                style = HEADING_STYLES[hlevel]
                add_para(t, size=style['size'], bold=style['bold'],
                         color=style['color'])
            else:
                add_para(t)

doc.save(OUTPUT)
print(f'\nSaved: {OUTPUT}')
print(f'Images inserted: {images_inserted}')
