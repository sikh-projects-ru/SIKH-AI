# chatgpt_darpan/bot.py

from __future__ import annotations

import argparse
import json
import os
import re
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from playwright.sync_api import Page, TimeoutError as PWTimeout, sync_playwright
from playwright_stealth import Stealth

_stealth = Stealth()

DARPAN_URL = "https://www.gurugranthdarpan.net/{page:04d}.html"
DEFAULT_CHAT_URL = "https://chatgpt.com/"
BOT_PROFILE = Path(__file__).parent / "bot_profile"

TITLE_TEXT = "Шри Гуру Грантх Сахиб — перевод на русский"
SUBTITLE_TEXT = "Источники: Guru Granth Sahib Darpan · Prof. Sahib Singh"
COMMENTARY_LABEL = "Комментарий Prof. Sahib Singh"

COLOR_GURBANI = RGBColor(0x55, 0x00, 0x00)
COLOR_ROMAN = RGBColor(0x44, 0x44, 0x44)
COLOR_LINE = RGBColor(0x00, 0x60, 0x00)
COLOR_ARTH = RGBColor(0x33, 0x33, 0x55)
COLOR_COMMENTARY = RGBColor(0x33, 0x33, 0x55)
COLOR_LABEL = RGBColor(0x99, 0x99, 0x99)
COLOR_HEADING = RGBColor(0x22, 0x22, 0x22)

ROMANIZATION_RULES = """\
Требования к полю "roman":
- нужна аккуратная панджабская romanization для русскоязычного читателя;
- сохраняй диакритику там, где это важно: ā, ī, ū, ṭ, ṭh, ḍ, ḍh, ṇ, ṛ, ñ, ś, ṃ, ṅ;
- не делай механическую школьную санскритизацию;
- roman должна читаться ближе к живому панджабскому чтению, а не как буквенная ловушка;
- не тащи автоматически конечные -u / -i, если в принятом чтении они не звучат отдельно;
- для форм вроде ਸੋਚੈ не пиши "soce"; пиши ближе к чтению: "sochey";
- для форм вроде ਗਾਵੈ не пиши "gāvai"; пиши ближе к чтению: "gāvē";
- для SGGS Darpan используй практичную научную romanization с учётом чтения.

Предпочтительные примеры:
- ਸੋਚੈ → sochey
- ਗਾਵੈ → gāvē
- ਹੁਕਮਿ → hukam
- ਸਤਿਨਾਮੁ → sat nām
- ਜਪੁ → jap
- ਗੁਰ ਪ੍ਰਸਾਦਿ → gur prasād
- ਅਜੂਨੀ → ajūnī
- ਸੈਭੰ → saibhaṅ
- ਕਰਤਾ ਪੁਰਖੁ → kartā purakh
- ਨਿਰਭਉ → nirbhau
- ਨਿਰਵੈਰੁ → nirvair
"""

RUSSIAN_GLOSSARY = """\
Предпочтительные русские формы в переводе:
- ਹਉਮੈ → Хоумэ
- ਜਪੁ → Джап
- ਜਪੁ ਜੀ → Джап Джи
- ਸਤਿਗੁਰੁ / ਸਤਿਗੁਰੂ → Сатгуру
- ਸਤਿਨਾਮੁ → Сат Нам
- ਏਕੰਕਾਰੁ / ਓਅੰਕਾਰੁ → Онкар / Эконкар по контексту, без механического финального "у"
- ਗੁਰ ਪ੍ਰਸਾਦਿ → Гур прасад
- ਅਜੂਨੀ → Аджуни
- ਸੈਭੰ → Сэйбханг
- ਕਰਤਾ ਪੁਰਖੁ → Карта Пуракх
- ਨਿਰਭਉ → Нирбхау
- ਨਿਰਵੈਰੁ → Нирвэр

Общее правило:
- roman: аккуратно и единообразно;
- русский текст: естественно, по панджабскому чтению и смыслу.
"""

PROMPT_TEMPLATE = """\
Игнорируй весь предыдущий контекст этого чата.
Используй только текст текущего сообщения.

Переведи следующий текст из комментария к Гуру Грантх Сахиб (Guru Granth Sahib Darpan, проф. Сахиб Сингх) на русский язык.

Нужно подготовить материал для DOCX по одному ангу.

В этом анге ровно {expected_verses} блоков [BANI].
Верни ровно {expected_verses} элементов в массиве "verses".

Для каждой строфы верни четыре уровня:
1) "gurmukhi" — текст строфы на панджаби;
2) "roman" — аккуратная romanization;
3) "line_translation" — ровный, короткий, цельный перевод самой строки;
4) "arth" — развёрнутый смысловой разбор на основе pad-arth / arath:
   - объясняй ключевые слова,
   - поясняй связки и смысл,
   - можно писать в духе:
     "По Хукаму возникают тела, но сам Хукам не поддаётся описанию. 'Ākār' — это формы, тела..."
   - если в материале есть разбор слов, используй его;
   - arth не должен быть просто повтором line_translation.

Также верни:
5) "commentary" — общий комментарий к ангу, если он есть.
Если общего комментария нет, верни пустую строку.

Как понимать вход:
- блоки [BANI] — это строфы;
- блоки [ARATH] рядом с ними — это толкование, pad-arth, пояснения;
- длинные пояснительные фрагменты без собственной строфы можно перенести в commentary;
- не выдумывай gurmukhi, если его нет.

{romanization_rules}

{russian_glossary}

Правила:
- сохраняй порядок;
- не пропускай фрагменты;
- не добавляй вступлений, пояснений от себя, markdown и code fences;
- не пиши заголовки вроде "Транслитерация", "Перевод", "Артх";
- поля "roman", "line_translation" и "arth" обязательны для каждой строфы;
- ответ верни строго между BEGIN_DARPAN_JSON и END_DARPAN_JSON.

Формат ответа:

BEGIN_DARPAN_JSON
{{
  "verses": [
    {{
      "gurmukhi": "...",
      "roman": "...",
      "line_translation": "...",
      "arth": "..."
    }}
  ],
  "commentary": "..."
}}
END_DARPAN_JSON

Текст анга:

{content}
"""

REPAIR_PROMPT_TEMPLATE = """\
Игнорируй весь предыдущий контекст этого чата, кроме твоего последнего ответа.
Используй только смысл своего последнего ответа и перепиши его в правильном формате.

В этом анге должно быть ровно {expected_verses} элементов в массиве "verses".

Нужен только валидный JSON между BEGIN_DARPAN_JSON и END_DARPAN_JSON.
Без markdown, без code fences, без пояснений.

Обязательные поля для каждой строфы:
- gurmukhi
- roman
- line_translation
- arth

Также:
- commentary

roman должна учитывать панджабское чтение для русскоязычного читателя.
Используй формы вроде:
- sochey
- gāvē
- hukam
- sat nām
- gur prasād
- saibhaṅ

Формат:

BEGIN_DARPAN_JSON
{{
  "verses": [
    {{
      "gurmukhi": "...",
      "roman": "...",
      "line_translation": "...",
      "arth": "..."
    }}
  ],
  "commentary": "..."
}}
END_DARPAN_JSON
"""


@dataclass
class Verse:
    gurmukhi: str
    roman: str
    line_translation: str
    arth: str


@dataclass
class AngTranslation:
    verses: list[Verse]
    commentary: str


@dataclass
class RuntimeConfig:
    page_timeout_ms: int
    input_timeout_ms: int
    response_timeout_ms: int
    new_message_timeout_ms: int
    max_retries: int
    retry_delay_s: float
    raw_log_dir: Path | None
    json_dir: Path


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def safe_filename_part(name: str) -> str:
    return re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", name)


def build_progress_path(output_path: Path) -> Path:
    return output_path.parent / f".progress_{safe_filename_part(output_path.name)}.txt"


def load_progress(progress_file: Path) -> int:
    if not progress_file.exists():
        return 0
    try:
        return int(progress_file.read_text(encoding="utf-8").strip() or "0")
    except ValueError:
        return 0


def save_progress(progress_file: Path, ang: int) -> None:
    progress_file.write_text(str(ang), encoding="utf-8")


def ensure_output_doc_exists(docx_path: Path) -> None:
    if docx_path.exists():
        return

    doc = Document()
    doc.core_properties.author = "ChatGPT Darpan Bot"

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title = doc.add_heading(TITLE_TEXT, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(SUBTITLE_TEXT)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = COLOR_LABEL

    doc.add_paragraph()
    doc.save(str(docx_path))


def reset_output_doc(docx_path: Path) -> None:
    if docx_path.exists():
        docx_path.unlink()
    ensure_output_doc_exists(docx_path)


def add_run(
    para,
    text: str,
    *,
    color=None,
    bold: bool = False,
    italic: bool = False,
    size_pt: int | None = None,
):
    run = para.add_run(text)
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def add_ang_heading(doc: Document, ang: int) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(8)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run(f"АНГ {ang}")
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_HEADING


def add_vertical_verse_block(doc: Document, verse: Verse) -> None:
    p_g = doc.add_paragraph()
    p_g.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_g.paragraph_format.space_before = Pt(8)
    p_g.paragraph_format.space_after = Pt(2)
    add_run(p_g, verse.gurmukhi, color=COLOR_GURBANI, size_pt=13, bold=True)

    p_r = doc.add_paragraph()
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r.paragraph_format.space_after = Pt(3)
    add_run(p_r, verse.roman, color=COLOR_ROMAN, size_pt=10, italic=True)

    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t.paragraph_format.space_after = Pt(5)
    add_run(p_t, verse.line_translation, color=COLOR_LINE, size_pt=11)

    p_a = doc.add_paragraph()
    p_a.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_a.paragraph_format.left_indent = Inches(0.2)
    p_a.paragraph_format.right_indent = Inches(0.2)
    p_a.paragraph_format.space_after = Pt(10)
    add_run(p_a, verse.arth, color=COLOR_ARTH, size_pt=9)

    doc.add_paragraph()


def add_commentary_block(doc: Document, commentary: str) -> None:
    commentary = normalize_text(commentary)
    if not commentary:
        return

    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.paragraph_format.space_before = Pt(10)
    label.paragraph_format.space_after = Pt(4)
    run = label.add_run(COMMENTARY_LABEL)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_LABEL

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.left_indent = Inches(0.2)
    para.paragraph_format.right_indent = Inches(0.2)
    run2 = para.add_run(commentary)
    run2.font.size = Pt(9)
    run2.font.color.rgb = COLOR_COMMENTARY


def append_ang_to_docx(docx_path: Path, ang: int, ang_data: AngTranslation) -> None:
    ensure_output_doc_exists(docx_path)
    doc = Document(str(docx_path))

    add_ang_heading(doc, ang)

    for verse in ang_data.verses:
        add_vertical_verse_block(doc, verse)

    add_commentary_block(doc, ang_data.commentary)
    doc.save(str(docx_path))
    print(f"  ✓ DOCX дописан: {docx_path.name}")


def ang_to_dict(ang: int, ang_data: AngTranslation) -> dict[str, Any]:
    return {
        "ang": ang,
        "verse_count": len(ang_data.verses),
        "verses": [
            {
                "gurmukhi": verse.gurmukhi,
                "roman": verse.roman,
                "line_translation": verse.line_translation,
                "arth": verse.arth,
            }
            for verse in ang_data.verses
        ],
        "commentary": ang_data.commentary,
    }


def verse_from_dict(item: dict[str, Any]) -> Verse:
    return Verse(
        gurmukhi=normalize_text(str(item.get("gurmukhi", ""))),
        roman=normalize_text(str(item.get("roman", ""))),
        line_translation=normalize_text(str(item.get("line_translation", ""))),
        arth=normalize_text(str(item.get("arth", ""))),
    )


def ang_translation_from_dict(data: dict[str, Any]) -> AngTranslation:
    verses = [verse_from_dict(item) for item in data.get("verses", []) if isinstance(item, dict)]
    commentary = normalize_text(str(data.get("commentary", "")))
    return AngTranslation(verses=verses, commentary=commentary)


def ang_json_path(json_dir: Path, ang: int) -> Path:
    return json_dir / f"ang_{ang:04d}.json"


def save_ang_json(json_dir: Path, ang: int, ang_data: AngTranslation) -> Path:
    json_dir.mkdir(parents=True, exist_ok=True)
    json_path = ang_json_path(json_dir, ang)
    json_path.write_text(
        json.dumps(ang_to_dict(ang, ang_data), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return json_path


def load_ang_json(json_dir: Path, ang: int) -> AngTranslation | None:
    json_path = ang_json_path(json_dir, ang)
    if not json_path.exists():
        return None
    try:
        data = json.loads(json_path.read_text(encoding="utf-8"))
    except Exception:
        return None
    return ang_translation_from_dict(data)


def rebuild_docx_from_json(output_path: Path, json_dir: Path, start: int, end: int) -> int:
    reset_output_doc(output_path)
    rebuilt = 0

    for ang in range(start, end + 1):
        ang_data = load_ang_json(json_dir, ang)
        if not ang_data or not ang_data.verses:
            print(f"  ↷ Нет JSON для анга {ang}, пропускаю")
            continue

        append_ang_to_docx(output_path, ang, ang_data)
        rebuilt += 1

    return rebuilt


def save_raw_text(raw_log_dir: Path | None, filename: str, text: str) -> None:
    if raw_log_dir is None:
        return
    raw_log_dir.mkdir(parents=True, exist_ok=True)
    (raw_log_dir / filename).write_text(text, encoding="utf-8")


def extract_darpan_blocks(page: Page, ang: int, page_timeout_ms: int) -> list[dict[str, str]]:
    url = DARPAN_URL.format(page=ang)
    print(f"  → Загружаю {url}")

    try:
        page.goto(url, wait_until="domcontentloaded", timeout=page_timeout_ms)
    except PWTimeout:
        print(f"  ✗ Таймаут при загрузке {url}")
        return []

    page.wait_for_timeout(1200)

    blocks = page.evaluate(
        """() => {
            return Array.from(document.querySelectorAll("p.Bani, p.Arath"))
                .map((el) => {
                    const kind = el.classList.contains("Bani") ? "BANI" : "ARATH";
                    const text = (el.innerText || "").trim();
                    return { kind, text };
                })
                .filter((item) => item.text);
        }"""
    )

    if not blocks:
        print(f"  ✗ Нет блоков p.Bani / p.Arath на странице {ang}")
        return []

    return blocks


def count_bani_blocks(blocks: list[dict[str, str]]) -> int:
    return sum(1 for block in blocks if block.get("kind", "").upper() == "BANI")


def build_prompt_input(blocks: list[dict[str, str]]) -> str:
    parts: list[str] = []

    for block in blocks:
        kind = block["kind"].strip().upper()
        text = normalize_text(block["text"])
        parts.append(f"[{kind}]\n{text}\n[/{kind}]")

    return "\n\n".join(parts).strip()


def build_prompt(content: str, expected_verses: int) -> str:
    return PROMPT_TEMPLATE.format(
        content=content,
        expected_verses=expected_verses,
        romanization_rules=ROMANIZATION_RULES,
        russian_glossary=RUSSIAN_GLOSSARY,
    )


def build_repair_prompt(expected_verses: int) -> str:
    return REPAIR_PROMPT_TEMPLATE.format(expected_verses=expected_verses)


def open_chat_if_needed(page: Page, chat_url: str, page_timeout_ms: int) -> None:
    current = page.url or ""
    if "chatgpt.com" in current or "chat.openai.com" in current:
        return
    page.goto(chat_url, wait_until="domcontentloaded", timeout=page_timeout_ms)
    page.wait_for_timeout(2000)


def assistant_message_count(page: Page) -> int:
    return page.locator('[data-message-author-role="assistant"]').count()


def focus_textarea(page: Page, input_timeout_ms: int):
    textarea = page.locator("#prompt-textarea")
    textarea.wait_for(state="visible", timeout=input_timeout_ms)
    textarea.click()
    return textarea


def clear_textarea(page: Page, textarea) -> None:
    modifier = "Meta" if os.name == "darwin" else "Control"
    textarea.click()
    page.keyboard.press(f"{modifier}+a")
    page.keyboard.press("Delete")
    page.wait_for_timeout(250)


def insert_prompt(page: Page, prompt: str) -> None:
    page.evaluate(
        """(text) => {
            const el = document.getElementById("prompt-textarea");
            if (!el) return;
            el.focus();
            document.execCommand("selectAll", false, null);
            document.execCommand("insertText", false, text);
        }""",
        prompt,
    )
    page.wait_for_timeout(500)


def click_send(page: Page, textarea) -> None:
    send_btn = page.locator('button[data-testid="send-button"]')
    try:
        send_btn.wait_for(state="visible", timeout=5_000)
        send_btn.click()
    except PWTimeout:
        textarea.press("Enter")


def wait_generation_finished(page: Page, response_timeout_ms: int) -> None:
    stop_btn = page.locator('button[data-testid="stop-button"]')

    try:
        stop_btn.wait_for(state="visible", timeout=20_000)
    except PWTimeout:
        pass

    try:
        stop_btn.wait_for(state="hidden", timeout=response_timeout_ms)
    except PWTimeout:
        print(f"  ⚠ Таймаут ожидания ответа ({response_timeout_ms / 1000:.0f}с) — беру то, что есть")

    page.wait_for_timeout(1200)


def click_continue_if_present(page: Page) -> bool:
    candidates = [
        page.get_by_role("button", name=re.compile(r"Continue generating", re.I)),
        page.get_by_role("button", name=re.compile(r"Продолжить", re.I)),
        page.locator("button").filter(has_text=re.compile(r"Continue generating|Продолжить", re.I)),
    ]

    for locator in candidates:
        try:
            if locator.count() > 0 and locator.first.is_visible():
                locator.first.click()
                page.wait_for_timeout(600)
                print("  → Нажимаю Continue generating...")
                return True
        except Exception:
            continue

    return False


def drain_generation(page: Page, response_timeout_ms: int, max_continues: int = 5) -> None:
    for _ in range(max_continues + 1):
        wait_generation_finished(page, response_timeout_ms)
        if not click_continue_if_present(page):
            return


def get_new_last_assistant_message(page: Page, before_count: int, new_message_timeout_ms: int) -> str | None:
    try:
        page.wait_for_function(
            """(before) => {
                return document.querySelectorAll('[data-message-author-role="assistant"]').length > before;
            }""",
            arg=before_count,
            timeout=new_message_timeout_ms,
        )
    except PWTimeout:
        pass

    messages = page.locator('[data-message-author-role="assistant"]')
    after_count = messages.count()

    if after_count == 0 or after_count <= before_count:
        return None

    last = messages.nth(after_count - 1)

    try:
        last.wait_for(state="visible", timeout=5_000)
    except PWTimeout:
        return None

    text = normalize_text(last.inner_text())
    return text or None


def send_prompt_and_get_answer(page: Page, chat_url: str, prompt: str, cfg: RuntimeConfig) -> str | None:
    open_chat_if_needed(page, chat_url, cfg.page_timeout_ms)

    try:
        textarea = focus_textarea(page, cfg.input_timeout_ms)
    except PWTimeout:
        print("  ✗ Поле ввода ChatGPT не найдено")
        return None

    before_count = assistant_message_count(page)
    clear_textarea(page, textarea)
    insert_prompt(page, prompt)
    click_send(page, textarea)

    print("  → Жду ответа ChatGPT...")
    drain_generation(page, cfg.response_timeout_ms)

    return get_new_last_assistant_message(page, before_count, cfg.new_message_timeout_ms)


def extract_json_candidate(text: str) -> str | None:
    text = text.strip()

    match = re.search(r"BEGIN_DARPAN_JSON\s*(.*?)\s*END_DARPAN_JSON", text, flags=re.S)
    if not match:
        return None

    candidate = match.group(1).strip()
    candidate = re.sub(r"^```json\s*", "", candidate, flags=re.I)
    candidate = re.sub(r"^```\s*", "", candidate)
    candidate = re.sub(r"\s*```$", "", candidate)

    if candidate.startswith("{") and candidate.endswith("}"):
        return candidate

    return None


def validate_translation_payload(data: Any, expected_verses: int) -> tuple[AngTranslation | None, str]:
    if not isinstance(data, dict):
        return None, "payload is not a dict"

    verses_raw = data.get("verses")
    commentary_raw = data.get("commentary", "")

    if not isinstance(verses_raw, list) or not verses_raw:
        return None, "verses is empty or not a list"

    if len(verses_raw) != expected_verses:
        return None, f"verse count mismatch: expected {expected_verses}, got {len(verses_raw)}"

    verses: list[Verse] = []

    for index, item in enumerate(verses_raw, start=1):
        if not isinstance(item, dict):
            return None, f"verse #{index} is not a dict"

        gurmukhi = normalize_text(str(item.get("gurmukhi", "")))
        roman = normalize_text(str(item.get("roman", "")))
        line_translation = normalize_text(str(item.get("line_translation", "")))
        arth = normalize_text(str(item.get("arth", "")))

        if not gurmukhi:
            return None, f"verse #{index}: missing gurmukhi"
        if not roman:
            return None, f"verse #{index}: missing roman"
        if not line_translation:
            return None, f"verse #{index}: missing line_translation"
        if not arth:
            return None, f"verse #{index}: missing arth"

        verses.append(
            Verse(
                gurmukhi=gurmukhi,
                roman=roman,
                line_translation=line_translation,
                arth=arth,
            )
        )

    commentary = normalize_text(str(commentary_raw))
    return AngTranslation(verses=verses, commentary=commentary), ""


def parse_structured_answer(answer: str, expected_verses: int) -> tuple[AngTranslation | None, str]:
    candidate = extract_json_candidate(answer)
    if not candidate:
        return None, "BEGIN_DARPAN_JSON / END_DARPAN_JSON not found"

    try:
        data = json.loads(candidate)
    except json.JSONDecodeError as exc:
        return None, f"json decode error: {exc}"

    return validate_translation_payload(data, expected_verses)


def request_structured_translation(
    page: Page,
    chat_url: str,
    content: str,
    ang: int,
    expected_verses: int,
    cfg: RuntimeConfig,
) -> AngTranslation | None:
    prompt = build_prompt(content, expected_verses)
    repair_prompt = build_repair_prompt(expected_verses)

    save_raw_text(cfg.raw_log_dir, f"ang_{ang:04d}_source.txt", content)
    save_raw_text(cfg.raw_log_dir, f"ang_{ang:04d}_prompt.txt", prompt)

    for attempt in range(1, cfg.max_retries + 1):
        print(f"  → Попытка {attempt}/{cfg.max_retries}")

        current_prompt = prompt if attempt == 1 else repair_prompt
        answer = send_prompt_and_get_answer(page, chat_url, current_prompt, cfg)

        if answer:
            save_raw_text(cfg.raw_log_dir, f"ang_{ang:04d}_attempt_{attempt}_answer.txt", answer)
            parsed, reason = parse_structured_answer(answer, expected_verses)
            if parsed:
                print(f"  ✓ Ответ распознан: строф {len(parsed.verses)}")
                return parsed
            print(f"  ⚠ Ответ получен, но не прошёл проверку: {reason}")
        else:
            save_raw_text(cfg.raw_log_dir, f"ang_{ang:04d}_attempt_{attempt}_answer.txt", "[NO_ANSWER]")
            print("  ⚠ Пустой ответ или не удалось получить новый ответ ассистента")

        if attempt < cfg.max_retries:
            print(f"  ⏳ Повтор через {cfg.retry_delay_s}с...")
            time.sleep(cfg.retry_delay_s)

    return None


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Darpan → ChatGPT → DOCX + JSON")

    parser.add_argument("--start", type=int, default=1, help="Начальный анг")
    parser.add_argument("--end", type=int, default=1430, help="Конечный анг")
    parser.add_argument("--output", type=str, default="darpan_chatgpt.docx", help="Имя выходного DOCX")
    parser.add_argument("--delay", type=float, default=3.0, help="Пауза между ангами в секундах")

    parser.add_argument("--max-retries", type=int, default=3, help="Максимум попыток на один анг")
    parser.add_argument("--retry-delay", type=float, default=10.0, help="Пауза между повторами в секундах")

    parser.add_argument("--page-timeout", type=int, default=30, help="Таймаут загрузки страницы в секундах")
    parser.add_argument("--input-timeout", type=int, default=15, help="Таймаут ожидания поля ввода в секундах")
    parser.add_argument("--response-timeout", type=int, default=900, help="Таймаут ожидания ответа модели в секундах")
    parser.add_argument("--new-message-timeout", type=int, default=30, help="Таймаут ожидания нового assistant-сообщения в секундах")

    parser.add_argument(
        "--chat-url",
        type=str,
        default=None,
        help="URL конкретного чата ChatGPT. Лучше использовать отдельный чистый чат.",
    )
    parser.add_argument(
        "--raw-log-dir",
        type=str,
        default="raw_logs",
        help='Папка для сырых логов prompt/answer. Пустая строка "" отключает логи.',
    )
    parser.add_argument(
        "--json-dir",
        type=str,
        default="ang_json",
        help="Папка для JSON по ангам",
    )
    parser.add_argument(
        "--rebuild-missing-docx-from-json",
        action="store_true",
        help="Пересобрать DOCX из уже сохранённых JSON в указанном диапазоне, без GPT.",
    )

    return parser.parse_args()


def main() -> None:
    args = parse_args()

    if args.start < 1 or args.end < 1 or args.start > args.end:
        print("  ✗ Некорректный диапазон: проверь --start и --end")
        return

    output_path = (Path(__file__).parent / args.output).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    progress_file = build_progress_path(output_path)

    raw_log_dir = None
    if args.raw_log_dir.strip():
        raw_log_dir = (Path(__file__).parent / args.raw_log_dir).resolve()

    json_dir = (Path(__file__).parent / args.json_dir).resolve()

    cfg = RuntimeConfig(
        page_timeout_ms=args.page_timeout * 1000,
        input_timeout_ms=args.input_timeout * 1000,
        response_timeout_ms=args.response_timeout * 1000,
        new_message_timeout_ms=args.new_message_timeout * 1000,
        max_retries=args.max_retries,
        retry_delay_s=args.retry_delay,
        raw_log_dir=raw_log_dir,
        json_dir=json_dir,
    )

    if args.rebuild_missing_docx_from_json:
        print("Пересобираю DOCX из JSON...")
        rebuilt = rebuild_docx_from_json(output_path, cfg.json_dir, args.start, args.end)
        if rebuilt > 0:
            save_progress(progress_file, args.end)
        print(f"✓ Пересобрано ангов: {rebuilt}")
        print(f"✓ Результат: {output_path}")
        print(f"✓ JSON по ангам: {cfg.json_dir}")
        return

    ensure_output_doc_exists(output_path)

    chat_url = (args.chat_url or "").strip()
    if not chat_url:
        print("URL чата не задан — открою ChatGPT по умолчанию.")
        print("Лучше использовать отдельный чистый чат без старой истории.")
        chat_url = DEFAULT_CHAT_URL

    if not chat_url.startswith("http"):
        print("  ✗ Некорректный URL, выхожу.")
        return

    BOT_PROFILE.mkdir(exist_ok=True)
    first_run = not (BOT_PROFILE / "Default").exists()

    with sync_playwright() as pw:
        print("Запускаю браузер...")
        ctx = pw.chromium.launch_persistent_context(
            user_data_dir=str(BOT_PROFILE),
            headless=False,
            args=["--start-maximized"],
            no_viewport=True,
        )

        darpan_page = ctx.new_page()
        _stealth.apply_stealth_sync(darpan_page)

        chatgpt_page = ctx.new_page()
        _stealth.apply_stealth_sync(chatgpt_page)
        chatgpt_page.goto(chat_url, wait_until="domcontentloaded", timeout=cfg.page_timeout_ms)

        if first_run:
            print("\n⚠ Первый запуск: войди в ChatGPT в открывшемся браузере.")
            print("После входа нажми Enter здесь...")
            input()
        else:
            chatgpt_page.wait_for_timeout(2000)
            print("ChatGPT открыт.\n")

        saved = load_progress(progress_file)
        if saved:
            print(f"Последний сохранённый прогресс: анг {saved}")
        print(f"Работаю строго по диапазону: {args.start}..{args.end}")

        for ang in range(args.start, args.end + 1):
            print(f"\n══ Анг {ang}/{args.end} ══")

            existing_json = ang_json_path(cfg.json_dir, ang)
            if existing_json.exists():
                print(f"  ↷ JSON уже существует, пропускаю анг {ang}: {existing_json.name}")
                continue

            blocks = extract_darpan_blocks(darpan_page, ang, cfg.page_timeout_ms)
            if not blocks:
                print(f"  ⚠ Пропускаю анг {ang}")
                continue

            expected_verses = count_bani_blocks(blocks)
            if expected_verses == 0:
                print(f"  ⚠ На странице нет блоков [BANI], пропускаю анг {ang}")
                continue

            print(f"  → Ожидаю строф: {expected_verses}")
            source_text = build_prompt_input(blocks)

            chatgpt_page.bring_to_front()
            ang_data = request_structured_translation(
                chatgpt_page,
                chat_url,
                source_text,
                ang,
                expected_verses,
                cfg,
            )

            if not ang_data:
                print(f"  ⚠ Не удалось получить валидный результат для анга {ang}, пропускаю")
                continue

            json_path = save_ang_json(cfg.json_dir, ang, ang_data)
            print(f"  ✓ JSON сохранён: {json_path.name}")

            append_ang_to_docx(output_path, ang, ang_data)
            save_progress(progress_file, ang)

            if ang < args.end:
                print(f"  ⏳ Пауза {args.delay}с...")
                time.sleep(args.delay)

        print(f"\n✓ Готово! Результат: {output_path}")
        print(f"✓ Прогресс: {progress_file}")
        print(f"✓ JSON по ангам: {cfg.json_dir}")
        if cfg.raw_log_dir:
            print(f"✓ Сырые логи: {cfg.raw_log_dir}")

        ctx.close()


if __name__ == "__main__":
    main()
