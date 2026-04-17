#!/usr/bin/env python3
# chatgpt_khojgurbani_sahibsingh_bot.py

from __future__ import annotations

import argparse
import json
import os
import re
import time
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from playwright.sync_api import Page, TimeoutError as PWTimeout, sync_playwright
from playwright_stealth import Stealth

_stealth = Stealth()

API_BASE = "https://apiprod.khojgurbani.org/api/v1"
HEADERS = {
    "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36",
    "Origin": "https://khojgurbani.org",
    "Referer": "https://khojgurbani.org/",
    "Accept": "application/json",
}

DEFAULT_CHAT_URL = "https://chatgpt.com/"
BOT_PROFILE = Path(__file__).parent / "bot_profile"

TITLE_TEXT = "Шри Гуру Грантх Сахиб — перевод на русский"
SUBTITLE_TEXT = "Источник: KhojGurbani · Prof. Sahib Singh"

COLOR_GURBANI = RGBColor(0x55, 0x00, 0x00)
COLOR_ROMAN = RGBColor(0x44, 0x44, 0x44)
COLOR_TRANSLATION = RGBColor(0x00, 0x60, 0x00)
COLOR_LABEL = RGBColor(0x99, 0x99, 0x99)
COLOR_HEADING = RGBColor(0x22, 0x22, 0x22)

ROMANIZATION_RULES = """\
Требования к полю "roman":
- нужна аккуратная панджабская romanization для русскоязычного читателя;
- сохраняй диакритику там, где это важно: ā, ī, ū, ṭ, ṭh, ḍ, ḍh, ṇ, ṛ, ñ, ś, ṃ, ṅ;
- не делай механическую школьную санскритизацию;
- roman должна читаться ближе к живому панджабскому чтению, а не как буквенная ловушка;
- не тащи автоматически конечные -u / -i, если в принятом чтении они не звучат отдельно;
- для этого проекта букву ਚ передавай как "ch", а не как "c";
- букву ਛ передавай как "chh";
- для форм вроде ਸੋਚੈ не пиши "soce"; пиши ближе к чтению: "sochey";
- для форм вроде ਗਾਵੈ не пиши "gāvai"; пиши ближе к чтению: "gāvē";
- используй практичную научную romanization с учётом чтения.

Предпочтительные примеры:
- ਸੋਚੈ → sochey
- ਵਿਚਾਰ → vichār
- ਚਾਰ → chār
- ਚੁਪ → chup
- ਸਚੁ → sach
- ਛੇ → chhē
- ਗਾਵੈ → gāvē
- ਹੁਕਮਿ → hukam
- ਸਤਿਨਾਮੁ → sat nām
- ਜਪੁ → jap
- ਗੁਰ ਪ੍ਰਸਾਦਿ → gur prasād
- ਅਜੂਨੀ → ajūnī
- ਸੈਭੰ → saibhaṅ
- ਕਰਤਾ ਪੁਰਖੁ → kartā purakh
- ਨਿਰਭਉ → nirbhau
- ਨਿਰਵੈਰੁ → nirvair
"""

RUSSIAN_GLOSSARY = """\
Предпочтительные русские формы:
- ਹਉਮੈ → Хоумэ
- ਜਪੁ → Джап
- ਜਪੁ ਜੀ → Джап Джи
- ਸਤਿਗੁਰੁ / ਸਤਿਗੁਰੂ → Сатгуру
- ਸਤਿਨਾਮੁ → Сат Нам
- ਏਕੰਕਾਰੁ / ਓਅੰਕਾਰੁ → Онкар / Эконкар по контексту, без механического финального "у"
- ਗੁਰ ਪ੍ਰਸਾਦਿ → Гур прасад
- ਅਜੂਨੀ → Аджуни
- ਸੈਭੰ → Сэйбханг
- ਕਰਤਾ ਪੁਰਖੁ → Карта Пуракх
- ਨਿਰਭਉ → Нирбхау
- ਨਿਰਵੈਰੁ → Нирвэр

Общее правило:
- roman: аккуратно и единообразно;
- русский текст: естественно, по панджабскому чтению и смыслу;
- переводить только смысл Prof. Sahib Singh из поля sahib_singh_pa;
- не добавлять свои комментарии, толкования и расширения.
"""

PROMPT_TEMPLATE = """\
Игнорируй весь предыдущий контекст этого чата.
Используй только текст текущего сообщения.

Это один анг из KhojGurbani.
Для каждой строки у тебя есть:
- verse_id
- gurmukhi
- sahib_singh_pa

Нужно вернуть материал строго по строкам этого анга.
Источник смысла для "translation_ru" — только поле "sahib_singh_pa" у соответствующей строки.
Нельзя опираться на внешний контекст, на другие переводы, на собственные догадки и на комментарии от себя.

В этом анге ровно {expected_lines} строк.
Верни ровно {expected_lines} элементов в массиве "lines" и сохрани порядок строк.

Для каждой строки верни:
1) "verse_id" — тот же integer, что во входе;
2) "roman" — аккуратная romanization строки gurmukhi;
3) "translation_ru" — только русский перевод панджабского перевода Prof. Sahib Singh из поля "sahib_singh_pa", без добавлений от себя.

Если для строки на сайте нет перевода Prof. Sahib Singh или переводная строка реально отсутствует, оставь "translation_ru" пустой строкой "".
Это допустимо. Не выдумывай перевод для таких строк.

Не добавляй:
- пояснений;
- комментариев;
- markdown;
- code fences;
- лишних полей.

{romanization_rules}

{russian_glossary}

Правила:
- сохраняй порядок строк;
- не пропускай строки;
- не меняй verse_id;
- если строка короткая, не раздувай перевод;
- не вставляй заголовки;
- ответ верни строго между BEGIN_KG_JSON и END_KG_JSON;
- ВАЖНО: никогда не используй ASCII-кавычки " " внутри значений строк — это ломает JSON; если нужно выделить слово или цитату, используй «ёлочки» или одинарные кавычки.

Формат ответа:

BEGIN_KG_JSON
{{
  "ang": {ang},
  "lines": [
    {{
      "verse_id": 123,
      "roman": "...",
      "translation_ru": "..."
    }}
  ]
}}
END_KG_JSON

Текст анга:

{content}
"""

REPAIR_PROMPT_TEMPLATE = """\
Игнорируй весь предыдущий контекст этого чата, кроме своего последнего ответа.
Используй только смысл своего последнего ответа и перепиши его в правильном формате.

Должно быть ровно {expected_lines} элементов в массиве "lines".

Нужен только валидный JSON между BEGIN_KG_JSON и END_KG_JSON.
Без markdown, без code fences, без пояснений.

Обязательные поля для каждой строки:
- verse_id
- roman
- translation_ru

Критично:
- не меняй verse_id;
- не добавляй новые поля;
- translation_ru должен быть только переводом поля sahib_singh_pa без собственных добавлений;
- если у строки нет перевода на сайте, допускается пустая строка "";
- никогда не используй ASCII-кавычки " " внутри значений строк — это ломает JSON; используй «ёлочки» или одинарные кавычки;
- roman должна учитывать проектную норму:
  - ਚ = ch
  - ਛ = chh
  - sochey
  - vichār
  - chār
  - chup
  - sach
  - chhē
  - gāvē
  - hukam
  - sat nām
  - gur prasād
  - saibhaṅ

Формат:

BEGIN_KG_JSON
{{
  "lines": [
    {{
      "verse_id": 123,
      "roman": "...",
      "translation_ru": "..."
    }}
  ]
}}
END_KG_JSON
"""


@dataclass
class SourceLine:
    index: int
    verse_id: int
    shabad_num: int
    shabad_id: int | None
    gurmukhi: str
    site_roman: str
    sahib_singh_pa: str


@dataclass
class OutputLine:
    index: int
    verse_id: int
    shabad_num: int
    shabad_id: int | None
    gurmukhi: str
    site_roman: str
    sahib_singh_pa: str
    roman: str
    translation_ru: str


@dataclass
class AngTranslation:
    ang: int
    lines: list[OutputLine]


@dataclass
class RuntimeConfig:
    page_timeout_ms: int
    input_timeout_ms: int
    response_timeout_ms: int
    new_message_timeout_ms: int
    max_retries: int
    retry_delay_s: float
    raw_log_dir: Path | None
    json_dir: Path
    keep_chat_tabs: bool


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def safe_filename_part(name: str) -> str:
    return re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", name)


def build_progress_path(output_path: Path) -> Path:
    return output_path.parent / f".progress_{safe_filename_part(output_path.name)}.txt"


def load_progress(progress_file: Path) -> int:
    if not progress_file.exists():
        return 0
    try:
        return int(progress_file.read_text(encoding="utf-8").strip() or "0")
    except ValueError:
        return 0


def save_progress(progress_file: Path, ang: int) -> None:
    progress_file.write_text(str(ang), encoding="utf-8")


def reset_progress(progress_file: Path) -> None:
    if progress_file.exists():
        progress_file.unlink()


def save_raw_text(raw_log_dir: Path | None, filename: str, text: str) -> None:
    if raw_log_dir is None:
        return
    raw_log_dir.mkdir(parents=True, exist_ok=True)
    (raw_log_dir / filename).write_text(text, encoding="utf-8")


def api_get(path: str) -> dict[str, Any]:
    req = urllib.request.Request(f"{API_BASE}{path}", headers=HEADERS)
    with urllib.request.urlopen(req, timeout=20) as response:
        return json.loads(response.read())


def fetch_shabad(ang: int, shabad_num: int, *, quiet: bool = False) -> dict[str, Any] | None:
    try:
        data = api_get(f"/shabad/{ang}/{shabad_num}")
        if data.get("status") == "success":
            payload = data.get("data")
            if isinstance(payload, dict):
                return payload
    except Exception as exc:
        if not quiet:
            print(f"  [!] fetch_shabad({ang}/{shabad_num}): {exc}")
    return None


def _normalize_int_list(values: Any) -> list[int]:
    result: list[int] = []
    if not isinstance(values, list):
        return result

    for item in values:
        try:
            result.append(int(item))
        except Exception:
            continue

    return sorted(set(result))


def discover_shabad_numbers_for_ang(ang: int, max_probe: int = 1500) -> list[int]:
    found_by_scan: list[int] = []
    misses_after_first = 0

    for probe in range(1, max_probe + 1):
        shabad = fetch_shabad(ang, probe, quiet=True)
        if shabad:
            pages = _normalize_int_list(shabad.get("pages"))
            if pages:
                return pages

            found_by_scan.append(probe)
            misses_after_first = 0
            continue

        if found_by_scan:
            misses_after_first += 1
            if misses_after_first >= 30:
                break

    return found_by_scan


def fetch_ang_source_lines(ang: int) -> list[SourceLine]:
    shabad_numbers = discover_shabad_numbers_for_ang(ang)
    if not shabad_numbers:
        print(f"  ✗ Не удалось определить шабады для анга {ang}")
        return []

    lines: list[SourceLine] = []
    seen_verse_ids: set[int] = set()

    for shabad_num in shabad_numbers:
        shabad = fetch_shabad(ang, shabad_num)
        if not shabad:
            continue

        shabad_id = shabad.get("id")
        scriptures = shabad.get("scriptures", [])
        if not isinstance(scriptures, list):
            continue

        for verse in scriptures:
            try:
                verse_id = int(verse["id"])
            except Exception:
                continue

            if verse_id in seen_verse_ids:
                continue

            gurmukhi = normalize_text(str(verse.get("Scripture", "")))
            site_roman = normalize_text(str(verse.get("ScriptureRoman", "")))
            translations = verse.get("translation", {}) or {}
            sahib_singh_pa = normalize_text(str(translations.get("SahibSinghPunjabi", "")))

            if not gurmukhi or not sahib_singh_pa:
                continue

            seen_verse_ids.add(verse_id)
            lines.append(
                SourceLine(
                    index=len(lines) + 1,
                    verse_id=verse_id,
                    shabad_num=shabad_num,
                    shabad_id=int(shabad_id) if isinstance(shabad_id, int) else None,
                    gurmukhi=gurmukhi,
                    site_roman=site_roman,
                    sahib_singh_pa=sahib_singh_pa,
                )
            )

    return lines


def build_prompt_input(lines: list[SourceLine]) -> str:
    parts: list[str] = []

    for line in lines:
        parts.append(
            "\n".join(
                [
                    "[LINE]",
                    f"index: {line.index}",
                    f"verse_id: {line.verse_id}",
                    f"gurmukhi: {line.gurmukhi}",
                    f"sahib_singh_pa: {line.sahib_singh_pa}",
                    "[/LINE]",
                ]
            )
        )

    return "\n\n".join(parts).strip()


def build_prompt(ang: int, lines: list[SourceLine]) -> str:
    return PROMPT_TEMPLATE.format(
        ang=ang,
        expected_lines=len(lines),
        romanization_rules=ROMANIZATION_RULES,
        russian_glossary=RUSSIAN_GLOSSARY,
        content=build_prompt_input(lines),
    )


def build_repair_prompt(expected_lines: int) -> str:
    return REPAIR_PROMPT_TEMPLATE.format(expected_lines=expected_lines)


def add_run(
    para,
    text: str,
    *,
    color=None,
    bold: bool = False,
    italic: bool = False,
    size_pt: int | None = None,
):
    run = para.add_run(text)
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def ensure_output_doc_exists(docx_path: Path) -> None:
    if docx_path.exists():
        return

    doc = Document()
    doc.core_properties.author = "KhojGurbani Sahib Singh Bot"

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title = doc.add_heading(TITLE_TEXT, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(SUBTITLE_TEXT)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = COLOR_LABEL

    doc.add_paragraph()
    doc.save(str(docx_path))


def reset_output_doc(docx_path: Path) -> None:
    if docx_path.exists():
        docx_path.unlink()
    ensure_output_doc_exists(docx_path)


def add_ang_heading(doc: Document, ang: int) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run(f"АНГ {ang}")
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_HEADING


def add_centered_line_block(doc: Document, line: OutputLine) -> None:
    p_g = doc.add_paragraph()
    p_g.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_g.paragraph_format.space_before = Pt(8)
    p_g.paragraph_format.space_after = Pt(2)
    add_run(p_g, line.gurmukhi, color=COLOR_GURBANI, size_pt=13, bold=True)

    p_r = doc.add_paragraph()
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r.paragraph_format.space_after = Pt(3)
    add_run(p_r, line.roman, color=COLOR_ROMAN, size_pt=10, italic=True)

    if normalize_text(line.translation_ru):
        p_t = doc.add_paragraph()
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_t.paragraph_format.space_after = Pt(10)
        add_run(p_t, line.translation_ru, color=COLOR_TRANSLATION, size_pt=11)

    doc.add_paragraph()


def append_ang_to_docx(docx_path: Path, ang_data: AngTranslation) -> None:
    ensure_output_doc_exists(docx_path)
    doc = Document(str(docx_path))

    add_ang_heading(doc, ang_data.ang)

    for line in ang_data.lines:
        add_centered_line_block(doc, line)

    doc.save(str(docx_path))
    print(f"  ✓ DOCX дописан: {docx_path.name}")


def ang_to_dict(ang_data: AngTranslation) -> dict[str, Any]:
    return {
        "ang": ang_data.ang,
        "line_count": len(ang_data.lines),
        "lines": [
            {
                "index": line.index,
                "verse_id": line.verse_id,
                "shabad_num": line.shabad_num,
                "shabad_id": line.shabad_id,
                "gurmukhi": line.gurmukhi,
                "site_roman": line.site_roman,
                "sahib_singh_pa": line.sahib_singh_pa,
                "roman": line.roman,
                "translation_ru": line.translation_ru,
            }
            for line in ang_data.lines
        ],
    }


def output_line_from_dict(item: dict[str, Any]) -> OutputLine:
    shabad_id_raw = item.get("shabad_id")
    shabad_id = int(shabad_id_raw) if isinstance(shabad_id_raw, int) else None

    return OutputLine(
        index=int(item.get("index", 0)),
        verse_id=int(item.get("verse_id", 0)),
        shabad_num=int(item.get("shabad_num", 0)),
        shabad_id=shabad_id,
        gurmukhi=normalize_text(str(item.get("gurmukhi", ""))),
        site_roman=normalize_text(str(item.get("site_roman", ""))),
        sahib_singh_pa=normalize_text(str(item.get("sahib_singh_pa", ""))),
        roman=normalize_text(str(item.get("roman", ""))),
        translation_ru=normalize_text(str(item.get("translation_ru", ""))),
    )


def ang_translation_from_dict(data: dict[str, Any]) -> AngTranslation:
    raw_lines = data.get("lines", [])
    lines: list[OutputLine] = []
    if isinstance(raw_lines, list):
        for item in raw_lines:
            if not isinstance(item, dict):
                continue
            lines.append(output_line_from_dict(item))

    return AngTranslation(ang=int(data.get("ang", 0)), lines=lines)


def ang_json_path(json_dir: Path, ang: int) -> Path:
    return json_dir / f"ang_{ang:04d}.json"


def save_ang_json(json_dir: Path, ang_data: AngTranslation) -> Path:
    json_dir.mkdir(parents=True, exist_ok=True)
    json_path = ang_json_path(json_dir, ang_data.ang)
    json_path.write_text(
        json.dumps(ang_to_dict(ang_data), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return json_path


def load_ang_json(json_dir: Path, ang: int) -> AngTranslation | None:
    json_path = ang_json_path(json_dir, ang)
    if not json_path.exists():
        return None

    try:
        data = json.loads(json_path.read_text(encoding="utf-8"))
    except Exception:
        return None

    return ang_translation_from_dict(data)


def reset_json_range(json_dir: Path, start: int, end: int) -> int:
    removed = 0
    for ang in range(start, end + 1):
        json_path = ang_json_path(json_dir, ang)
        if json_path.exists():
            json_path.unlink()
            removed += 1
    return removed


def rebuild_docx_from_json(output_path: Path, json_dir: Path, start: int, end: int) -> int:
    reset_output_doc(output_path)
    rebuilt = 0

    for ang in range(start, end + 1):
        ang_data = load_ang_json(json_dir, ang)
        if not ang_data or not ang_data.lines:
            print(f"  ↷ Нет JSON для анга {ang}, пропускаю")
            continue

        append_ang_to_docx(output_path, ang_data)
        rebuilt += 1

    return rebuilt


_META_GUESS_PATTERNS = [
    re.compile(r"\bя\s+думаю\b", re.IGNORECASE),
    re.compile(r"\bмне\s+кажется\b", re.IGNORECASE),
    re.compile(r"\bвероятно\b", re.IGNORECASE),
    re.compile(r"\bпохоже\b", re.IGNORECASE),
    re.compile(r"\bпо-видимому\b", re.IGNORECASE),
]


def repair_json_quotes(text: str) -> str:
    """Fix unescaped double-quotes inside JSON string values.

    ChatGPT sometimes writes "я" or "слово" inside a translation string,
    breaking the JSON. This scanner detects those and escapes them.
    """
    result: list[str] = []
    i = 0
    in_string = False
    escape_next = False

    while i < len(text):
        c = text[i]

        if escape_next:
            result.append(c)
            escape_next = False
            i += 1
            continue

        if c == "\\":
            result.append(c)
            escape_next = True
            i += 1
            continue

        if c == '"':
            if not in_string:
                in_string = True
                result.append(c)
            else:
                # Determine if this quote closes the current string or is an inner quote.
                # Look ahead past whitespace to see what follows.
                j = i + 1
                while j < len(text) and text[j] in " \t\n\r":
                    j += 1
                next_ch = text[j] if j < len(text) else ""
                if next_ch in (",", "}", "]", ":"):
                    in_string = False
                    result.append(c)
                else:
                    # Unescaped inner quote — escape it
                    result.append("\\")
                    result.append(c)
            i += 1
            continue

        result.append(c)
        i += 1

    return "".join(result)


def looks_like_model_guess(text: str) -> bool:
    sample = normalize_text(text)
    if not sample:
        return False
    return any(pattern.search(sample) for pattern in _META_GUESS_PATTERNS)


def open_chat_tab(context, chat_url: str, page_timeout_ms: int) -> Page:
    page = context.new_page()
    _stealth.apply_stealth_sync(page)
    page.goto(chat_url, wait_until="domcontentloaded", timeout=page_timeout_ms)
    page.wait_for_timeout(2000)
    return page


def assistant_message_count(page: Page) -> int:
    return page.locator('[data-message-author-role="assistant"]').count()


def focus_textarea(page: Page, input_timeout_ms: int):
    textarea = page.locator("#prompt-textarea")
    textarea.wait_for(state="visible", timeout=input_timeout_ms)
    textarea.click()
    return textarea


def clear_textarea(page: Page, textarea) -> None:
    modifier = "Meta" if os.name == "darwin" else "Control"
    textarea.click()
    page.keyboard.press(f"{modifier}+a")
    page.keyboard.press("Delete")
    page.wait_for_timeout(250)


def insert_prompt(page: Page, prompt: str) -> None:
    page.evaluate(
        """(text) => {
            const el = document.getElementById("prompt-textarea");
            if (!el) return;
            el.focus();
            document.execCommand("selectAll", false, null);
            document.execCommand("insertText", false, text);
        }""",
        prompt,
    )
    page.wait_for_timeout(500)


def click_send(page: Page, textarea) -> None:
    send_btn = page.locator('button[data-testid="send-button"]')
    try:
        send_btn.wait_for(state="visible", timeout=5_000)
        send_btn.click()
    except PWTimeout:
        textarea.press("Enter")


def wait_generation_finished(page: Page, response_timeout_ms: int) -> None:
    stop_btn = page.locator('button[data-testid="stop-button"]')

    try:
        stop_btn.wait_for(state="visible", timeout=20_000)
    except PWTimeout:
        pass

    try:
        stop_btn.wait_for(state="hidden", timeout=response_timeout_ms)
    except PWTimeout:
        print(f"  ⚠ Таймаут ожидания ответа ({response_timeout_ms / 1000:.0f}с) — беру то, что есть")

    page.wait_for_timeout(1200)


def click_continue_if_present(page: Page) -> bool:
    candidates = [
        page.get_by_role("button", name=re.compile(r"Continue generating", re.I)),
        page.get_by_role("button", name=re.compile(r"Продолжить", re.I)),
        page.locator("button").filter(has_text=re.compile(r"Continue generating|Продолжить", re.I)),
    ]

    for locator in candidates:
        try:
            if locator.count() > 0 and locator.first.is_visible():
                locator.first.click()
                page.wait_for_timeout(600)
                print("  → Нажимаю Continue generating...")
                return True
        except Exception:
            continue

    return False


def drain_generation(page: Page, response_timeout_ms: int, max_continues: int = 5) -> None:
    for _ in range(max_continues + 1):
        wait_generation_finished(page, response_timeout_ms)
        if not click_continue_if_present(page):
            return


def get_new_last_assistant_message(page: Page, before_count: int, new_message_timeout_ms: int) -> str | None:
    try:
        page.wait_for_function(
            """(before) => {
                return document.querySelectorAll('[data-message-author-role="assistant"]').length > before;
            }""",
            arg=before_count,
            timeout=new_message_timeout_ms,
        )
    except PWTimeout:
        pass

    messages = page.locator('[data-message-author-role="assistant"]')
    after_count = messages.count()

    if after_count == 0 or after_count <= before_count:
        return None

    last = messages.nth(after_count - 1)
    try:
        last.wait_for(state="visible", timeout=5_000)
    except PWTimeout:
        return None

    text = normalize_text(last.inner_text())
    return text or None


def send_prompt_and_get_answer(page: Page, prompt: str, cfg: RuntimeConfig) -> str | None:
    try:
        textarea = focus_textarea(page, cfg.input_timeout_ms)
    except PWTimeout:
        print("  ✗ Поле ввода ChatGPT не найдено")
        return None

    before_count = assistant_message_count(page)
    clear_textarea(page, textarea)
    insert_prompt(page, prompt)
    click_send(page, textarea)

    print("  → Жду ответа ChatGPT...")
    drain_generation(page, cfg.response_timeout_ms)

    return get_new_last_assistant_message(page, before_count, cfg.new_message_timeout_ms)


def extract_json_candidate(text: str) -> str | None:
    text = text.strip()

    match = re.search(r"BEGIN_KG_JSON\s*(.*?)\s*END_KG_JSON", text, flags=re.S)
    if not match:
        return None

    candidate = match.group(1).strip()
    candidate = re.sub(r"^```json\s*", "", candidate, flags=re.I)
    candidate = re.sub(r"^```\s*", "", candidate)
    candidate = re.sub(r"\s*```$", "", candidate)

    if candidate.startswith("{") and candidate.endswith("}"):
        return candidate

    return None


def merge_model_lines(
    ang: int,
    source_lines: list[SourceLine],
    model_lines: list[dict[str, Any]],
) -> tuple[AngTranslation | None, str]:
    if len(model_lines) != len(source_lines):
        return None, f"line count mismatch: expected {len(source_lines)}, got {len(model_lines)}"

    expected_ids = [line.verse_id for line in source_lines]
    received_ids: list[int] = []
    merged_lines: list[OutputLine] = []

    for source, item in zip(source_lines, model_lines):
        if not isinstance(item, dict):
            return None, f"line #{source.index} is not a dict"

        try:
            verse_id = int(item.get("verse_id"))
        except Exception:
            return None, f"line #{source.index}: invalid verse_id"

        received_ids.append(verse_id)
        if verse_id != source.verse_id:
            return None, f"line #{source.index}: verse_id mismatch, expected {source.verse_id}, got {verse_id}"

        roman = normalize_text(str(item.get("roman", "")))
        translation_ru = normalize_text(str(item.get("translation_ru", "")))

        if not roman:
            return None, f"line #{source.index}: missing roman"

        merged_lines.append(
            OutputLine(
                index=source.index,
                verse_id=source.verse_id,
                shabad_num=source.shabad_num,
                shabad_id=source.shabad_id,
                gurmukhi=source.gurmukhi,
                site_roman=source.site_roman,
                sahib_singh_pa=source.sahib_singh_pa,
                roman=roman,
                translation_ru=translation_ru,
            )
        )

    if received_ids != expected_ids:
        return None, "line order mismatch"

    return AngTranslation(ang=ang, lines=merged_lines), ""


def parse_structured_answer(answer: str, ang: int, source_lines: list[SourceLine]) -> tuple[AngTranslation | None, str]:
    candidate = extract_json_candidate(answer)
    if not candidate:
        return None, "BEGIN_KG_JSON / END_KG_JSON not found"

    try:
        data = json.loads(candidate)
    except json.JSONDecodeError:
        repaired = repair_json_quotes(candidate)
        try:
            data = json.loads(repaired)
        except json.JSONDecodeError as exc:
            return None, f"json decode error: {exc}"

    if not isinstance(data, dict):
        return None, "payload is not a dict"

    lines = data.get("lines")
    if not isinstance(lines, list) or not lines:
        return None, "lines is empty or not a list"

    payload_ang = data.get("ang")
    if payload_ang not in (None, ang):
        try:
            if int(payload_ang) != ang:
                return None, f"ang mismatch: expected {ang}, got {payload_ang}"
        except Exception:
            return None, f"invalid ang value: {payload_ang}"

    return merge_model_lines(ang, source_lines, lines)


def request_structured_translation(
    page: Page,
    ang: int,
    source_lines: list[SourceLine],
    cfg: RuntimeConfig,
) -> AngTranslation | None:
    prompt = build_prompt(ang, source_lines)
    repair_prompt = build_repair_prompt(len(source_lines))

    save_raw_text(cfg.raw_log_dir, f"ang_{ang:04d}_source_lines.json", json.dumps(
        [
            {
                "index": line.index,
                "verse_id": line.verse_id,
                "shabad_num": line.shabad_num,
                "shabad_id": line.shabad_id,
                "gurmukhi": line.gurmukhi,
                "site_roman": line.site_roman,
                "sahib_singh_pa": line.sahib_singh_pa,
            }
            for line in source_lines
        ],
        ensure_ascii=False,
        indent=2,
    ))
    save_raw_text(cfg.raw_log_dir, f"ang_{ang:04d}_prompt.txt", prompt)

    for attempt in range(1, cfg.max_retries + 1):
        print(f"  → Попытка {attempt}/{cfg.max_retries}")
        current_prompt = prompt if attempt == 1 else repair_prompt
        answer = send_prompt_and_get_answer(page, current_prompt, cfg)

        if answer:
            save_raw_text(cfg.raw_log_dir, f"ang_{ang:04d}_attempt_{attempt}_answer.txt", answer)
            parsed, reason = parse_structured_answer(answer, ang, source_lines)
            if parsed:
                print(f"  ✓ Ответ распознан: строк {len(parsed.lines)}")
                return parsed
            print(f"  ⚠ Ответ получен, но не прошёл проверку: {reason}")
        else:
            save_raw_text(cfg.raw_log_dir, f"ang_{ang:04d}_attempt_{attempt}_answer.txt", "[NO_ANSWER]")
            print("  ⚠ Пустой ответ или не удалось получить новый ответ ассистента")

        if attempt < cfg.max_retries:
            print(f"  ⏳ Повтор через {cfg.retry_delay_s}с...")
            time.sleep(cfg.retry_delay_s)

    return None


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="KhojGurbani → ChatGPT → JSON + DOCX (по ангам, только Sahib Singh)")

    parser.add_argument("--start", type=int, default=1, help="Начальный анг")
    parser.add_argument("--end", type=int, default=1430, help="Конечный анг")
    parser.add_argument("--output", type=str, default="khojgurbani_sahibsingh_chatgpt.docx", help="Имя выходного DOCX")
    parser.add_argument("--delay", type=float, default=3.0, help="Пауза между ангами в секундах")

    parser.add_argument("--max-retries", type=int, default=3, help="Максимум попыток на один анг")
    parser.add_argument("--retry-delay", type=float, default=10.0, help="Пауза между повторами в секундах")

    parser.add_argument("--page-timeout", type=int, default=30, help="Таймаут загрузки страницы в секундах")
    parser.add_argument("--input-timeout", type=int, default=15, help="Таймаут ожидания поля ввода в секундах")
    parser.add_argument("--response-timeout", type=int, default=900, help="Таймаут ожидания ответа модели в секундах")
    parser.add_argument("--new-message-timeout", type=int, default=30, help="Таймаут ожидания нового assistant-сообщения в секундах")

    parser.add_argument(
        "--chat-url",
        type=str,
        default=None,
        help="URL ChatGPT. Обычно достаточно https://chatgpt.com/ — для каждого анга откроется новая вкладка.",
    )
    parser.add_argument(
        "--raw-log-dir",
        type=str,
        default="raw_logs",
        help='Папка для сырых логов prompt/answer. Пустая строка "" отключает логи.',
    )
    parser.add_argument(
        "--json-dir",
        type=str,
        default="ang_json",
        help="Папка для JSON-дампа по ангам",
    )
    parser.add_argument(
        "--rebuild-docx-from-json",
        action="store_true",
        help="Пересобрать DOCX из уже сохранённых JSON в указанном диапазоне, без GPT.",
    )
    parser.add_argument(
        "--reset-progress",
        action="store_true",
        help="Сбросить .progress_<output>.txt перед запуском.",
    )
    parser.add_argument(
        "--force-retranslate",
        action="store_true",
        help="Не пропускать анги из-за уже существующих JSON; перевести диапазон заново и перезаписать JSON.",
    )
    parser.add_argument(
        "--reset-json-range",
        action="store_true",
        help="Перед запуском удалить JSON текущего диапазона --start..--end.",
    )
    parser.add_argument(
        "--keep-chat-tabs",
        action="store_true",
        help="Не закрывать вкладки ChatGPT после обработки анга.",
    )
    parser.add_argument(
        "--scan-missing",
        action="store_true",
        help=(
            "Сканировать диапазон --start..--end, показать отсутствующие JSON, "
            "спросить подтверждение и дополнить их через ChatGPT."
        ),
    )

    return parser.parse_args()


def main() -> None:
    args = parse_args()

    if args.start < 1 or args.end < 1 or args.start > args.end:
        print("  ✗ Некорректный диапазон: проверь --start и --end")
        return

    output_path = (Path(__file__).parent / args.output).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    progress_file = build_progress_path(output_path)
    if args.reset_progress:
        reset_progress(progress_file)

    raw_log_dir = None
    if args.raw_log_dir.strip():
        raw_log_dir = (Path(__file__).parent / args.raw_log_dir).resolve()

    json_dir = (Path(__file__).parent / args.json_dir).resolve()

    cfg = RuntimeConfig(
        page_timeout_ms=args.page_timeout * 1000,
        input_timeout_ms=args.input_timeout * 1000,
        response_timeout_ms=args.response_timeout * 1000,
        new_message_timeout_ms=args.new_message_timeout * 1000,
        max_retries=args.max_retries,
        retry_delay_s=args.retry_delay,
        raw_log_dir=raw_log_dir,
        json_dir=json_dir,
        keep_chat_tabs=args.keep_chat_tabs,
    )

    if args.reset_json_range:
        removed_json = reset_json_range(cfg.json_dir, args.start, args.end)
        print(f"Сброс JSON в диапазоне {args.start}..{args.end}: удалено {removed_json}")

    if args.scan_missing:
        missing = [
            ang for ang in range(args.start, args.end + 1)
            if not ang_json_path(cfg.json_dir, ang).exists()
        ]
        if not missing:
            print(f"✓ В диапазоне {args.start}..{args.end} все JSON на месте.")
            return
        print(f"Пропущено ангов ({len(missing)}) в диапазоне {args.start}..{args.end}:")
        print("  " + ", ".join(str(a) for a in missing))
        answer = input("\nДополнить через ChatGPT? [y/N] ").strip().lower()
        if answer != "y":
            print("Отменено.")
            return
        # Override start/end to cover only missing angs as a contiguous fill.
        # We keep the full range but skip existing JSONs (default behavior).
        args.start = missing[0]
        args.end = missing[-1]
        print(f"Запускаю перевод пропущенных ангов {args.start}..{args.end}...\n")

    if args.rebuild_docx_from_json:
        print("Пересобираю DOCX из JSON...")
        rebuilt = rebuild_docx_from_json(output_path, cfg.json_dir, args.start, args.end)
        if rebuilt > 0:
            save_progress(progress_file, args.end)
        print(f"✓ Пересобрано ангов: {rebuilt}")
        print(f"✓ Результат: {output_path}")
        print(f"✓ JSON по ангам: {cfg.json_dir}")
        return

    ensure_output_doc_exists(output_path)

    chat_url = (args.chat_url or "").strip()
    if not chat_url:
        print("URL чата не задан — открою ChatGPT по умолчанию.")
        chat_url = DEFAULT_CHAT_URL

    if not chat_url.startswith("http"):
        print("  ✗ Некорректный URL, выхожу.")
        return

    BOT_PROFILE.mkdir(exist_ok=True)
    first_run = not (BOT_PROFILE / "Default").exists()

    with sync_playwright() as pw:
        print("Запускаю браузер...")
        context = pw.chromium.launch_persistent_context(
            user_data_dir=str(BOT_PROFILE),
            headless=False,
            args=["--start-maximized"],
            no_viewport=True,
        )

        if first_run:
            login_page = open_chat_tab(context, chat_url, cfg.page_timeout_ms)
            input("\n>> Press ENTER when ChatGPT is ready (logged in, new chat open)…\n")
            if not cfg.keep_chat_tabs:
                login_page.close()
        else:
            warmup_page = open_chat_tab(context, chat_url, cfg.page_timeout_ms)
            print("ChatGPT открыт.\n")
            if not cfg.keep_chat_tabs:
                warmup_page.close()

        saved = load_progress(progress_file)
        if saved:
            print(f"Последний сохранённый прогресс: анг {saved}")
        else:
            print("Последний сохранённый прогресс: нет")
        print(f"Работаю строго по диапазону: {args.start}..{args.end}")
        print(f"Файл прогресса: {progress_file}")
        print(f"Папка JSON: {cfg.json_dir}")

        for ang in range(args.start, args.end + 1):
            print(f"\n══ Анг {ang}/{args.end} ══")

            existing_json = ang_json_path(cfg.json_dir, ang)
            if existing_json.exists():
                if args.force_retranslate:
                    existing_json.unlink()
                    print(f"  ↺ Перевожу заново анг {ang}: удалён старый {existing_json.name}")
                else:
                    print(f"  ↷ JSON уже существует, пропускаю анг {ang}: {existing_json.name}")
                    continue

            source_lines = fetch_ang_source_lines(ang)
            if not source_lines:
                print(f"  ⚠ Не удалось собрать строки из KhojGurbani для анга {ang}, пропускаю")
                continue

            print(f"  → Собрано строк: {len(source_lines)}")

            chat_page = open_chat_tab(context, chat_url, cfg.page_timeout_ms)
            chat_page.bring_to_front()

            ang_data = None
            try:
                ang_data = request_structured_translation(chat_page, ang, source_lines, cfg)

                if ang_data:
                    json_path = save_ang_json(cfg.json_dir, ang_data)
                    print(f"  ✓ JSON сохранён: {json_path.name}")
                    append_ang_to_docx(output_path, ang_data)
                    save_progress(progress_file, ang)
            finally:
                if not cfg.keep_chat_tabs:
                    try:
                        chat_page.close()
                    except Exception:
                        pass

            if not ang_data:
                print(f"  ⚠ Не удалось получить валидный результат для анга {ang}, пропускаю")
                continue

            if ang < args.end:
                print(f"  ⏳ Пауза {args.delay}с...")
                time.sleep(args.delay)

        print(f"\n✓ Готово! Результат: {output_path}")
        print(f"✓ Прогресс: {progress_file}")
        print(f"✓ JSON по ангам: {cfg.json_dir}")
        if cfg.raw_log_dir:
            print(f"✓ Сырые логи: {cfg.raw_log_dir}")

        context.close()


if __name__ == "__main__":
    main()
