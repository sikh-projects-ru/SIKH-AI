#!/usr/bin/env python3
# chatgpt_khojgurbani_sahibsingh_bot.py

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import time
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from playwright.sync_api import Page, TimeoutError as PWTimeout, sync_playwright
from playwright_stealth import Stealth

_stealth = Stealth()

API_BASE = "https://apiprod.khojgurbani.org/api/v1"
HEADERS = {
    "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36",
    "Origin": "https://khojgurbani.org",
    "Referer": "https://khojgurbani.org/",
    "Accept": "application/json",
}

DEFAULT_CHAT_URL = "https://chatgpt.com/"
BOT_PROFILE = Path(__file__).parent / "bot_profile"
SHABAD_MAP_PATH = Path(__file__).parent / "shabad_map.json"

TITLE_TEXT = "Шри Гуру Грантх Сахиб — перевод на русский"
SUBTITLE_TEXT = "Источник: KhojGurbani · Prof. Sahib Singh"

COLOR_GURBANI = RGBColor(0x55, 0x00, 0x00)
COLOR_ROMAN = RGBColor(0x44, 0x44, 0x44)
COLOR_TRANSLATION = RGBColor(0x00, 0x60, 0x00)
COLOR_LABEL = RGBColor(0x99, 0x99, 0x99)
COLOR_HEADING = RGBColor(0x22, 0x22, 0x22)

ROMANIZATION_RULES = """\
Требования к полю "roman":
- нужна аккуратная панджабская romanization для русскоязычного читателя;
- сохраняй диакритику там, где это важно: ā, ī, ū, ṭ, ṭh, ḍ, ḍh, ṇ, ṛ, ñ, ś, ṃ, ṅ;
- не делай механическую школьную санскритизацию;
- roman должна читаться ближе к живому панджабскому чтению, а не как буквенная ловушка;
- не тащи автоматически конечные -u / -i, если в принятом чтении они не звучат отдельно;
- для этого проекта букву ਚ передавай как "ch", а не как "c";
- букву ਛ передавай как "chh";
- для форм вроде ਸੋਚੈ не пиши "soce"; пиши ближе к чтению: "sochey";
- для форм вроде ਗਾਵੈ не пиши "gāvai"; пиши ближе к чтению: "gāvē";
- используй практичную научную romanization с учётом чтения.

Предпочтительные примеры:
- ਸੋਚੈ → sochey
- ਵਿਚਾਰ → vichār
- ਚਾਰ → chār
- ਚੁਪ → chup
- ਸਚੁ → sach
- ਛੇ → chhē
- ਗਾਵੈ → gāvē
- ਹੁਕਮਿ → hukam
- ਸਤਿਨਾਮੁ → sat nām
- ਜਪੁ → jap
- ਗੁਰ ਪ੍ਰਸਾਦਿ → gur prasād
- ਅਜੂਨੀ → ajūnī
- ਸੈਭੰ → saibhaṅ
- ਕਰਤਾ ਪੁਰਖੁ → kartā purakh
- ਨਿਰਭਉ → nirbhau
- ਨਿਰਵੈਰੁ → nirvair
"""

RUSSIAN_GLOSSARY = """\
Предпочтительные русские формы:
- ਹਉਮੈ → Хоумэ
- ਜਪੁ → Джап
- ਜਪੁ ਜੀ → Джап Джи
- ਸਤਿਗੁਰੁ / ਸਤਿਗੁਰੂ → Сатгуру
- ਸਤਿਨਾਮੁ → Сат Нам
- ਏਕੰਕਾਰੁ / ਓਅੰਕਾਰੁ → Онкар / Эконкар по контексту, без механического финального "у"
- ਗੁਰ ਪ੍ਰਸਾਦਿ → Гур прасад
- ਅਜੂਨੀ → Аджуни
- ਸੈਭੰ → Сэйбханг
- ਕਰਤਾ ਪੁਰਖੁ → Карта Пуракх
- ਨਿਰਭਉ → Нирбхау
- ਨਿਰਵੈਰੁ → Нирвэр

Общее правило:
- roman: аккуратно и единообразно;
- русский текст: естественно, по панджабскому чтению и смыслу;
- переводить только смысл Prof. Sahib Singh из поля sahib_singh_pa;
- не добавлять свои комментарии, толкования и расширения.
"""

PROMPT_TEMPLATE = """\
Игнорируй весь предыдущий контекст этого чата.
Используй только текст текущего сообщения.

Это один анг из KhojGurbani.
Для каждой строки у тебя есть:
- verse_id
- gurmukhi
- sahib_singh_pa

Нужно вернуть материал строго по строкам этого анга.
Источник смысла для "translation_ru" — только поле "sahib_singh_pa" у соответствующей строки.
Нельзя опираться на внешний контекст, на другие переводы, на собственные догадки и на комментарии от себя.

В этом анге ровно {expected_lines} строк.
Верни ровно {expected_lines} элементов в массиве "lines" и сохрани порядок строк.

Для каждой строки верни:
1) "verse_id" — тот же integer, что во входе;
2) "roman" — аккуратная romanization строки gurmukhi;
3) "translation_ru" — только русский перевод панджабского перевода Prof. Sahib Singh из поля "sahib_singh_pa", без добавлений от себя.

Если для строки на сайте нет перевода Prof. Sahib Singh или переводная строка реально отсутствует, оставь "translation_ru" пустой строкой "".
Это допустимо. Не выдумывай перевод для таких строк.

Не добавляй:
- пояснений;
- комментариев;
- markdown;
- code fences;
- лишних полей.

{romanization_rules}

{russian_glossary}

Правила:
- сохраняй порядок строк;
- не пропускай строки;
- не меняй verse_id;
- если строка короткая, не раздувай перевод;
- не вставляй заголовки;
- ответ верни строго между BEGIN_KG_JSON и END_KG_JSON;
- ВАЖНО: никогда не используй ASCII-кавычки " " внутри значений строк — это ломает JSON; если нужно выделить слово или цитату, используй «ёлочки» или одинарные кавычки.

Формат ответа:

BEGIN_KG_JSON
{{
  "ang": {ang},
  "lines": [
    {{
      "verse_id": 123,
      "roman": "...",
      "translation_ru": "..."
    }}
  ]
}}
END_KG_JSON

Текст анга:

{content}
"""

REPAIR_PROMPT_TEMPLATE = """\
Игнорируй весь предыдущий контекст этого чата, кроме своего последнего ответа.
Используй только смысл своего последнего ответа и перепиши его в правильном формате.

Должно быть ровно {expected_lines} элементов в массиве "lines".

Нужен только валидный JSON между BEGIN_KG_JSON и END_KG_JSON.
Без markdown, без code fences, без пояснений.

Обязательные поля для каждой строки:
- verse_id
- roman
- translation_ru

Критично:
- не меняй verse_id;
- не добавляй новые поля;
- translation_ru должен быть только переводом поля sahib_singh_pa без собственных добавлений;
- если у строки нет перевода на сайте, допускается пустая строка "";
- никогда не используй ASCII-кавычки " " внутри значений строк — это ломает JSON; используй «ёлочки» или одинарные кавычки;
- roman должна учитывать проектную норму:
  - ਚ = ch
  - ਛ = chh
  - sochey
  - vichār
  - chār
  - chup
  - sach
  - chhē
  - gāvē
  - hukam
  - sat nām
  - gur prasād
  - saibhaṅ

Формат:

BEGIN_KG_JSON
{{
  "lines": [
    {{
      "verse_id": 123,
      "roman": "...",
      "translation_ru": "..."
    }}
  ]
}}
END_KG_JSON
"""


@dataclass
class SourceLine:
    index: int
    verse_id: int
    shabad_num: int
    shabad_id: int | None
    gurmukhi: str
    site_roman: str
    sahib_singh_pa: str


@dataclass
class OutputLine:
    index: int
    verse_id: int
    shabad_num: int
    shabad_id: int | None
    gurmukhi: str
    site_roman: str
    sahib_singh_pa: str
    roman: str
    translation_ru: str


@dataclass
class AngTranslation:
    ang: int
    lines: list[OutputLine]


@dataclass
class RuntimeConfig:
    page_timeout_ms: int
    input_timeout_ms: int
    response_timeout_ms: int
    new_message_timeout_ms: int
    max_retries: int
    retry_delay_s: float
    raw_log_dir: Path | None
    json_dir: Path
    keep_chat_tabs: bool
    chunk_size: int = 0  # 0 = disabled; >0 = split large angs into chunks of this size
    banidb_path: Path | None = None  # if set, filter collected verse_ids to only those in banidb


@dataclass
class CorruptLineInfo:
    ang: int
    line_index: int
    verse_id: int
    issue: str
    snippet: str


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def safe_filename_part(name: str) -> str:
    return re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", name)


def build_progress_path(output_path: Path) -> Path:
    return output_path.parent / f".progress_{safe_filename_part(output_path.name)}.txt"


def load_progress(progress_file: Path) -> int:
    if not progress_file.exists():
        return 0
    try:
        return int(progress_file.read_text(encoding="utf-8").strip() or "0")
    except ValueError:
        return 0


def save_progress(progress_file: Path, ang: int) -> None:
    progress_file.write_text(str(ang), encoding="utf-8")


def reset_progress(progress_file: Path) -> None:
    if progress_file.exists():
        progress_file.unlink()


def load_shabad_map() -> dict[int, int]:
    if not SHABAD_MAP_PATH.exists():
        return {}
    try:
        raw = json.loads(SHABAD_MAP_PATH.read_text(encoding="utf-8"))
        return {int(k): int(v) for k, v in raw.items()}
    except Exception:
        return {}


def save_shabad_map(shabad_map: dict[int, int]) -> None:
    SHABAD_MAP_PATH.write_text(
        json.dumps({str(k): v for k, v in sorted(shabad_map.items())}, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


def build_shabad_map_from_json(json_dir: Path) -> dict[int, int]:
    shabad_map: dict[int, int] = {}
    for p in sorted(json_dir.glob("ang_*.json")):
        stem = p.stem[4:]
        if not stem.isdigit():
            continue
        ang = int(stem)
        data = load_ang_json(json_dir, ang)
        if data and data.lines:
            shabad_map[ang] = min(line.shabad_num for line in data.lines)
    return shabad_map


def estimate_start_probe(ang: int, shabad_map: dict[int, int]) -> int:
    if ang in shabad_map:
        return shabad_map[ang]

    sorted_entries = sorted(shabad_map.items())
    lower = [(a, s) for a, s in sorted_entries if a < ang]
    if not lower:
        return 1

    ref_ang, ref_shabad = lower[-1]
    first_ang, first_shabad = sorted_entries[0]

    # Overall rate from first to last known entry — much more stable than local rate
    if ref_ang != first_ang:
        rate = (ref_shabad - first_shabad) / (ref_ang - first_ang)
    else:
        rate = ref_shabad / max(1, ref_ang)

    estimate = int(ref_shabad + (ang - ref_ang) * rate)
    # Start 100 below estimate: estimate tends to overshoot slightly
    return max(1, estimate - 100)


def save_raw_text(raw_log_dir: Path | None, filename: str, text: str) -> None:
    if raw_log_dir is None:
        return
    raw_log_dir.mkdir(parents=True, exist_ok=True)
    (raw_log_dir / filename).write_text(text, encoding="utf-8")


def banidb_verse_ids_for_ang(banidb_path: Path, ang: int) -> set[int]:
    """Return the set of verse_ids that banidb assigns to this ang."""
    import sqlite3
    try:
        conn = sqlite3.connect(str(banidb_path))
        cur = conn.cursor()
        cur.execute("SELECT verse_id FROM verses WHERE ang = ?", (ang,))
        result = {row[0] for row in cur.fetchall()}
        conn.close()
        return result
    except Exception:
        return set()


def api_get(path: str) -> dict[str, Any]:
    req = urllib.request.Request(f"{API_BASE}{path}", headers=HEADERS)
    with urllib.request.urlopen(req, timeout=20) as response:
        return json.loads(response.read())


def fetch_shabad(ang: int, shabad_num: int, *, quiet: bool = False) -> dict[str, Any] | None:
    try:
        data = api_get(f"/shabad/{ang}/{shabad_num}")
        if data.get("status") == "success":
            payload = data.get("data")
            if isinstance(payload, dict):
                return payload
    except Exception as exc:
        if not quiet:
            print(f"  [!] fetch_shabad({ang}/{shabad_num}): {exc}")
    return None


def _normalize_int_list(values: Any) -> list[int]:
    result: list[int] = []
    if not isinstance(values, list):
        return result

    for item in values:
        try:
            result.append(int(item))
        except Exception:
            continue

    return sorted(set(result))


def discover_shabad_numbers_for_ang(
    ang: int,
    max_probe: int = 1500,
    start_probe: int = 1,
    initial_miss_limit: int = 10,
) -> list[int]:
    found_by_scan: list[int] = []
    misses_after_first = 0
    initial_misses = 0

    for probe in range(start_probe, start_probe + max_probe):
        shabad = fetch_shabad(ang, probe, quiet=True)
        if shabad:
            pages = _normalize_int_list(shabad.get("pages"))
            if pages:
                return pages

            found_by_scan.append(probe)
            misses_after_first = 0
            initial_misses = 0
            continue

        if found_by_scan:
            misses_after_first += 1
            if misses_after_first >= 2:
                break
        else:
            initial_misses += 1
            if initial_misses >= initial_miss_limit:
                break

    return found_by_scan


def fetch_ang_source_lines(
    ang: int,
    start_probe: int = 1,
    initial_miss_limit: int = 10,
    valid_verse_ids: set[int] | None = None,
) -> list[SourceLine]:
    if start_probe > 1:
        print(f"  → Зондирование шабадов с №{start_probe}")

    shabad_numbers = discover_shabad_numbers_for_ang(ang, start_probe=start_probe, initial_miss_limit=initial_miss_limit)
    if not shabad_numbers:
        print(f"  ✗ Не удалось определить шабады для анга {ang}")
        return []

    lines: list[SourceLine] = []
    seen_verse_ids: set[int] = set()

    for shabad_num in shabad_numbers:
        shabad = fetch_shabad(ang, shabad_num)
        if not shabad:
            continue

        shabad_id = shabad.get("id")
        scriptures = shabad.get("scriptures", [])
        if not isinstance(scriptures, list):
            continue

        for verse in scriptures:
            try:
                verse_id = int(verse["id"])
            except Exception:
                continue

            if verse_id in seen_verse_ids:
                continue

            if valid_verse_ids is not None and verse_id not in valid_verse_ids:
                continue

            gurmukhi = normalize_text(str(verse.get("Scripture", "")))
            site_roman = normalize_text(str(verse.get("ScriptureRoman", "")))
            translations = verse.get("translation", {}) or {}
            sahib_singh_pa = normalize_text(str(translations.get("SahibSinghPunjabi", "")))

            if not gurmukhi or not sahib_singh_pa:
                continue

            seen_verse_ids.add(verse_id)
            lines.append(
                SourceLine(
                    index=len(lines) + 1,
                    verse_id=verse_id,
                    shabad_num=shabad_num,
                    shabad_id=int(shabad_id) if isinstance(shabad_id, int) else None,
                    gurmukhi=gurmukhi,
                    site_roman=site_roman,
                    sahib_singh_pa=sahib_singh_pa,
                )
            )

    return lines


def build_prompt_input(lines: list[SourceLine]) -> str:
    parts: list[str] = []

    for line in lines:
        parts.append(
            "\n".join(
                [
                    "[LINE]",
                    f"index: {line.index}",
                    f"verse_id: {line.verse_id}",
                    f"gurmukhi: {line.gurmukhi}",
                    f"sahib_singh_pa: {line.sahib_singh_pa}",
                    "[/LINE]",
                ]
            )
        )

    return "\n\n".join(parts).strip()


def build_prompt(ang: int, lines: list[SourceLine]) -> str:
    return PROMPT_TEMPLATE.format(
        ang=ang,
        expected_lines=len(lines),
        romanization_rules=ROMANIZATION_RULES,
        russian_glossary=RUSSIAN_GLOSSARY,
        content=build_prompt_input(lines),
    )


def build_repair_prompt(expected_lines: int) -> str:
    return REPAIR_PROMPT_TEMPLATE.format(expected_lines=expected_lines)


def add_run(
    para,
    text: str,
    *,
    color=None,
    bold: bool = False,
    italic: bool = False,
    size_pt: int | None = None,
):
    run = para.add_run(text)
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def ensure_output_doc_exists(docx_path: Path) -> None:
    if docx_path.exists():
        return

    doc = Document()
    doc.core_properties.author = "KhojGurbani Sahib Singh Bot"

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title = doc.add_heading(TITLE_TEXT, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(SUBTITLE_TEXT)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = COLOR_LABEL

    doc.add_paragraph()
    doc.save(str(docx_path))


def reset_output_doc(docx_path: Path) -> None:
    if docx_path.exists():
        docx_path.unlink()
    ensure_output_doc_exists(docx_path)


def add_ang_heading(doc: Document, ang: int) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run(f"АНГ {ang}")
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_HEADING


def add_centered_line_block(doc: Document, line: OutputLine) -> None:
    p_g = doc.add_paragraph()
    p_g.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_g.paragraph_format.space_before = Pt(8)
    p_g.paragraph_format.space_after = Pt(2)
    add_run(p_g, line.gurmukhi, color=COLOR_GURBANI, size_pt=13, bold=True)

    p_r = doc.add_paragraph()
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r.paragraph_format.space_after = Pt(3)
    add_run(p_r, line.roman, color=COLOR_ROMAN, size_pt=10, italic=True)

    if normalize_text(line.translation_ru):
        p_t = doc.add_paragraph()
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_t.paragraph_format.space_after = Pt(10)
        add_run(p_t, line.translation_ru, color=COLOR_TRANSLATION, size_pt=11)

    doc.add_paragraph()


def append_ang_to_docx(docx_path: Path, ang_data: AngTranslation) -> None:
    ensure_output_doc_exists(docx_path)
    doc = Document(str(docx_path))

    add_ang_heading(doc, ang_data.ang)

    for line in ang_data.lines:
        add_centered_line_block(doc, line)

    doc.save(str(docx_path))
    print(f"  ✓ DOCX дописан: {docx_path.name}")


def ang_to_dict(ang_data: AngTranslation) -> dict[str, Any]:
    return {
        "ang": ang_data.ang,
        "line_count": len(ang_data.lines),
        "lines": [
            {
                "index": line.index,
                "verse_id": line.verse_id,
                "shabad_num": line.shabad_num,
                "shabad_id": line.shabad_id,
                "gurmukhi": line.gurmukhi,
                "site_roman": line.site_roman,
                "sahib_singh_pa": line.sahib_singh_pa,
                "roman": line.roman,
                "translation_ru": line.translation_ru,
            }
            for line in ang_data.lines
        ],
    }


def output_line_from_dict(item: dict[str, Any]) -> OutputLine:
    shabad_id_raw = item.get("shabad_id")
    shabad_id = int(shabad_id_raw) if isinstance(shabad_id_raw, int) else None

    return OutputLine(
        index=int(item.get("index", 0)),
        verse_id=int(item.get("verse_id", 0)),
        shabad_num=int(item.get("shabad_num", 0)),
        shabad_id=shabad_id,
        gurmukhi=normalize_text(str(item.get("gurmukhi", ""))),
        site_roman=normalize_text(str(item.get("site_roman", ""))),
        sahib_singh_pa=normalize_text(str(item.get("sahib_singh_pa", ""))),
        roman=normalize_text(str(item.get("roman", ""))),
        translation_ru=normalize_text(str(item.get("translation_ru", ""))),
    )


def ang_translation_from_dict(data: dict[str, Any]) -> AngTranslation:
    raw_lines = data.get("lines", [])
    lines: list[OutputLine] = []
    if isinstance(raw_lines, list):
        for item in raw_lines:
            if not isinstance(item, dict):
                continue
            lines.append(output_line_from_dict(item))

    return AngTranslation(ang=int(data.get("ang", 0)), lines=lines)


def ang_json_path(json_dir: Path, ang: int) -> Path:
    return json_dir / f"ang_{ang:04d}.json"


def save_ang_json(json_dir: Path, ang_data: AngTranslation) -> Path:
    json_dir.mkdir(parents=True, exist_ok=True)
    json_path = ang_json_path(json_dir, ang_data.ang)
    json_path.write_text(
        json.dumps(ang_to_dict(ang_data), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return json_path


def load_ang_json(json_dir: Path, ang: int) -> AngTranslation | None:
    json_path = ang_json_path(json_dir, ang)
    if not json_path.exists():
        return None

    try:
        data = json.loads(json_path.read_text(encoding="utf-8"))
    except Exception:
        return None

    return ang_translation_from_dict(data)


def reset_json_range(json_dir: Path, start: int, end: int) -> int:
    removed = 0
    for ang in range(start, end + 1):
        json_path = ang_json_path(json_dir, ang)
        if json_path.exists():
            json_path.unlink()
            removed += 1
    return removed


def rebuild_docx_from_json(output_path: Path, json_dir: Path, start: int, end: int) -> int:
    reset_output_doc(output_path)
    rebuilt = 0

    for ang in range(start, end + 1):
        ang_data = load_ang_json(json_dir, ang)
        if not ang_data or not ang_data.lines:
            print(f"  ↷ Нет JSON для анга {ang}, пропускаю")
            continue

        append_ang_to_docx(output_path, ang_data)
        rebuilt += 1

    return rebuilt


def _has_foreign_script(text: str) -> bool:
    """Возвращает True, если в тексте есть деванагари или гурмукхи."""
    return bool(_RE_DEVANAGARI.search(text) or _RE_GURMUKHI_OUT.search(text))


# ---------------------------------------------------------------------------
# Таблица: деванагари / гурмукхи → латиница (для авто-фикса поля roman)
# Покрывает типичные вставки ChatGPT (согласные, гласные знаки, анусвара).
# ---------------------------------------------------------------------------
_SCRIPT_TO_LATIN: dict[str, str] = {
    # ── Деванагари: согласные ──────────────────────────────────────────────
    "क": "k",  "ख": "kh", "ग": "g",  "घ": "gh", "ङ": "ṅ",
    "च": "ch", "छ": "chh","ज": "j",  "झ": "jh", "ञ": "ñ",
    "ट": "ṭ",  "ठ": "ṭh", "ड": "ḍ",  "ढ": "ḍh", "ण": "ṇ",
    "त": "t",  "थ": "th", "द": "d",  "ध": "dh", "न": "n",
    "प": "p",  "फ": "ph", "ब": "b",  "भ": "bh", "म": "m",
    "य": "y",  "र": "r",  "ल": "l",  "व": "v",
    "श": "sh", "ष": "ṣh", "स": "s",  "ह": "h",
    "ळ": "ḷ",  "क्ष": "kṣh", "ज्ञ": "jñ",
    # ── Деванагари: гласные (самостоятельные) ──────────────────────────────
    "अ": "a",  "आ": "ā",  "इ": "i",  "ई": "ī",
    "उ": "u",  "ऊ": "ū",  "ए": "e",  "ऐ": "ai",
    "ओ": "o",  "औ": "au", "ऋ": "ri",
    # ── Деванагари: знаки гласных (матры) ─────────────────────────────────
    "\u093E": "ā",  # ा
    "\u093F": "i",  # ि
    "\u0940": "ī",  # ी
    "\u0941": "u",  # ु
    "\u0942": "ū",  # ू
    "\u0947": "e",  # े
    "\u0948": "ai", # ै
    "\u094B": "o",  # ो
    "\u094C": "au", # ौ
    "\u0943": "ri", # ृ
    # ── Деванагари: прочие знаки ───────────────────────────────────────────
    "\u0902": "ṃ",  # ं  анусвара
    "\u0903": "ḥ",  # ः  висарга
    "\u094D": "",   # ्  халант (вирама — подавляет гласную)
    "\u0901": "ṃ",  # ँ  чандрабинду
    # ── Гурмукхи: согласные ───────────────────────────────────────────────
    "ਕ": "k",  "ਖ": "kh", "ਗ": "g",  "ਘ": "gh", "ਙ": "ṅ",
    "ਚ": "ch", "ਛ": "chh","ਜ": "j",  "ਝ": "jh", "ਞ": "ñ",
    "ਟ": "ṭ",  "ਠ": "ṭh", "ਡ": "ḍ",  "ਢ": "ḍh", "ਣ": "ṇ",
    "ਤ": "t",  "ਥ": "th", "ਦ": "d",  "ਧ": "dh", "ਨ": "n",
    "ਪ": "p",  "ਫ": "ph", "ਬ": "b",  "ਭ": "bh", "ਮ": "m",
    "ਯ": "y",  "ਰ": "r",  "ਲ": "l",  "ਵ": "v",
    "ਸ": "s",  "ਹ": "h",  "ਲ਼": "ḷ",
    "ੜ": "ṛ",  "ਸ਼": "sh", "ਖ਼": "kh", "ਗ਼": "g",
    "ਜ਼": "z",  "ਫ਼": "f",
    # ── Гурмукхи: гласные (самостоятельные) ──────────────────────────────
    "ਅ": "a",  "ਆ": "ā",  "ਇ": "i",  "ਈ": "ī",
    "ਉ": "u",  "ਊ": "ū",  "ਏ": "e",  "ਐ": "ai",
    "ਓ": "o",  "ਔ": "au",
    # ── Гурмукхи: знаки гласных (лагама) ─────────────────────────────────
    "\u0A3E": "ā",  # ਾ
    "\u0A3F": "i",  # ਿ
    "\u0A40": "ī",  # ੀ
    "\u0A41": "u",  # ੁ
    "\u0A42": "ū",  # ੂ
    "\u0A47": "e",  # ੇ
    "\u0A48": "ai", # ੈ
    "\u0A4B": "o",  # ੋ
    "\u0A4C": "au", # ੌ
    # ── Гурмукхи: прочие знаки ────────────────────────────────────────────
    "\u0A02": "ṃ",  # ਂ  биндӣ
    "\u0A70": "ṃ",  # ੰ  типпа
    "\u0A71": "",   # ੱ  аддак (удвоение согласной — игнорируем)
    "\u0A3C": "",   # ਼  нуктā
    "\u0A4D": "",   # ੍  вирама
}


def _fix_foreign_chars_in_roman(text: str) -> str:
    """Заменяет деванагари/гурмукхи символы в поле roman их латинскими эквивалентами.

    Сначала пробует найти многосимвольные последовательности (например ਸ਼ → sh),
    затем одиночные символы из таблицы. Неизвестные символы оставляет как есть.
    """
    # Сортируем ключи по убыванию длины: многосимвольные замены в первую очередь
    sorted_keys = sorted(_SCRIPT_TO_LATIN.keys(), key=len, reverse=True)
    result = text
    for key in sorted_keys:
        result = result.replace(key, _SCRIPT_TO_LATIN[key])
    return result


def fix_corrupt_roman_in_json(json_dir: Path, start: int, end: int) -> dict[int, int]:
    """Авто-фиксирует поле roman в JSON-файлах: заменяет деванагари/гурмукхи → латиница.

    Возвращает словарь {ang: кол-во исправленных строк}.
    Файлы с изменениями перезаписываются.
    """
    fixed: dict[int, int] = {}
    for ang in range(start, end + 1):
        ang_data = load_ang_json(json_dir, ang)
        if not ang_data:
            continue
        changed = 0
        for line in ang_data.lines:
            if _has_foreign_script(line.roman):
                new_roman = _fix_foreign_chars_in_roman(line.roman)
                if new_roman != line.roman:
                    line.roman = new_roman
                    changed += 1
        if changed:
            save_ang_json(json_dir, ang_data)
            fixed[ang] = changed
    return fixed


def scan_corrupt_angs(json_dir: Path, start: int, end: int) -> dict[int, list[CorruptLineInfo]]:
    """Сканирует JSON-файлы в диапазоне и возвращает анги с битыми строками.

    Битая строка — та, в которой translation_ru или roman содержит деванагари
    или гурмукхи (ChatGPT иногда вставляет хинди или панджаби вместо русского).
    """
    result: dict[int, list[CorruptLineInfo]] = {}
    for ang in range(start, end + 1):
        ang_data = load_ang_json(json_dir, ang)
        if not ang_data:
            continue
        issues: list[CorruptLineInfo] = []
        for line in ang_data.lines:
            if _has_foreign_script(line.translation_ru):
                snippet = line.translation_ru[:120].replace("\n", " ")
                issues.append(CorruptLineInfo(
                    ang=ang,
                    line_index=line.index,
                    verse_id=line.verse_id,
                    issue="деванагари/гурмукхи в translation_ru",
                    snippet=snippet,
                ))
            elif _has_foreign_script(line.roman):
                snippet = f"roman: {line.roman[:80].replace(chr(10), ' ')}"
                issues.append(CorruptLineInfo(
                    ang=ang,
                    line_index=line.index,
                    verse_id=line.verse_id,
                    issue="деванагари/гурмукхи в roman",
                    snippet=snippet,
                ))
        if issues:
            result[ang] = issues
    return result


_META_GUESS_PATTERNS = [
    re.compile(r"\bя\s+думаю\b", re.IGNORECASE),
    re.compile(r"\bмне\s+кажется\b", re.IGNORECASE),
    re.compile(r"\bвероятно\b", re.IGNORECASE),
    re.compile(r"\bпохоже\b", re.IGNORECASE),
    re.compile(r"\bпо-видимому\b", re.IGNORECASE),
]

# Скрипты, которые не должны появляться в translation_ru или roman.
# Исключения (встречаются в нормальных строках):
#   U+0964 (।), U+0965 (॥) — деванагари-данды, маркеры стихов
#   U+0A66–U+0A6F (੦-੯)   — гурмукхи-цифры в нумерации стихов
_RE_DEVANAGARI = re.compile(r"[\u0900-\u0963\u0966-\u097F]")
_RE_GURMUKHI_OUT = re.compile(r"[\u0A00-\u0A65\u0A70-\u0A7F]")
# TODO: extend _has_foreign_script to also catch Telugu (U+0C00-U+0C7F) and
# Kannada (U+0C80-U+0CFF) \u2014 ChatGPT occasionally outputs "\u0C26\u0C30\u0C4D\u0C36" (Telugu darsha)
# or "\u0C82\u0C95" (Kannada) instead of Cyrillic when translating darshan/Nirankar.
# Known fix: replace "\u0C26\u0C30\u0C4D\u0C36*" \u2192 "\u0432\u0438\u0434\u0435\u043D\u0438*" and "\u0C82\u0C95" \u2192 "\u0430\u043D\u043A".


def repair_json_quotes(text: str) -> str:
    """Fix unescaped double-quotes inside JSON string values.

    ChatGPT sometimes writes "я" or "слово" inside a translation string,
    breaking the JSON. This scanner detects those and escapes them.
    """
    result: list[str] = []
    i = 0
    in_string = False
    escape_next = False

    while i < len(text):
        c = text[i]

        if escape_next:
            result.append(c)
            escape_next = False
            i += 1
            continue

        if c == "\\":
            result.append(c)
            escape_next = True
            i += 1
            continue

        if c == '"':
            if not in_string:
                in_string = True
                result.append(c)
            else:
                # Determine if this quote closes the current string or is an inner quote.
                # Look ahead past whitespace to see what follows.
                j = i + 1
                while j < len(text) and text[j] in " \t\n\r":
                    j += 1
                next_ch = text[j] if j < len(text) else ""
                if next_ch in (",", "}", "]", ":"):
                    in_string = False
                    result.append(c)
                else:
                    # Unescaped inner quote — escape it
                    result.append("\\")
                    result.append(c)
            i += 1
            continue

        result.append(c)
        i += 1

    return "".join(result)


def looks_like_model_guess(text: str) -> bool:
    sample = normalize_text(text)
    if not sample:
        return False
    return any(pattern.search(sample) for pattern in _META_GUESS_PATTERNS)


def open_chat_tab(context, chat_url: str, page_timeout_ms: int) -> Page:
    page = context.new_page()
    _stealth.apply_stealth_sync(page)
    page.goto(chat_url, wait_until="domcontentloaded", timeout=page_timeout_ms)
    page.wait_for_timeout(2000)
    return page


def assistant_message_count(page: Page) -> int:
    return page.locator('[data-message-author-role="assistant"]').count()


def focus_textarea(page: Page, input_timeout_ms: int):
    textarea = page.locator("#prompt-textarea")
    textarea.wait_for(state="visible", timeout=input_timeout_ms)
    textarea.click()
    return textarea


def clear_textarea(page: Page, textarea) -> None:
    modifier = "Meta" if os.name == "darwin" else "Control"
    textarea.click()
    page.keyboard.press(f"{modifier}+a")
    page.keyboard.press("Delete")
    page.wait_for_timeout(250)


def insert_prompt(page: Page, prompt: str) -> None:
    page.evaluate(
        """(text) => {
            const el = document.getElementById("prompt-textarea");
            if (!el) return;
            el.focus();
            document.execCommand("selectAll", false, null);
            document.execCommand("insertText", false, text);
        }""",
        prompt,
    )
    page.wait_for_timeout(500)


def click_send(page: Page, textarea) -> None:
    send_btn = page.locator('button[data-testid="send-button"]')
    try:
        send_btn.wait_for(state="visible", timeout=5_000)
        send_btn.click()
    except PWTimeout:
        textarea.press("Enter")


def wait_generation_finished(page: Page, response_timeout_ms: int) -> None:
    stop_btn = page.locator('button[data-testid="stop-button"]')

    try:
        stop_btn.wait_for(state="visible", timeout=20_000)
    except PWTimeout:
        pass

    try:
        stop_btn.wait_for(state="hidden", timeout=response_timeout_ms)
    except PWTimeout:
        print(f"  ⚠ Таймаут ожидания ответа ({response_timeout_ms / 1000:.0f}с) — беру то, что есть")

    page.wait_for_timeout(1200)


def click_continue_if_present(page: Page) -> bool:
    candidates = [
        page.get_by_role("button", name=re.compile(r"Continue generating", re.I)),
        page.get_by_role("button", name=re.compile(r"Продолжить", re.I)),
        page.locator("button").filter(has_text=re.compile(r"Continue generating|Продолжить", re.I)),
    ]

    for locator in candidates:
        try:
            if locator.count() > 0 and locator.first.is_visible():
                locator.first.click()
                page.wait_for_timeout(600)
                print("  → Нажимаю Continue generating...")
                return True
        except Exception:
            continue

    return False


def drain_generation(page: Page, response_timeout_ms: int, max_continues: int = 5) -> None:
    for _ in range(max_continues + 1):
        wait_generation_finished(page, response_timeout_ms)
        if not click_continue_if_present(page):
            return


def get_new_last_assistant_message(page: Page, before_count: int, new_message_timeout_ms: int) -> str | None:
    try:
        page.wait_for_function(
            """(before) => {
                return document.querySelectorAll('[data-message-author-role="assistant"]').length > before;
            }""",
            arg=before_count,
            timeout=new_message_timeout_ms,
        )
    except PWTimeout:
        pass

    messages = page.locator('[data-message-author-role="assistant"]')
    after_count = messages.count()

    if after_count == 0 or after_count <= before_count:
        return None

    last = messages.nth(after_count - 1)
    try:
        last.wait_for(state="visible", timeout=5_000)
    except PWTimeout:
        return None

    text = normalize_text(last.inner_text())
    return text or None


def send_prompt_and_get_answer(page: Page, prompt: str, cfg: RuntimeConfig) -> str | None:
    try:
        textarea = focus_textarea(page, cfg.input_timeout_ms)
    except PWTimeout:
        print("  ✗ Поле ввода ChatGPT не найдено")
        return None

    before_count = assistant_message_count(page)
    clear_textarea(page, textarea)
    insert_prompt(page, prompt)
    click_send(page, textarea)

    print("  → Жду ответа ChatGPT...")
    drain_generation(page, cfg.response_timeout_ms)

    return get_new_last_assistant_message(page, before_count, cfg.new_message_timeout_ms)


def extract_json_candidate(text: str) -> str | None:
    text = text.strip()

    match = re.search(r"BEGIN_KG_JSON\s*(.*?)\s*END_KG_JSON", text, flags=re.S)
    if not match:
        return None

    candidate = match.group(1).strip()
    candidate = re.sub(r"^```json\s*", "", candidate, flags=re.I)
    candidate = re.sub(r"^```\s*", "", candidate)
    candidate = re.sub(r"\s*```$", "", candidate)

    if candidate.startswith("{") and candidate.endswith("}"):
        return candidate

    return None


def merge_model_lines(
    ang: int,
    source_lines: list[SourceLine],
    model_lines: list[dict[str, Any]],
) -> tuple[AngTranslation | None, str]:
    if len(model_lines) != len(source_lines):
        return None, f"line count mismatch: expected {len(source_lines)}, got {len(model_lines)}"

    expected_ids = [line.verse_id for line in source_lines]
    received_ids: list[int] = []
    merged_lines: list[OutputLine] = []

    for source, item in zip(source_lines, model_lines):
        if not isinstance(item, dict):
            return None, f"line #{source.index} is not a dict"

        try:
            verse_id = int(item.get("verse_id"))
        except Exception:
            return None, f"line #{source.index}: invalid verse_id"

        received_ids.append(verse_id)
        if verse_id != source.verse_id:
            return None, f"line #{source.index}: verse_id mismatch, expected {source.verse_id}, got {verse_id}"

        roman = normalize_text(str(item.get("roman", "")))
        translation_ru = normalize_text(str(item.get("translation_ru", "")))

        if not roman:
            return None, f"line #{source.index}: missing roman"

        merged_lines.append(
            OutputLine(
                index=source.index,
                verse_id=source.verse_id,
                shabad_num=source.shabad_num,
                shabad_id=source.shabad_id,
                gurmukhi=source.gurmukhi,
                site_roman=source.site_roman,
                sahib_singh_pa=source.sahib_singh_pa,
                roman=roman,
                translation_ru=translation_ru,
            )
        )

    if received_ids != expected_ids:
        return None, "line order mismatch"

    return AngTranslation(ang=ang, lines=merged_lines), ""


def parse_structured_answer(answer: str, ang: int, source_lines: list[SourceLine]) -> tuple[AngTranslation | None, str]:
    candidate = extract_json_candidate(answer)
    if not candidate:
        return None, "BEGIN_KG_JSON / END_KG_JSON not found"

    try:
        data = json.loads(candidate)
    except json.JSONDecodeError:
        repaired = repair_json_quotes(candidate)
        try:
            data = json.loads(repaired)
        except json.JSONDecodeError as exc:
            return None, f"json decode error: {exc}"

    if not isinstance(data, dict):
        return None, "payload is not a dict"

    lines = data.get("lines")
    if not isinstance(lines, list) or not lines:
        return None, "lines is empty or not a list"

    payload_ang = data.get("ang")
    if payload_ang not in (None, ang):
        try:
            if int(payload_ang) != ang:
                return None, f"ang mismatch: expected {ang}, got {payload_ang}"
        except Exception:
            return None, f"invalid ang value: {payload_ang}"

    return merge_model_lines(ang, source_lines, lines)


def request_structured_translation(
    page: Page,
    ang: int,
    source_lines: list[SourceLine],
    cfg: RuntimeConfig,
) -> AngTranslation | None:
    prompt = build_prompt(ang, source_lines)
    repair_prompt = build_repair_prompt(len(source_lines))

    save_raw_text(cfg.raw_log_dir, f"ang_{ang:04d}_source_lines.json", json.dumps(
        [
            {
                "index": line.index,
                "verse_id": line.verse_id,
                "shabad_num": line.shabad_num,
                "shabad_id": line.shabad_id,
                "gurmukhi": line.gurmukhi,
                "site_roman": line.site_roman,
                "sahib_singh_pa": line.sahib_singh_pa,
            }
            for line in source_lines
        ],
        ensure_ascii=False,
        indent=2,
    ))
    save_raw_text(cfg.raw_log_dir, f"ang_{ang:04d}_prompt.txt", prompt)

    for attempt in range(1, cfg.max_retries + 1):
        print(f"  → Попытка {attempt}/{cfg.max_retries}")
        current_prompt = prompt if attempt == 1 else repair_prompt
        answer = send_prompt_and_get_answer(page, current_prompt, cfg)

        if answer:
            save_raw_text(cfg.raw_log_dir, f"ang_{ang:04d}_attempt_{attempt}_answer.txt", answer)
            parsed, reason = parse_structured_answer(answer, ang, source_lines)
            if parsed:
                print(f"  ✓ Ответ распознан: строк {len(parsed.lines)}")
                return parsed
            print(f"  ⚠ Ответ получен, но не прошёл проверку: {reason}")
        else:
            save_raw_text(cfg.raw_log_dir, f"ang_{ang:04d}_attempt_{attempt}_answer.txt", "[NO_ANSWER]")
            print("  ⚠ Пустой ответ или не удалось получить новый ответ ассистента")

        if attempt < cfg.max_retries:
            print(f"  ⏳ Повтор через {cfg.retry_delay_s}с...")
            time.sleep(cfg.retry_delay_s)

    return None


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="KhojGurbani → ChatGPT → JSON + DOCX (по ангам, только Sahib Singh)")

    parser.add_argument("--start", type=int, default=1, help="Начальный анг")
    parser.add_argument("--end", type=int, default=1430, help="Конечный анг")
    parser.add_argument("--output", type=str, default="khojgurbani_sahibsingh_chatgpt.docx", help="Имя выходного DOCX")
    parser.add_argument("--delay", type=float, default=3.0, help="Пауза между ангами в секундах")

    parser.add_argument("--max-retries", type=int, default=3, help="Максимум попыток на один анг")
    parser.add_argument("--retry-delay", type=float, default=10.0, help="Пауза между повторами в секундах")

    parser.add_argument("--page-timeout", type=int, default=30, help="Таймаут загрузки страницы в секундах")
    parser.add_argument("--input-timeout", type=int, default=15, help="Таймаут ожидания поля ввода в секундах")
    parser.add_argument("--response-timeout", type=int, default=900, help="Таймаут ожидания ответа модели в секундах")
    parser.add_argument("--new-message-timeout", type=int, default=30, help="Таймаут ожидания нового assistant-сообщения в секундах")

    parser.add_argument(
        "--chat-url",
        type=str,
        default=None,
        help="URL ChatGPT. Обычно достаточно https://chatgpt.com/ — для каждого анга откроется новая вкладка.",
    )
    parser.add_argument(
        "--raw-log-dir",
        type=str,
        default="raw_logs",
        help='Папка для сырых логов prompt/answer. Пустая строка "" отключает логи.',
    )
    parser.add_argument(
        "--json-dir",
        type=str,
        default="ang_json",
        help="Папка для JSON-дампа по ангам",
    )
    parser.add_argument(
        "--rebuild-docx-from-json",
        action="store_true",
        help="Пересобрать DOCX из уже сохранённых JSON в указанном диапазоне, без GPT.",
    )
    parser.add_argument(
        "--reset-progress",
        action="store_true",
        help="Сбросить .progress_<output>.txt перед запуском.",
    )
    parser.add_argument(
        "--force-retranslate",
        action="store_true",
        help="Не пропускать анги из-за уже существующих JSON; перевести диапазон заново и перезаписать JSON.",
    )
    parser.add_argument(
        "--reset-json-range",
        action="store_true",
        help="Перед запуском удалить JSON текущего диапазона --start..--end.",
    )
    parser.add_argument(
        "--keep-chat-tabs",
        action="store_true",
        help="Не закрывать вкладки ChatGPT после обработки анга.",
    )
    parser.add_argument(
        "--scan-missing",
        action="store_true",
        help=(
            "Сканировать диапазон --start..--end, показать отсутствующие JSON, "
            "спросить подтверждение и дополнить их через ChatGPT."
        ),
    )
    parser.add_argument(
        "--scan-corrupt",
        action="store_true",
        help=(
            "Сканировать диапазон --start..--end, найти строки с деванагари/гурмукхи "
            "в translation_ru или roman, спросить подтверждение и переперевести их через ChatGPT."
        ),
    )
    parser.add_argument(
        "--fix-corrupt-roman",
        action="store_true",
        help=(
            "Авто-исправить поле roman в диапазоне --start..--end: "
            "заменить деванагари/гурмукхи символы их латинскими эквивалентами прямо в JSON, "
            "без ChatGPT. Работает только для поля roman — translation_ru не трогает."
        ),
    )
    parser.add_argument(
        "--menu",
        action="store_true",
        help="Показать интерактивное меню для выбора режима работы.",
    )
    parser.add_argument(
        "--update-shabad-map",
        action="store_true",
        help="Пересобрать shabad_map.json из ang_json/ и сохранить.",
    )
    parser.add_argument(
        "--banidb",
        type=str,
        default="",
        help=(
            "Путь к sggs.db (banidb). Если указан, при сборе строк из KhojGurbani "
            "оставляет только те verse_id, которые banidb относит к данному ангу. "
            "Устраняет проблему сбора 400+ строк и дублей."
        ),
    )
    parser.add_argument(
        "--chunk-size",
        type=int,
        default=0,
        help=(
            "Дробить анг на чанки по N строк если строк больше N (0 = отключено). "
            "Рекомендуется 50 для больших ангов: --chunk-size 50"
        ),
    )

    return parser.parse_args()


def _translate_source_lines(
    context,
    chat_url: str,
    ang: int,
    source_lines: list[SourceLine],
    cfg: RuntimeConfig,
) -> AngTranslation | None:
    """Translate source_lines for an ang, splitting into chunks if cfg.chunk_size is set."""
    chunk_size = cfg.chunk_size
    if chunk_size <= 0 or len(source_lines) <= chunk_size:
        page = open_chat_tab(context, chat_url, cfg.page_timeout_ms)
        page.bring_to_front()
        try:
            return request_structured_translation(page, ang, source_lines, cfg)
        finally:
            if not cfg.keep_chat_tabs:
                try:
                    page.close()
                except Exception:
                    pass

    chunks = [source_lines[i:i + chunk_size] for i in range(0, len(source_lines), chunk_size)]
    total_chunks = len(chunks)
    print(f"  → Дробим на {total_chunks} чанков по ≤{chunk_size} строк")

    all_lines: list[OutputLine] = []
    for chunk_idx, chunk in enumerate(chunks, 1):
        print(f"  → Чанк {chunk_idx}/{total_chunks} (строки {chunk[0].index}–{chunk[-1].index}, {len(chunk)} строк)")
        page = open_chat_tab(context, chat_url, cfg.page_timeout_ms)
        page.bring_to_front()
        try:
            result = request_structured_translation(page, ang, chunk, cfg)
        finally:
            if not cfg.keep_chat_tabs:
                try:
                    page.close()
                except Exception:
                    pass
        if result is None:
            print(f"  ⚠ Чанк {chunk_idx}/{total_chunks} не удался — отменяю анг {ang}")
            return None
        all_lines.extend(result.lines)

    return AngTranslation(ang=ang, lines=all_lines)


def run_browser_session(
    args: argparse.Namespace,
    cfg: RuntimeConfig,
    output_path: Path,
    progress_file: Path,
) -> None:
    """Открывает браузер и переводит анги в диапазоне args.start..args.end."""
    ensure_output_doc_exists(output_path)

    chat_url = (args.chat_url or "").strip()
    if not chat_url:
        print("URL чата не задан — открою ChatGPT по умолчанию.")
        chat_url = DEFAULT_CHAT_URL

    if not chat_url.startswith("http"):
        print("  ✗ Некорректный URL, выхожу.")
        return

    BOT_PROFILE.mkdir(exist_ok=True)
    first_run = not (BOT_PROFILE / "Default").exists()

    with sync_playwright() as pw:
        print("Запускаю браузер...")
        context = pw.chromium.launch_persistent_context(
            user_data_dir=str(BOT_PROFILE),
            headless=False,
            args=["--start-maximized"],
            no_viewport=True,
        )

        if first_run:
            login_page = open_chat_tab(context, chat_url, cfg.page_timeout_ms)
            input("\n>> Press ENTER when ChatGPT is ready (logged in, new chat open)…\n")
            if not cfg.keep_chat_tabs:
                login_page.close()
        else:
            warmup_page = open_chat_tab(context, chat_url, cfg.page_timeout_ms)
            print("ChatGPT открыт.\n")
            if not cfg.keep_chat_tabs:
                warmup_page.close()

        saved = load_progress(progress_file)
        if saved:
            print(f"Последний сохранённый прогресс: анг {saved}")
        else:
            print("Последний сохранённый прогресс: нет")
        print(f"Работаю строго по диапазону: {args.start}..{args.end}")
        print(f"Файл прогресса: {progress_file}")
        print(f"Папка JSON: {cfg.json_dir}")

        # Загружаем карту шабадов и обновляем из уже существующих JSON
        shabad_map = load_shabad_map()
        map_from_json = build_shabad_map_from_json(cfg.json_dir)
        if map_from_json != {k: shabad_map.get(k) for k in map_from_json}:
            shabad_map.update(map_from_json)
            save_shabad_map(shabad_map)

        # Ищем стартовый шабад: сначала карта, затем ближайший предыдущий JSON
        last_shabad = 0
        if args.start > 1:
            for look_back in range(args.start - 1, max(0, args.start - 100), -1):
                prev = load_ang_json(cfg.json_dir, look_back)
                if prev and prev.lines:
                    last_shabad = max(line.shabad_num for line in prev.lines)
                    break

        for ang in range(args.start, args.end + 1):
            print(f"\n══ Анг {ang}/{args.end} ══")

            existing_json = ang_json_path(cfg.json_dir, ang)
            if existing_json.exists():
                if args.force_retranslate:
                    existing_json.unlink()
                    print(f"  ↺ Перевожу заново анг {ang}: удалён старый {existing_json.name}")
                else:
                    cached = load_ang_json(cfg.json_dir, ang)
                    if cached and cached.lines:
                        last_shabad = max(line.shabad_num for line in cached.lines)
                    print(f"  ↷ JSON уже существует, пропускаю анг {ang}: {existing_json.name}")
                    continue

            # Определяем start_probe: точно из карты, иначе last_shabad, иначе интерполяция
            if ang in shabad_map:
                start_probe = shabad_map[ang]
                miss_limit = 10
            elif last_shabad > 0:
                start_probe = last_shabad
                miss_limit = 10
            else:
                start_probe = estimate_start_probe(ang, shabad_map)
                miss_limit = 300
                print(f"  → Интерполированный старт шабада: №{start_probe} (окно +300)")

            valid_verse_ids: set[int] | None = None
            if cfg.banidb_path:
                valid_verse_ids = banidb_verse_ids_for_ang(cfg.banidb_path, ang)
                if valid_verse_ids:
                    print(f"  → banidb: ожидается {len(valid_verse_ids)} verse_id для анга {ang}")

            source_lines = fetch_ang_source_lines(
                ang,
                start_probe=start_probe,
                initial_miss_limit=miss_limit,
                valid_verse_ids=valid_verse_ids,
            )
            if not source_lines:
                print(f"  ⚠ Не удалось собрать строки из KhojGurbani для анга {ang}, пропускаю")
                continue

            print(f"  → Собрано строк: {len(source_lines)}")

            ang_data = _translate_source_lines(context, chat_url, ang, source_lines, cfg)

            if ang_data:
                json_path = save_ang_json(cfg.json_dir, ang_data)
                print(f"  ✓ JSON сохранён: {json_path.name}")
                append_ang_to_docx(output_path, ang_data)
                save_progress(progress_file, ang)

            # Обновляем last_shabad и карту
            if source_lines:
                last_shabad = max(line.shabad_num for line in source_lines)
                min_sh = min(line.shabad_num for line in source_lines)
                if shabad_map.get(ang) != min_sh:
                    shabad_map[ang] = min_sh
                    save_shabad_map(shabad_map)

            if not ang_data:
                print(f"  ⚠ Не удалось получить валидный результат для анга {ang}, пропускаю")
                continue

            if ang < args.end:
                print(f"  ⏳ Пауза {args.delay}с...")
                time.sleep(args.delay)

        print(f"\n✓ Готово! Результат: {output_path}")
        print(f"✓ Прогресс: {progress_file}")
        print(f"✓ JSON по ангам: {cfg.json_dir}")
        if cfg.raw_log_dir:
            print(f"✓ Сырые логи: {cfg.raw_log_dir}")

        context.close()


def run_interactive_menu(
    args: argparse.Namespace,
    cfg: RuntimeConfig,
    output_path: Path,
    progress_file: Path,
) -> None:
    """Интерактивное меню: показывает текущий прогресс и предлагает режимы работы."""
    progress = load_progress(progress_file)
    existing_count = sum(
        1 for a in range(1, 1431) if ang_json_path(cfg.json_dir, a).exists()
    )
    found_angs = sorted(
        int(p.stem[4:]) for p in cfg.json_dir.glob("ang_*.json") if p.stem[4:].isdigit()
    )
    highest_json = found_angs[-1] if found_angs else 0
    last_done = max(progress, highest_json)
    next_ang = last_done + 1 if last_done < 1430 else None
    default_end = min(last_done + 15, 1430) if next_ang else None
    scan_up_to = max(last_done, 1)

    print("\n══ KhojGurbani Sahib Singh Bot ══\n")
    print(f"  JSON готово:               {existing_count} из 1430")
    if last_done:
        print(f"  Последний сохранённый анг: {last_done}")
    else:
        print("  Прогресс: не начат")
    print()
    print("  Что делаем?")
    if next_ang:
        print(f"  1) Продолжить перевод            → анги {next_ang}..{default_end}")
    print("  2) Свой диапазон ангов")
    print(f"  3) Пропущенные анги              → сканировать 1..{scan_up_to}")
    print(f"  4) Битые roman (авто-фикс)       → исправить 1..{scan_up_to} без ChatGPT")
    print(f"  5) Битые translation_ru          → переперевести через ChatGPT")
    print("  6) Пересобрать DOCX из JSON")
    print("  0) Выход\n")

    choice = input("  Выбор: ").strip()
    print()

    if choice == "0":
        print("Выход.")

    elif choice == "1" and next_ang:
        args.start = next_ang
        args.end = default_end
        print(f"Переводить анги {args.start}..{args.end}\n")
        run_browser_session(args, cfg, output_path, progress_file)

    elif choice == "2":
        try:
            args.start = int(input("  Начальный анг: ").strip())
            args.end = int(input("  Конечный анг:  ").strip())
        except ValueError:
            print("  ✗ Некорректный ввод")
            return
        if args.start < 1 or args.end < 1 or args.start > args.end:
            print("  ✗ Некорректный диапазон")
            return
        print(f"\nПереводить анги {args.start}..{args.end}\n")
        run_browser_session(args, cfg, output_path, progress_file)

    elif choice == "3":
        missing = [a for a in range(1, scan_up_to + 1) if not ang_json_path(cfg.json_dir, a).exists()]
        if not missing:
            print(f"✓ Пропущенных ангов нет в диапазоне 1..{scan_up_to}")
            return
        print(f"Пропущено ангов ({len(missing)}):")
        print("  " + ", ".join(str(a) for a in missing))
        if input("\nДополнить через ChatGPT? [y/N] ").strip().lower() != "y":
            print("Отменено.")
            return
        args.start = missing[0]
        args.end = missing[-1]
        print(f"\nЗапускаю перевод пропущенных ангов {args.start}..{args.end}...\n")
        run_browser_session(args, cfg, output_path, progress_file)

    elif choice == "4":
        print(f"\nАвто-фикс roman в диапазоне 1..{scan_up_to}...")
        fixed = fix_corrupt_roman_in_json(cfg.json_dir, 1, scan_up_to)
        if not fixed:
            print("✓ Битых строк в roman не найдено — ничего не изменено.")
        else:
            total = sum(fixed.values())
            print(f"✓ Исправлено {total} строк в {len(fixed)} ангах:")
            for ang_num, cnt in sorted(fixed.items()):
                print(f"  Анг {ang_num}: {cnt} строк")

    elif choice == "5":
        print(f"Сканирую анги 1..{scan_up_to} на битые translation_ru...")
        corrupt = scan_corrupt_angs(cfg.json_dir, 1, scan_up_to)
        # Оставляем только те, у кого translation_ru битый
        corrupt_tr = {a: [i for i in issues if "translation_ru" in i.issue]
                      for a, issues in corrupt.items()}
        corrupt_tr = {a: v for a, v in corrupt_tr.items() if v}
        if not corrupt_tr:
            print(f"✓ Битых translation_ru не найдено в диапазоне 1..{scan_up_to}")
            return
        total_lines = sum(len(v) for v in corrupt_tr.values())
        print(f"\nНайдено {total_lines} битых строк в {len(corrupt_tr)} ангах:\n")
        for ang_num, issues in sorted(corrupt_tr.items()):
            print(f"  Анг {ang_num} ({len(issues)} строк):")
            for info in issues[:3]:
                print(f"    #{info.line_index}: {info.snippet[:90]}")
            if len(issues) > 3:
                print(f"    ... ещё {len(issues) - 3} строк")
        if input("\nПереперевести эти анги через ChatGPT? [y/N] ").strip().lower() != "y":
            print("Отменено.")
            return
        corrupt_angs = sorted(corrupt_tr.keys())
        for ang_num in corrupt_angs:
            p = ang_json_path(cfg.json_dir, ang_num)
            if p.exists():
                p.unlink()
                print(f"  ✗ Удалён JSON анга {ang_num}")
        args.start = corrupt_angs[0]
        args.end = corrupt_angs[-1]
        print(f"\nЗапускаю переперевод ангов {args.start}..{args.end}...\n")
        run_browser_session(args, cfg, output_path, progress_file)

    elif choice == "6":
        rb_start_s = input("  Начальный анг [1]: ").strip() or "1"
        rb_end_s = input("  Конечный анг [1430]: ").strip() or "1430"
        try:
            rb_start, rb_end = int(rb_start_s), int(rb_end_s)
        except ValueError:
            print("  ✗ Некорректный ввод")
            return
        print(f"\nПересобираю DOCX из JSON ({rb_start}..{rb_end})...")
        rebuilt = rebuild_docx_from_json(output_path, cfg.json_dir, rb_start, rb_end)
        print(f"✓ Пересобрано ангов: {rebuilt}")
        print(f"✓ Результат: {output_path}")

    else:
        print("  ✗ Неизвестный выбор")


def main() -> None:
    args = parse_args()

    if len(sys.argv) == 1:
        args.menu = True

    output_path = (Path(__file__).parent / args.output).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    progress_file = build_progress_path(output_path)
    if args.reset_progress:
        reset_progress(progress_file)

    raw_log_dir = None
    if args.raw_log_dir.strip():
        raw_log_dir = (Path(__file__).parent / args.raw_log_dir).resolve()

    json_dir = (Path(__file__).parent / args.json_dir).resolve()

    cfg = RuntimeConfig(
        page_timeout_ms=args.page_timeout * 1000,
        input_timeout_ms=args.input_timeout * 1000,
        response_timeout_ms=args.response_timeout * 1000,
        new_message_timeout_ms=args.new_message_timeout * 1000,
        max_retries=args.max_retries,
        retry_delay_s=args.retry_delay,
        raw_log_dir=raw_log_dir,
        json_dir=json_dir,
        keep_chat_tabs=args.keep_chat_tabs,
        chunk_size=args.chunk_size,
        banidb_path=Path(args.banidb) if args.banidb.strip() else None,
    )

    if args.update_shabad_map:
        print("Пересобираю shabad_map.json из ang_json/...")
        shabad_map = build_shabad_map_from_json(json_dir)
        save_shabad_map(shabad_map)
        print(f"✓ Сохранено {len(shabad_map)} записей в {SHABAD_MAP_PATH}")
        return

    if args.menu:
        run_interactive_menu(args, cfg, output_path, progress_file)
        return

    if args.start < 1 or args.end < 1 or args.start > args.end:
        print("  ✗ Некорректный диапазон: проверь --start и --end")
        return

    if args.reset_json_range:
        removed_json = reset_json_range(cfg.json_dir, args.start, args.end)
        print(f"Сброс JSON в диапазоне {args.start}..{args.end}: удалено {removed_json}")

    if args.scan_missing:
        missing = [
            ang for ang in range(args.start, args.end + 1)
            if not ang_json_path(cfg.json_dir, ang).exists()
        ]
        if not missing:
            print(f"✓ В диапазоне {args.start}..{args.end} все JSON на месте.")
            return
        print(f"Пропущено ангов ({len(missing)}) в диапазоне {args.start}..{args.end}:")
        print("  " + ", ".join(str(a) for a in missing))
        answer = input("\nДополнить через ChatGPT? [y/N] ").strip().lower()
        if answer != "y":
            print("Отменено.")
            return
        args.start = missing[0]
        args.end = missing[-1]
        print(f"Запускаю перевод пропущенных ангов {args.start}..{args.end}...\n")

    if args.scan_corrupt:
        print(f"Сканирую анги {args.start}..{args.end} на наличие битых строк...")
        corrupt = scan_corrupt_angs(cfg.json_dir, args.start, args.end)
        if not corrupt:
            print(f"✓ Битых строк не найдено в диапазоне {args.start}..{args.end}")
            return
        total_lines = sum(len(v) for v in corrupt.values())
        print(f"\nНайдено {total_lines} битых строк в {len(corrupt)} ангах:\n")
        for ang_num, issues in sorted(corrupt.items()):
            print(f"  Анг {ang_num} ({len(issues)} строк):")
            for info in issues[:3]:
                print(f"    #{info.line_index} (verse {info.verse_id}): {info.issue}")
                print(f"      {info.snippet[:90]}")
            if len(issues) > 3:
                print(f"    ... ещё {len(issues) - 3} строк")
        answer = input("\nПереперевести эти анги через ChatGPT? [y/N] ").strip().lower()
        if answer != "y":
            print("Отменено.")
            return
        corrupt_angs = sorted(corrupt.keys())
        for ang_num in corrupt_angs:
            p = ang_json_path(cfg.json_dir, ang_num)
            if p.exists():
                p.unlink()
                print(f"  ✗ Удалён JSON анга {ang_num}")
        args.start = corrupt_angs[0]
        args.end = corrupt_angs[-1]
        print(f"Запускаю переперевод ангов {args.start}..{args.end}...\n")

    if args.fix_corrupt_roman:
        print(f"Авто-фикс roman в диапазоне {args.start}..{args.end}...")
        fixed = fix_corrupt_roman_in_json(cfg.json_dir, args.start, args.end)
        if not fixed:
            print("✓ Битых строк в roman не найдено — ничего не изменено.")
        else:
            total = sum(fixed.values())
            print(f"✓ Исправлено {total} строк в {len(fixed)} ангах:")
            for ang_num, cnt in sorted(fixed.items()):
                print(f"  Анг {ang_num}: {cnt} строк")
        return

    if args.rebuild_docx_from_json:
        print("Пересобираю DOCX из JSON...")
        rebuilt = rebuild_docx_from_json(output_path, cfg.json_dir, args.start, args.end)
        if rebuilt > 0:
            save_progress(progress_file, args.end)
        print(f"✓ Пересобрано ангов: {rebuilt}")
        print(f"✓ Результат: {output_path}")
        print(f"✓ JSON по ангам: {cfg.json_dir}")
        return

    run_browser_session(args, cfg, output_path, progress_file)


if __name__ == "__main__":
    main()
