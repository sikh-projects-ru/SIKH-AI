"""
darpan_rebuild.py

Читает локальный docx Guru Granth Darpan (проф. Сахиб Сингх) и пересобирает
его в структурированный русский docx.

Источник переводов: gpt_darpan_python.docx
  - Бани-строки → русский перевод (тёмно-красный) + транслитерация (коричневый)
  - PadArath    → пад-артх из gpt docx (фиолетовый)
  - Arath       → перевод (синий)
  - Bhav        → перевод (бирюзовый)
  - Note        → перевод (зелёный)

Запуск:
    python3 darpan_rebuild.py
    python3 darpan_rebuild.py --start 448 --end 746
    python3 darpan_rebuild.py --reset
"""

import argparse
import json
import re
import unicodedata
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SOURCE_DOCX  = Path(__file__).parent / 'reference_material' / 'GuruGranth Darpan by Prof Sahib Singh (Uni).docx'
GPT_DOCX     = Path(__file__).parent / 'gpt_darpan_python.docx'
OUTPUT_DEFAULT = 'Darpan_Rebuilt.docx'
CONTENT_START  = 448
SAVE_INTERVAL  = 500

# ── Маппинг стилей ─────────────────────────────────────────────────────────────
STYLE_TO_CLS = {
    'Bani2':         'bani',
    'Bani-Centered': 'bani',
    'Arath':         'arath',
    'PadArath':      'padarth',
    'Note':          'note',
    'Bhav':          'bhav',
    'Body Text1':    'blackuni',
    'text':          'blackuni',
    'text bold':     'baniblack',
    'side title':    'sidetitle',
    'title1':        'title1',
}

# ── Цвета ──────────────────────────────────────────────────────────────────────
COLORS = {
    'bani':      RGBColor(0x80, 0x00, 0x00),
    'translit':  RGBColor(0x8B, 0x45, 0x13),
    'padarth':   RGBColor(0x80, 0x00, 0x80),
    'arath':     RGBColor(0x00, 0x00, 0x80),
    'bhav':      RGBColor(0x00, 0x80, 0x80),
    'note':      RGBColor(0x00, 0x80, 0x00),
    'sidetitle': RGBColor(0x00, 0x00, 0x00),
    'title1':    RGBColor(0x00, 0x00, 0x00),
    'blackuni':  RGBColor(0x00, 0x00, 0x00),
}

CLS_COLOR = {
    'padarth':  'padarth',
    'arath':    'arath',
    'bhav':     'bhav',
    'note':     'note',
    'blackuni': 'blackuni',
    'baniblack':'blackuni',
    'sidetitle':'sidetitle',
    'title1':   'title1',
}

LABEL = {
    'padarth': 'Пад-артх: ',
    'arath':   'Арат: ',
    'bhav':    'Бхав: ',
    'note':    'Примечание: ',
}


# ── Утилиты ───────────────────────────────────────────────────────────────────
def _norm(s: str) -> str:
    return unicodedata.normalize('NFC', s).replace('\u2019', "'").replace('\u2018', "'")

def _has_gurmukhi(text: str) -> bool:
    return any('\u0A00' <= c <= '\u0A7F' for c in text)

def _has_cyrillic(text: str) -> bool:
    return any('\u0400' <= c <= '\u04FF' for c in text)

def _guru_only(text: str) -> str:
    """Оставляем только символы гурмукхи и пробелы (нормализуем пробелы)."""
    raw = ''.join(c for c in text if '\u0A00' <= c <= '\u0A7F' or c == ' ')
    return ' '.join(raw.split())

def _starts_guru(text: str) -> bool:
    """True если строка начинается с символа гурмукхи.
    Допускает префикс ॥ (U+0965, Devanagari double danda) — встречается в заголовках типа '॥ ਜਪੁ ॥'.
    """
    if not text:
        return False
    # Пропускаем dandas и пробелы в начале
    t = text.lstrip('\u0964\u0965 ')
    return bool(t) and '\u0A00' <= t[0] <= '\u0A7F'

def _is_pure_guru_line(text: str) -> bool:
    """True если строка — чистая бани-строка (только гурмукхи + знаки препинания).
    Отфильтровывает смешанные строки вида "ਵਾਰਿਆ (vāriā) — я не могу...".
    """
    if not _starts_guru(text):
        return False
    if _has_cyrillic(text):
        return False
    # Если есть латинские буквы (признак транслитерации в скобках), отклоняем
    if any('a' <= c.lower() <= 'z' for c in text):
        return False
    return True


# ── Парсинг gpt_darpan_python.docx ───────────────────────────────────────────
# Структура GPT-документа — два формата:
#
#  Format A (паури 1-2, подробный разбор):
#    - параграф: только гурмукхи (может быть многострочным через \n)
#    - параграф: "Transliteration:\n translit text"
#    - ... много комментариев ...
#    - параграф/заголовок: "Перевод:..." / "Смысл..." с русским переводом
#
#  Format B (паури 3+, компактный):
#    - параграф: "гурмукхи\ntranslit\nрусский перевод" (всё в одном)
#
def load_gpt_translations(docx_path: Path):
    """
    Возвращает:
      bani_pairs   — list[(gurmukhi_key, russian_text)], сортировка: длиннее вперёд
      translit_map — dict[gurmukhi_key → roman_text]
      padarth_pairs— list[(gurmukhi_word, russian_explanation)], длиннее вперёд
    """
    doc = Document(str(docx_path))
    paras = doc.paragraphs
    n = len(paras)

    bani_dict    = {}   # gurmukhi_key → russian_text (дедупликация через dict)
    translit_map = {}
    padarth_dict = {}
    commentary_dict: dict[str, list[str]] = {}  # gurmukhi_key → список комментариев

    # Метки перевода в Format A
    PEREVOD_PREFIXES = (
        'Перевод на русский:', 'Перевод:',
        'Перевод смысла', 'Перевод (ਅਰਥ)',
        'Перевод пары строк:', 'Перевод этих двух строк:',
        'Смысл первых двух строк:', 'Смысл двух строк:', 'Смысл:',
    )

    in_padarth = False
    # Состояние Format A: список (gurmukhi_key, translit_or_None)
    pending: list[tuple[str, str | None]] = []
    pending_commentary: list[str] = []  # комментарии, накопленные до перевода

    def _commit(ru_text: str) -> None:
        for key, tr in pending:
            bani_dict.setdefault(key, ru_text)
            if tr:
                translit_map.setdefault(key, tr)
            if pending_commentary:
                commentary_dict.setdefault(key, []).extend(pending_commentary)

    i = 0
    while i < n:
        raw   = paras[i].text
        text  = raw.strip()
        style = paras[i].style.name

        if not text:
            i += 1
            continue

        # ── Заголовки ─────────────────────────────────────────────────────────
        if style.startswith('Heading'):
            tl = text.lower()
            is_padarth = ('ਪਦ ਅਰਥ' in text or ('Разбор' in text and 'слов' in text)
                          or ('пад' in tl and 'арт' in tl))
            is_perevod = any(text.startswith(p.rstrip(':')) for p in PEREVOD_PREFIXES)

            if is_padarth:
                in_padarth = True
            elif is_perevod and pending:
                # Заголовок "Перевод смысла (ਅਰਥ)" — перевод в следующих параграфах
                j, parts = i + 1, []
                while j < n:
                    t2, s2 = paras[j].text.strip(), paras[j].style.name
                    if not t2:
                        j += 1; continue
                    if s2.startswith('Heading'):
                        break
                    if _is_pure_guru_line(t2):
                        break
                    if _has_cyrillic(t2):
                        parts.append(t2.replace('\n', ' '))
                    j += 1
                if parts:
                    _commit(' '.join(parts))
                    pending.clear()
                    pending_commentary.clear()
                in_padarth = False
            else:
                in_padarth = False
            i += 1
            continue

        # ── Пад-артх ──────────────────────────────────────────────────────────
        # Минимальная длина ключа 8 симв. — фильтрует короткие слова, которые
        # матчатся в цитатах в тексте Сахиб Сингха (напр. ਪਦ ਅਰਥ = 6 симв.)
        MIN_PADARTH_KEY = 4
        if in_padarth:
            if _has_gurmukhi(text) and ('—' in text or '–' in text):
                dash = '—' if '—' in text else '–'
                left, _, right = text.partition(dash)
                right = right.strip()
                if _has_cyrillic(right):
                    key = _norm(_guru_only(left))
                    if key and len(key) >= MIN_PADARTH_KEY:
                        padarth_dict.setdefault(key, left.strip() + ' — ' + right)
            i += 1
            continue

        # ── Параграф с гурмукхи ───────────────────────────────────────────────
        if _starts_guru(text) and _has_gurmukhi(text):
            lines = [l.strip() for l in raw.split('\n') if l.strip()] if '\n' in raw else [text]
            # Чистые бани-строки: только гурмукхи, без кириллицы и латиницы
            pure_lines = [l for l in lines if _is_pure_guru_line(l)]
            # Прочие строки (транслит, перевод, смешанные)
            other_lines = [l for l in lines if not _is_pure_guru_line(l)]

            if other_lines and pure_lines:
                # ── Format B: гурмукхи + транслит + русский в одном параграфе ──
                tr_lines, ru_lines = [], []
                state = 'guru'
                for line in lines:
                    if line.startswith('Перевод на русский:'):
                        state = 'ru'
                        rest = line[len('Перевод на русский:'):].strip()
                        if rest:
                            ru_lines.append(rest)
                    elif line.startswith('Transliteration:'):
                        state = 'tr'
                        rest = line[len('Transliteration:'):].strip()
                        if rest:
                            tr_lines.append(rest)
                    elif state == 'guru':
                        if _is_pure_guru_line(line):
                            pass  # уже в pure_lines
                        elif _has_cyrillic(line) and not _has_gurmukhi(line):
                            state = 'ru'; ru_lines.append(line)
                        elif line:
                            state = 'tr'; tr_lines.append(line)
                    elif state == 'tr':
                        if _has_cyrillic(line) and not _has_gurmukhi(line):
                            state = 'ru'; ru_lines.append(line)
                        else:
                            tr_lines.append(line)
                    else:
                        if _has_cyrillic(line):
                            ru_lines.append(line)

                if pure_lines:
                    ru_text = ' '.join(ru_lines) if ru_lines else None
                    combined_key = _norm(' '.join(pure_lines))

                    # Сохраняем каждую строку гурмукхи как отдельный ключ
                    for j2, gline in enumerate(pure_lines):
                        key = _norm(gline)
                        if ru_text:
                            bani_dict.setdefault(key, ru_text)
                            # Транслит — только если есть русский перевод
                            # (иначе загрязняем translit пенджабскими объяснительными текстами)
                            if tr_lines:
                                tr = tr_lines[j2] if j2 < len(tr_lines) else tr_lines[-1]
                                translit_map.setdefault(key, tr)

                    # И комбинированный ключ (обе строки вместе)
                    if len(pure_lines) > 1 and ru_text:
                        bani_dict.setdefault(combined_key, ru_text)
                        if tr_lines:
                            translit_map.setdefault(combined_key, ' '.join(tr_lines))

                # Format B завершён — сбрасываем Format A ожидание
                pending.clear()
                pending_commentary.clear()

            elif '—' in text or '–' in text:
                dash = '—' if '—' in text else '–'
                left, _, right = text.partition(dash)
                right = right.strip()
                if _has_cyrillic(right):
                    key = _norm(_guru_only(left))
                    if key:
                        left_has_latin = any('a' <= c.lower() <= 'z' for c in left)
                        if left_has_latin:
                            # "ਆਦਿ (āḍi) — с самого начала" — словарное объяснение
                            # → padarth_dict, не bani_dict (иначе загрязняет матчинг бани)
                            if len(key) >= MIN_PADARTH_KEY:
                                padarth_dict.setdefault(key, left.strip() + ' — ' + right)
                        else:
                            # "॥ ਜਪੁ ॥ — «Джапу»" — заголовок/бани → bani_dict
                            bani_dict.setdefault(key, right)

            elif pure_lines:
                # ── Format A: только чистые строки гурмукхи, ждём транслит и перевод ──
                new_pending = []
                for gline in pure_lines:
                    new_pending.append((_norm(gline), None))
                if len(pure_lines) > 1:
                    # Добавляем комбинированный ключ
                    new_pending.append((_norm(' '.join(pure_lines)), None))
                pending = new_pending
                pending_commentary.clear()

            i += 1
            continue

        # ── Транслитерация (Format A) ─────────────────────────────────────────
        if text.startswith('Transliteration:') and pending:
            tr_raw = raw.replace('Transliteration:', '', 1).strip()
            tr_lines = [l.strip() for l in tr_raw.split('\n') if l.strip()]
            tr_text_all = tr_raw.replace('\n', ' ')

            new_pending = []
            for j2, (key, _) in enumerate(pending):
                if j2 < len(tr_lines):
                    new_pending.append((key, tr_lines[j2]))
                else:
                    new_pending.append((key, tr_text_all))
            pending = new_pending
            i += 1
            continue

        # ── Перевод / Комментарий (Format A) ─────────────────────────────────
        if pending and _has_cyrillic(text):
            for pfx in PEREVOD_PREFIXES:
                if text.startswith(pfx):
                    ru_text = raw[len(pfx):].strip().replace('\n', ' ')
                    if not ru_text.strip():
                        # Перевод в следующих параграфах
                        j, parts = i + 1, []
                        while j < n:
                            t2, s2 = paras[j].text.strip(), paras[j].style.name
                            if not t2:
                                j += 1; continue
                            if s2.startswith('Heading'):
                                break
                            if _is_pure_guru_line(t2):
                                break
                            if _has_cyrillic(t2):
                                parts.append(t2.replace('\n', ' '))
                            j += 1
                        ru_text = ' '.join(parts)
                    if ru_text:
                        _commit(ru_text)
                        pending.clear()
                        pending_commentary.clear()
                    break
            else:
                # Нет метки перевода → это комментарий к текущему бани-блоку
                pending_commentary.append(text.replace('\n', ' '))

        i += 1

    # Преобразуем dict → отсортированный список (длиннее ключи — точнее совпадение)
    bani_pairs     = sorted(bani_dict.items(),     key=lambda x: -len(x[0]))
    translit_pairs = sorted(translit_map.items(),  key=lambda x: -len(x[0]))
    padarth_pairs  = sorted(padarth_dict.items(),  key=lambda x: -len(x[0]))

    return bani_pairs, translit_pairs, padarth_pairs, commentary_dict


# ── Функции поиска ─────────────────────────────────────────────────────────────
def get_translation(text: str, pairs: list) -> str | None:
    text_n = _norm(text)
    for key_n, val in pairs:
        if key_n and key_n in text_n:
            return val
    return None

def get_translit(text: str, translit_pairs: list) -> str | None:
    text_n = _norm(text)
    for key_n, val in translit_pairs:
        if key_n and key_n in text_n:
            return val
    return None

def get_all_padarth(text: str, padarth_pairs: list) -> list[str]:
    """Все пад-артх объяснения, чей ключ есть в тексте."""
    text_n = _norm(text)
    results = []
    seen = set()
    for key_n, val in padarth_pairs:
        if key_n and key_n not in seen and key_n in text_n:
            results.append(val)
            seen.add(key_n)
    return results

def get_commentary(text: str, commentary_dict: dict) -> list[str]:
    """Все комментарии из GPT-документа, чей бани-ключ входит в текст."""
    text_n = _norm(text)
    results = []
    seen_keys = set()
    for key_n, comments in sorted(commentary_dict.items(), key=lambda x: -len(x[0])):
        if key_n and key_n not in seen_keys and key_n in text_n:
            results.extend(comments)
            seen_keys.add(key_n)
    return results


# ── Запись параграфа ──────────────────────────────────────────────────────────
def add_para(doc: Document, text: str, color: RGBColor, size_pt: float = 11,
             bold: bool = False, italic: bool = False,
             space_before: int = 0, space_after: int = 0,
             align_center: bool = False) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align_center:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def add_separator(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Прогресс ──────────────────────────────────────────────────────────────────
def _progress_file(output_path: Path) -> Path:
    return output_path.with_suffix('.progress')

def load_progress(output_path: Path) -> dict:
    pf = _progress_file(output_path)
    if pf.exists():
        try:
            return json.loads(pf.read_text())
        except Exception:
            pass
    return {'last_para': -1, 'translated': 0, 'untranslated': 0, 'bani': 0}

def save_progress(output_path: Path, state: dict) -> None:
    _progress_file(output_path).write_text(json.dumps(state))


# ── Основная сборка ───────────────────────────────────────────────────────────
def build(para_start: int, para_end: int | None, output_path: Path, reset: bool = False) -> None:
    # Загружаем переводы из gpt_darpan_python.docx
    print(f"Загружаю переводы из: {GPT_DOCX.name}")
    bani_pairs, translit_pairs, padarth_pairs, commentary_dict = load_gpt_translations(GPT_DOCX)
    print(f"  Бани-переводов: {len(bani_pairs)}, транслитераций: {len(translit_pairs)}, пад-артх: {len(padarth_pairs)}, комментариев: {len(commentary_dict)}")

    print(f"Читаю источник: {SOURCE_DOCX.name}")
    src = Document(str(SOURCE_DOCX))
    paragraphs = src.paragraphs

    total = len(paragraphs)
    end = para_end if para_end is not None else total
    end = min(end, total)

    # ── Прогресс / resume ────────────────────────────────────────────────────
    state = load_progress(output_path)
    if reset or state['last_para'] < para_start:
        state = {'last_para': -1, 'translated': 0, 'untranslated': 0, 'bani': 0}
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
        if reset:
            print("⟳ Сброс прогресса, начинаю заново.")
        else:
            print(f"Параграфы {para_start}–{end} из {total}")
    else:
        para_start = state['last_para'] + 1
        print(f"↩ Возобновляю с параграфа {para_start} (до {end}), "
              f"уже переведено: {state['translated']}, пропущено: {state['untranslated']}")
        doc = Document(str(output_path)) if output_path.exists() else Document()
        if not output_path.exists():
            section = doc.sections[0]
            section.top_margin = Cm(2); section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    if para_start >= end:
        print("✓ Документ уже собран полностью.")
        return

    total_range = end - para_start
    processed_since_save = 0
    prev_cls = None
    last_shown_bani_ru: str | None = None  # против дублей при парных бани-строках
    current_bani_block: str | None = None  # накопленный текст текущего бани-блока
    shown_padarth: set[str] = set()        # против дублей пад-артха внутри блока
    shown_commentary: set[str] = set()     # против дублей комментариев внутри блока

    for i, para in enumerate(paragraphs[para_start:end], start=para_start):
        text = para.text.strip()
        if not text:
            continue

        done = i - para_start + 1
        if done % 200 == 0:
            pct = round(done / total_range * 100)
            total_done = state['translated'] + state['untranslated']
            tr_pct = round(state['translated'] / total_done * 100) if total_done else 0
            print(f"  [{pct}%] параграф {i}/{end} | переведено {state['translated']} ({tr_pct}%)")

        text = _norm(text)
        cls = STYLE_TO_CLS.get(para.style.name, 'blackuni')

        # ── Bani ─────────────────────────────────────────────────────────────
        if cls in ('bani', 'banicenter'):
            # Паури могут быть одним параграфом с несколькими строками через \n
            raw_lines = [l.strip() for l in para.text.split('\n') if l.strip()]
            if not raw_lines:
                raw_lines = [text]

            for line in raw_lines:
                line = _norm(line)
                # Пропускаем нечистые строки: пад-артх, смешанные, пустые
                if not line or not _is_pure_guru_line(line):
                    continue
                state['bani'] += 1
                add_para(doc, line, COLORS['bani'], size_pt=12, bold=True,
                         space_before=6, align_center=True)

                tr = get_translit(line, translit_pairs)
                if tr:
                    add_para(doc, tr, COLORS['translit'], size_pt=10, italic=True,
                             align_center=True)

                ru = get_translation(line, bani_pairs)
                if ru:
                    state['translated'] += 1
                    if ru != last_shown_bani_ru:
                        add_para(doc, ru, COLORS['bani'], size_pt=11, align_center=True,
                                 space_after=4)
                        last_shown_bani_ru = ru
                else:
                    state['untranslated'] += 1
                    last_shown_bani_ru = None

        # ── Всё остальное — пропускаем ────────────────────────────────────────
        elif cls in ('padarth', 'arath', 'bhav', 'note', 'blackuni', 'baniblack'):
            pass

        # ── Заголовки ─────────────────────────────────────────────────────────
        elif cls in ('sidetitle', 'title1'):
            ru = get_translation(text, bani_pairs)
            if ru:
                state['translated'] += 1
                doc.add_heading(ru, level=2)
            else:
                state['untranslated'] += 1
                doc.add_heading(text, level=2)

        # ── Обычный текст ─────────────────────────────────────────────────────
        else:
            ru = get_translation(text, bani_pairs)
            if ru:
                state['translated'] += 1
                add_para(doc, ru, COLORS['blackuni'], size_pt=11)
            else:
                state['untranslated'] += 1

        prev_cls = cls
        state['last_para'] = i
        processed_since_save += 1

        if processed_since_save >= SAVE_INTERVAL:
            doc.save(str(output_path))
            save_progress(output_path, state)
            processed_since_save = 0
            print(f"  💾 Сохранено ({output_path.name}), параграф {i}")

    doc.save(str(output_path))
    save_progress(output_path, state)

    total_blocks = state['translated'] + state['untranslated']
    pct = round(state['translated'] / total_blocks * 100) if total_blocks else 0
    print(f"\n✓ Готово: {output_path.name}")
    print(f"  Бани-блоков:    {state['bani']}")
    print(f"  Переведено:     {state['translated']}/{total_blocks} ({pct}%)")
    print(f"  Пропущено:      {state['untranslated']}")


# ── CLI ────────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description='Rebuild Darpan → structured Russian docx')
    parser.add_argument('--start', type=int, default=CONTENT_START)
    parser.add_argument('--end',   type=int, default=None)
    parser.add_argument('--output', type=str, default=OUTPUT_DEFAULT)
    parser.add_argument('--reset', action='store_true',
                        help='Начать заново, игнорируя сохранённый прогресс')
    args = parser.parse_args()

    output = Path(__file__).parent / args.output
    build(args.start, args.end, output, reset=args.reset)


if __name__ == '__main__':
    main()
