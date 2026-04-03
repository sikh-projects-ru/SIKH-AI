#!/usr/bin/env python3
"""
KhojGurbani Translation Tool — интерактивный режим
Workflow:
  1. python3 khojgurbani_translate.py fetch --ang 1 --shabad 1
     → выводит исходники стихов + комментарий KG
     → Claude переводит в чате, записывает в khojgurbani_cache.json

  2. python3 khojgurbani_translate.py build --ang 1 --shabad 1
     → читает кэш, генерирует KhojGurbani_Russian.docx

  Диапазон: --shabad 1 --end-shabad 4
"""

import argparse
import json
import re
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

# ── Пути ────────────────────────────────────────────────────────────────────
BASE_DIR    = Path(__file__).parent
CACHE_FILE  = BASE_DIR / "khojgurbani_cache.json"
OUTPUT_FILE = BASE_DIR / "KhojGurbani_Russian.docx"

# ── API ──────────────────────────────────────────────────────────────────────
API_BASE = "https://apiprod.khojgurbani.org/api/v1"
HEADERS = {
    "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36",
    "Origin": "https://khojgurbani.org",
    "Referer": "https://khojgurbani.org/",
    "Accept": "application/json",
}

# ── Цвета ────────────────────────────────────────────────────────────────────
COLOR_GURBANI    = RGBColor(0x55, 0x00, 0x00)
COLOR_ROMAN      = RGBColor(0x44, 0x44, 0x44)
COLOR_KG         = RGBColor(0x00, 0x00, 0x80)   # тёмно-синий — KhojGurbani
COLOR_SS         = RGBColor(0x00, 0x60, 0x00)   # тёмно-зелёный — Prof. Sahib Singh
COLOR_COMMENTARY = RGBColor(0x33, 0x33, 0x55)   # тёмно-синий/серый — комментарий
COLOR_LABEL      = RGBColor(0x99, 0x99, 0x99)
COLOR_HEADING    = RGBColor(0x22, 0x22, 0x22)


# ── HTTP ─────────────────────────────────────────────────────────────────────
def api_get(path: str) -> dict:
    req = urllib.request.Request(f"{API_BASE}{path}", headers=HEADERS)
    with urllib.request.urlopen(req, timeout=20) as r:
        return json.loads(r.read())


# ── Кэш ─────────────────────────────────────────────────────────────────────
def load_cache() -> dict:
    if CACHE_FILE.exists():
        return json.loads(CACHE_FILE.read_text(encoding="utf-8"))
    return {}


def save_cache(cache: dict):
    CACHE_FILE.write_text(
        json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8"
    )


# ── Данные ───────────────────────────────────────────────────────────────────
def fetch_shabad(ang: int, shabad_num: int) -> dict | None:
    try:
        data = api_get(f"/shabad/{ang}/{shabad_num}")
        if data.get("status") == "success":
            return data["data"]
    except Exception as e:
        print(f"  [!] fetch_shabad({ang}/{shabad_num}): {e}")
    return None


def get_ang_shabads(ang: int) -> list[int]:
    """Возвращает список номеров стихов на данном анге."""
    # Probe with a known-working shabad number from this or nearby ang
    for probe in range(1, 1500):
        try:
            data = api_get(f"/shabad/{ang}/{probe}")
            if data.get("status") == "success":
                pages = data.get("pages", [])
                if pages:
                    return sorted(pages)
        except:
            pass
    return []


def fetch_commentary_text(shabad_id: int) -> str:
    """Возвращает официальный комментарий KG (stripped HTML) или пустую строку."""
    try:
        data = api_get(f"/commentary/list/{shabad_id}")
        results = data.get("result", [])
        # Take any result with non-empty commentary (approve_status check not reliable)
        for r in sorted(results, key=lambda x: x.get("approve_status", 0), reverse=True):
            text = strip_html(r.get("commentary", ""))
            if text:
                return text
    except Exception as e:
        print(f"  [!] commentary({shabad_id}): {e}")
    return ""


def strip_html(html: str) -> str:
    text = re.sub(r"<[^>]+>", "", html or "")
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"&amp;", "&", text)
    return text.strip()


# ── Команда: show (из локального дампа) ──────────────────────────────────────
def cmd_show(args):
    """Показывает данные шабадов из локального dump-файла (без сетевых запросов)."""
    dump_file = Path(args.dump)
    if not dump_file.exists():
        print(f"[!] Файл дампа не найден: {dump_file}")
        return
    dump = json.loads(dump_file.read_text(encoding="utf-8"))

    for shabad_num in range(args.shabad, (args.end_shabad or args.shabad) + 1):
        key = f"{args.ang}_{shabad_num}"
        entry = dump.get(key)
        if not entry:
            print(f"\n[!] АНГ {args.ang} шабад {shabad_num}: нет в дампе")
            continue

        shabad_id = entry.get("shabad_id")
        print(f"\n{'='*70}")
        print(f"АНГ {args.ang}  |  СТИХ {shabad_num}  (shabad_id={shabad_id})")
        print(f"{'='*70}")

        for verse in entry.get("scriptures", []):
            sid = verse["id"]
            print(f"\n── Строфа ID={sid} ──")
            print(f"  Gurmukhi : {verse['gurmukhi']}")
            print(f"  Roman    : {verse['roman']}")
            print(f"  KG (EN)  : {verse['kg_en'] if verse['kg_en'] else '(нет)'}")
            print(f"  SS (PA)  : {verse['ss_pa']}")
            print(f"  Кэш-ключи: kg_{sid}  /  ss_{sid}")

        commentary = entry.get("commentary_en", "")
        if commentary:
            print(f"\n── Комментарий KG (shabad_id={shabad_id}) ──")
            print(f"  Кэш-ключ: commentary_{shabad_id}")
            print(f"\n{commentary}\n")
        else:
            print(f"\n  (Комментарий отсутствует)")


# ── Команда: fetch ────────────────────────────────────────────────────────────
def cmd_fetch(args):
    for shabad_num in range(args.shabad, (args.end_shabad or args.shabad) + 1):
        print(f"\n{'='*70}")
        print(f"АНГ {args.ang}  |  СТИХ {shabad_num}")
        print(f"{'='*70}")

        shabad = fetch_shabad(args.ang, shabad_num)
        if not shabad:
            print("  [!] Нет данных")
            continue

        shabad_id  = shabad.get("id")
        scriptures = shabad.get("scriptures", [])

        for verse in scriptures:
            sid      = verse["id"]
            gurmukhi = verse.get("Scripture", "").strip()
            roman    = verse.get("ScriptureRoman", "").strip()
            t        = verse.get("translation", {})
            kg_src   = (t.get("KhojgurbaaniEnglish") or "").strip()
            ss_src   = (t.get("SahibSinghPunjabi") or "").strip()

            print(f"\n── Строфа ID={sid} ──")
            print(f"  Gurmukhi : {gurmukhi}")
            print(f"  Roman    : {roman}")
            print(f"  KG (EN)  : {kg_src if kg_src else '(нет)'}")
            print(f"  SS (PA)  : {ss_src}")
            print(f"  Кэш-ключи: kg_{sid}  /  ss_{sid}")

        # Комментарий
        if shabad_id:
            commentary = fetch_commentary_text(shabad_id)
            if commentary:
                print(f"\n── Комментарий KG (shabad_id={shabad_id}) ──")
                print(f"  Кэш-ключ: commentary_{shabad_id}")
                print(f"\n{commentary}\n")
            else:
                print(f"\n  (Комментарий отсутствует для shabad_id={shabad_id})")


# ── DOCX helpers ──────────────────────────────────────────────────────────────
def add_run(para, text: str, color=None, bold=False, italic=False, size_pt=None):
    run = para.add_run(text)
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def add_section_heading(doc: Document, ang: int, shabad_num: int):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(f"АНГ {ang}  ·  Стих {shabad_num}")
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_HEADING


def add_commentary_block(doc: Document, text: str):
    """Добавляет блок комментария KG после таблиц стихов."""
    # Разделитель
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(10)
    hr.paragraph_format.space_after  = Pt(4)
    run = hr.add_run("Комментарий KhojGurbani")
    run.font.bold  = True
    run.font.size  = Pt(10)
    run.font.color.rgb = COLOR_LABEL

    # Текст комментария
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(12)
    para.paragraph_format.left_indent  = Inches(0.2)
    run2 = para.add_run(text)
    run2.font.size = Pt(9)
    run2.font.color.rgb = COLOR_COMMENTARY


def build_verse_table(doc: Document, gurmukhi: str, roman: str,
                       kg_ru: str, ss_ru: str, has_kg: bool):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    widths = [Inches(2.1), Inches(2.1), Inches(3.1)]
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = widths[i]

    row = table.rows[0]

    p0 = row.cells[0].paragraphs[0]
    add_run(p0, gurmukhi, color=COLOR_GURBANI, size_pt=11)

    p1 = row.cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p1, roman, color=COLOR_ROMAN, italic=True, size_pt=10)

    c2 = row.cells[2]
    p2 = c2.paragraphs[0]

    if has_kg and kg_ru:
        add_run(p2, kg_ru, color=COLOR_KG, size_pt=10)
        add_run(p2, "\n— KhojGurbani", color=COLOR_LABEL, italic=True, size_pt=8)
        if ss_ru:
            p3 = c2.add_paragraph()
            p3.paragraph_format.space_before = Pt(5)
            add_run(p3, ss_ru, color=COLOR_SS, size_pt=10)
            add_run(p3, "\n— Prof. Sahib Singh", color=COLOR_LABEL, italic=True, size_pt=8)
    elif ss_ru:
        add_run(p2, ss_ru, color=COLOR_SS, size_pt=10)
        add_run(p2, "\n— Prof. Sahib Singh", color=COLOR_LABEL, italic=True, size_pt=8)
    else:
        add_run(p2, "[перевод отсутствует]", color=COLOR_LABEL, italic=True, size_pt=10)

    doc.add_paragraph()


# ── Команда: build ────────────────────────────────────────────────────────────
def cmd_build(args):
    cache = load_cache()
    doc   = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    title = doc.add_heading("Шри Гуру Грантх Сахиб — перевод на русский", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sub, "Источники: KhojGurbani · Prof. Sahib Singh",
            color=COLOR_LABEL, italic=True, size_pt=9)
    doc.add_paragraph()

    missing = []

    for shabad_num in range(args.shabad, (args.end_shabad or args.shabad) + 1):
        print(f"\nСборка АНГ {args.ang}, стих {shabad_num}...")
        shabad = fetch_shabad(args.ang, shabad_num)
        if not shabad:
            print("  [!] Нет данных, пропуск")
            continue

        shabad_id = shabad.get("id")
        add_section_heading(doc, args.ang, shabad_num)

        for verse in shabad.get("scriptures", []):
            sid      = verse["id"]
            gurmukhi = verse.get("Scripture", "").strip()
            roman    = verse.get("ScriptureRoman", "").strip()
            t        = verse.get("translation", {})
            has_kg   = bool((t.get("KhojgurbaaniEnglish") or "").strip())

            kg_ru = cache.get(f"kg_{sid}", "")
            ss_ru = cache.get(f"ss_{sid}", "")

            if not kg_ru and has_kg:
                missing.append(f"kg_{sid}")
            if not ss_ru:
                missing.append(f"ss_{sid}")

            build_verse_table(doc, gurmukhi, roman, kg_ru, ss_ru, has_kg)
            print(f"  ID={sid}  KG={'✓' if kg_ru else '✗'}  SS={'✓' if ss_ru else '✗'}")

        # Комментарий
        if shabad_id:
            comm_key = f"commentary_{shabad_id}"
            comm_ru  = cache.get(comm_key, "")
            if comm_ru:
                add_commentary_block(doc, comm_ru)
                print(f"  commentary_{shabad_id}: ✓")
            else:
                # Проверим есть ли вообще комментарий на сайте
                raw = fetch_commentary_text(shabad_id)
                if raw:
                    missing.append(comm_key)
                    print(f"  commentary_{shabad_id}: ✗ (есть на сайте, нет в кэше)")

    out = Path(args.output)
    doc.save(str(out))
    print(f"\n✓ Сохранено: {out}")

    if missing:
        print(f"\n⚠ Отсутствуют в кэше ({len(missing)}):")
        for k in missing:
            print(f"   {k}")


# ── Команда: dump ─────────────────────────────────────────────────────────────
def cmd_dump(args):
    """Скачивает сырые данные шабадов + комментариев и сохраняет в JSON-файл."""
    out_file = Path(args.output)
    dump = {}

    for shabad_num in range(args.shabad, (args.end_shabad or args.shabad) + 1):
        print(f"Загружаю шабад {args.ang}/{shabad_num}...", end=" ", flush=True)
        shabad = fetch_shabad(args.ang, shabad_num)
        if not shabad:
            print("✗ нет данных")
            continue

        shabad_id = shabad.get("id")
        entry = {
            "ang": args.ang,
            "shabad_num": shabad_num,
            "shabad_id": shabad_id,
            "scriptures": [],
        }

        for verse in shabad.get("scriptures", []):
            sid = verse["id"]
            t   = verse.get("translation", {})
            entry["scriptures"].append({
                "id":       sid,
                "gurmukhi": verse.get("Scripture", "").strip(),
                "roman":    verse.get("ScriptureRoman", "").strip(),
                "kg_en":    (t.get("KhojgurbaaniEnglish") or "").strip(),
                "ss_pa":    (t.get("SahibSinghPunjabi") or "").strip(),
            })

        commentary = ""
        if shabad_id:
            commentary = fetch_commentary_text(shabad_id)

        entry["commentary_en"] = commentary
        dump[f"{args.ang}_{shabad_num}"] = entry

        n_verses = len(entry["scriptures"])
        has_comm = "✓" if commentary else "✗"
        print(f"✓ {n_verses} строф, комментарий {has_comm}")

    out_file.write_text(json.dumps(dump, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n✓ Сохранено {len(dump)} шабадов → {out_file}")


# ── main ─────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser()
    sub    = parser.add_subparsers(dest="cmd", required=True)

    p_fetch = sub.add_parser("fetch")
    p_fetch.add_argument("--ang",        type=int, required=True)
    p_fetch.add_argument("--shabad",     type=int, required=True)
    p_fetch.add_argument("--end-shabad", type=int, default=None)

    p_build = sub.add_parser("build")
    p_build.add_argument("--ang",        type=int, required=True)
    p_build.add_argument("--shabad",     type=int, required=True)
    p_build.add_argument("--end-shabad", type=int, default=None)
    p_build.add_argument("--output",     default=str(OUTPUT_FILE))

    p_dump = sub.add_parser("dump")
    p_dump.add_argument("--ang",        type=int, required=True)
    p_dump.add_argument("--shabad",     type=int, required=True)
    p_dump.add_argument("--end-shabad", type=int, default=None)
    p_dump.add_argument("--output",     default=str(BASE_DIR / "kg_raw.json"))

    p_show = sub.add_parser("show")
    p_show.add_argument("--ang",        type=int, required=True)
    p_show.add_argument("--shabad",     type=int, required=True)
    p_show.add_argument("--end-shabad", type=int, default=None)
    p_show.add_argument("--dump",       default=str(BASE_DIR / "kg_raw_all.json"))

    args = parser.parse_args()
    if args.cmd == "fetch":
        cmd_fetch(args)
    elif args.cmd == "build":
        cmd_build(args)
    elif args.cmd == "dump":
        cmd_dump(args)
    elif args.cmd == "show":
        cmd_show(args)


if __name__ == "__main__":
    main()
