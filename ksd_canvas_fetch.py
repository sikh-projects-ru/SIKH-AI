#!/usr/bin/env python3
"""
Fetch Asia Samachar article and convert to docx, preserving order:
text paragraphs, headings, images, links.
"""

import re
import os
import html as html_module
import urllib.request
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Config ────────────────────────────────────────────────────────────────────
URL = "https://asiasamachar.com/2021/05/05/guru-nanaks-canvas/"
OUTPUT = "/home/royal/Work/Spiritual/ksd_reference_material/Nanak_Canvas_Part1.docx"
IMG_DIR = "/tmp/nanak_canvas_imgs"
os.makedirs(IMG_DIR, exist_ok=True)


# ── 1. Fetch HTML ─────────────────────────────────────────────────────────────
print("Fetching article...")
req = urllib.request.Request(URL, headers={'User-Agent': 'Mozilla/5.0'})
with urllib.request.urlopen(req) as r:
    raw_html = r.read().decode('utf-8', errors='replace')
print(f"  Downloaded {len(raw_html)} bytes")


# ── 2. Extract article block ──────────────────────────────────────────────────
art_start = raw_html.find('<div class="td-post-content">')
art_end_marker = raw_html.find('class="td-related-row"', art_start)
if art_end_marker == -1:
    art_end_marker = raw_html.find('id="comments"', art_start)
article_html = raw_html[art_start:art_end_marker]
print(f"  Article block: {len(article_html)} bytes")

# Extract page title
title_m = re.search(r'<h1[^>]*class="[^"]*entry-title[^"]*"[^>]*>(.*?)</h1>', raw_html, re.DOTALL)
if not title_m:
    title_m = re.search(r'<h1[^>]*>(.*?)</h1>', raw_html, re.DOTALL)
page_title = html_module.unescape(re.sub(r'<[^>]+>', '', title_m.group(1))).strip() if title_m else "Guru Nanak's Canvas"
print(f"  Title: {page_title}")


# ── 3. HTML → content nodes (in document order) ───────────────────────────────
def strip_tags(s):
    """Remove HTML tags and unescape entities."""
    s = re.sub(r'<[^>]+>', '', s)
    return html_module.unescape(s).strip()

def extract_links(s):
    """Return list of (text, href) tuples from <a> tags."""
    return re.findall(r'<a[^>]+href="([^"]*)"[^>]*>(.*?)</a>', s, re.DOTALL)

def best_img_src(tag_html):
    """Pick full-size src from img tag (avoid srcset thumbnails)."""
    # Use src= attribute only (not srcset)
    m = re.search(r'\bsrc="([^"]+)"', tag_html)
    if not m:
        return None
    src = m.group(1)
    # Skip dimension-suffixed thumbnails
    if re.search(r'-\d+x\d+\.(jpg|png)', src):
        return None
    return src

def emit_img_node(tag_html):
    """Return ('img', src, alt) node or None if not an article image."""
    src = best_img_src(tag_html)
    if not src:
        return None
    if 'uploads/2021' not in src:
        return None
    alt_m = re.search(r'alt="([^"]*)"', tag_html)
    alt = html_module.unescape(alt_m.group(1)) if alt_m else ''
    return ('img', src, alt)

def split_by_imgs(inner_html):
    """
    Split an inner HTML block (inside p/h5/etc.) into a list of sub-nodes:
    ('img', src, alt) or ('text', raw_html) — in document order.
    """
    parts = []
    pos = 0
    for m in re.finditer(r'<img\b[^>]*/?>',  inner_html, re.IGNORECASE | re.DOTALL):
        before = inner_html[pos:m.start()]
        if before.strip():
            parts.append(('text', before))
        node = emit_img_node(m.group(0))
        if node:
            parts.append(node)
        pos = m.end()
    tail = inner_html[pos:]
    if tail.strip():
        parts.append(('text', tail))
    return parts

# Tokenize: find all block-level tags, split each by embedded img tags
nodes = []

block_re = re.compile(
    r'<(h[1-6]|p|hr)\b([^>]*)>(.*?)</\1>',
    re.IGNORECASE | re.DOTALL
)

for m in block_re.finditer(article_html):
    tag   = m.group(1).lower()
    attrs = m.group(2)
    inner = m.group(3)

    if tag == 'hr':
        nodes.append(('hr',))
        continue

    # Split inner HTML at img tags to preserve image order
    sub = split_by_imgs(inner)

    for part in sub:
        if part[0] == 'img':
            nodes.append(part)
        else:  # text
            text = strip_tags(part[1])
            if not text or text in ('\xa0', '&nbsp;'):
                continue
            links = extract_links(part[1])
            indent = 'padding-left' in attrs

            if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                nodes.append(('heading', int(tag[1]), text, links))
            else:
                nodes.append(('p', text, links, indent))

print(f"  Parsed nodes: {len(nodes)} ({sum(1 for n in nodes if n[0]=='img')} images, "
      f"{sum(1 for n in nodes if n[0]=='p')} paragraphs, "
      f"{sum(1 for n in nodes if n[0]=='heading')} headings)")


# ── 4. Download images ────────────────────────────────────────────────────────
def download_img(url):
    fname = os.path.join(IMG_DIR, os.path.basename(url.split('?')[0]))
    if os.path.exists(fname):
        return fname
    try:
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=15) as r:
            data = r.read()
        with open(fname, 'wb') as f:
            f.write(data)
        print(f"  Downloaded image: {os.path.basename(fname)}")
        return fname
    except Exception as e:
        print(f"  WARN: could not download {url}: {e}")
        return None


# ── 5. Build DOCX ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default style
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

def add_text_para(doc, text, links=None, bold=False, italic=False,
                  size=11, color=None, center=False, indent=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.left_indent = Inches(0.5)

    # If there are links, annotate them inline
    if links:
        # Build annotated text: replace linked text with "text [url]"
        annotated = text
        for href, ltext in links:
            clean_ltext = strip_tags(ltext)
            if clean_ltext and href:
                # Mark with footnote style: text → text [→ url]
                annotated = annotated.replace(clean_ltext, f"{clean_ltext} [→ {href}]", 1)
        text = annotated

    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(page_title)
run.font.name = 'Arial'
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)

p2 = doc.add_paragraph()
run2 = p2.add_run("By Karminder Singh Dhillon | Asia Samachar | Part 1 of 12")
run2.font.name = 'Arial'
run2.font.size = Pt(10)
run2.font.italic = True
run2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

p3 = doc.add_paragraph()
run3 = p3.add_run(f"Source: {URL}")
run3.font.name = 'Arial'
run3.font.size = Pt(9)
run3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # spacer

# Heading level → style params
HEADING_STYLES = {
    1: dict(size=16, bold=True, color=(0x00, 0x00, 0x80)),
    2: dict(size=14, bold=True, color=(0x00, 0x00, 0x80)),
    3: dict(size=13, bold=True, color=(0x20, 0x20, 0x60)),
    4: dict(size=12, bold=True, color=(0x20, 0x20, 0x60)),
    5: dict(size=11, bold=True, color=(0x40, 0x00, 0x40)),
    6: dict(size=11, bold=False, color=(0x40, 0x40, 0x40)),
}

img_count = 0
for node in nodes:
    kind = node[0]

    if kind == 'img':
        _, src, alt = node
        local = download_img(src)
        if local:
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(local, width=Inches(5.5))
                img_count += 1
                if alt:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cap.add_run(alt)
                    cr.font.size = Pt(9)
                    cr.font.italic = True
                    cr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            except Exception as e:
                print(f"  WARN: could not insert image {src}: {e}")
                add_text_para(doc, f"[Image: {src}]", size=9, color=(0x60, 0x60, 0x60))

    elif kind == 'heading':
        _, level, text, links = node
        st = HEADING_STYLES.get(level, HEADING_STYLES[4])
        add_text_para(doc, text, links=links, center=(level <= 2), **st)

    elif kind == 'p':
        _, text, links, indent = node
        # Detect special paragraphs
        is_call_to_action = 'CLICK' in text and 'HERE' in text
        add_text_para(doc, text, links=links,
                      bold=is_call_to_action,
                      size=10 if is_call_to_action else 11,
                      color=(0x00, 0x60, 0x00) if is_call_to_action else None,
                      center=is_call_to_action,
                      indent=indent)

    elif kind == 'hr':
        doc.add_paragraph('─' * 40)

doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")
print(f"Images inserted: {img_count}")
print(f"Total nodes processed: {len(nodes)}")
