#!/usr/bin/env python3
# we_are_not_hindu_bot.py
#
# Uploads each PDF page image to ChatGPT, gets Russian translation,
# saves page_json/page_XXXX.json, then assembles a DOCX.

from __future__ import annotations

import argparse
import json
import re
import time
from pathlib import Path

import fitz  # pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from playwright.sync_api import Page, TimeoutError as PWTimeout, sync_playwright
from playwright_stealth import Stealth

_stealth = Stealth()

# ── Paths ─────────────────────────────────────────────────────────────────────

HERE = Path(__file__).parent
PDF_PATH = HERE / "Sikhs_We_Are_Not_Hindus_By_Kahan_Singh_Nabha.pdf"
PAGES_DIR = HERE / "pages"
JSON_DIR = HERE / "page_json"
RAW_LOG_DIR = HERE / "raw_logs"
OUTPUT_DOCX = HERE / "we_are_not_hindu_ru.docx"
PROGRESS_FILE = HERE / ".progress_we_are_not_hindu.txt"

# Reuse bot_profile from custom_khoj_sahib_singh (already logged in to ChatGPT)
BOT_PROFILE = HERE.parent / "custom_khoj_sahib_singh" / "bot_profile"

CHAT_URL = "https://chatgpt.com/"

# ── Colours for DOCX ──────────────────────────────────────────────────────────

COLOR_RUSSIAN   = RGBColor(0x00, 0x00, 0x00)
COLOR_PAGE      = RGBColor(0x99, 0x99, 0x99)
COLOR_HINDU     = RGBColor(0x8B, 0x00, 0x00)   # тёмно-красный — Хинду
COLOR_SIKH      = RGBColor(0x00, 0x40, 0x80)   # тёмно-синий  — Сикх
COLOR_FOOTNOTE  = RGBColor(0x77, 0x77, 0x77)   # серый — сноски
COLOR_MARKER    = RGBColor(0xAA, 0xAA, 0xAA)   # светло-серый — [→]
COLOR_FLAG      = RGBColor(0xCC, 0x66, 0x00)   # оранжевый — ⚠ метка качества

# ── Noise patterns to strip from translation ──────────────────────────────────

NOISE_PATTERNS = [
    r"sikhbookclub\.com",
    r"www\.sikhbookclub\.com",
    r"SIKHBOOKCLUB\.COM",
    r"WWW\.SIKHBOOKCLUB\.COM",
]
_NOISE_RE = re.compile("|".join(NOISE_PATTERNS), re.IGNORECASE)

# ── Quality-flag patterns ─────────────────────────────────────────────────────

# GPT meta-phrases that signal a non-translation response
_GPT_META_PATTERNS = re.compile(
    r"анализ изображени|i cannot|i'm unable|sorry|as an ai|"
    r"на изображении (видно|представлен|показан)|это изображение",
    re.IGNORECASE,
)

# Only page-marker / noise content (page entirely empty or watermark-only)
_PAGE_MARKER_RE = re.compile(
    r"^[\s\d\.\-—–]*"
    r"(страница\s*\d+|page\s*\d+|www\.|http)?"
    r"[\s\d\.\-—–]*$",
    re.IGNORECASE,
)


def quality_check(text: str, page_num: int) -> list[str]:
    """Return list of quality flag strings (empty = looks fine)."""
    flags: list[str] = []
    stripped = text.strip()

    if not stripped:
        return ["empty"]

    # Essentially only a page marker / watermark
    if _PAGE_MARKER_RE.match(stripped) or len(stripped) < 40:
        flags.append("short")

    # Noise still present
    if _NOISE_RE.search(stripped):
        flags.append("noise")

    # GPT refused or described the image instead of translating
    if _GPT_META_PATTERNS.search(stripped):
        flags.append("gpt_meta")

    # Too much Latin in what should be Russian text
    # (skip pages <11 — intro/prefatory pages legitimately have more Latin)
    if page_num >= 11:
        latin = sum(1 for c in stripped if "a" <= c.lower() <= "z")
        cyrillic = sum(1 for c in stripped if "\u0400" <= c <= "\u04ff")
        total = latin + cyrillic
        if total > 0 and latin / total > 0.35:
            flags.append("english_leak")

    return flags


def clean_translation(text: str) -> str:
    text = _NOISE_RE.sub("", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── Prompt ────────────────────────────────────────────────────────────────────

PROMPT_TEMPLATE = """\
Это страница {page} из книги «Ham Hindu Nahin» («Мы — не индуисты») Кан Сингх Набха. Книга на английском. Это диалогическая книга: в ней ведётся диалог Сикха и Хинду.

Переведи текст на этой странице на русский язык. Строго по содержанию, без добавлений.

Правила:
- Переводи только смысловой текст. Не включай в перевод адреса сайтов, издательские пометки, колонтитулы, номера страниц.
- Диалог: «Hindu:» → «Хинду:», «Sikh:» → «Сикх:».
- Цитаты из СГГС или других источников сохраняй как цитаты.
- Если текст начинается с середины предложения (продолжение предыдущей страницы) — начни перевод с [→].
- Если текст обрывается на середине предложения — заверши перевод на [→].
- Переводи понятно, как книжный текст, не дословно.

Верни ответ строго между BEGIN_JSON и END_JSON, без markdown, без пояснений:

BEGIN_JSON
{{
  "page": {page},
  "translation_ru": "...русский перевод..."
}}
END_JSON
"""

REPAIR_PROMPT = """\
Перепиши свой последний ответ в правильном JSON-формате.

Нужен только валидный JSON между BEGIN_JSON и END_JSON.
Без markdown, без code fences, без пояснений.

Обязательные поля: page (integer), translation_ru (string).

BEGIN_JSON
{
  "page": ...,
  "translation_ru": "..."
}
END_JSON
"""

# ── PDF → images ──────────────────────────────────────────────────────────────

def pdf_to_images(pdf_path: Path, pages_dir: Path, dpi: int = 200) -> list[Path]:
    pages_dir.mkdir(exist_ok=True)
    doc = fitz.open(str(pdf_path))
    total = len(doc)
    paths = []
    for i in range(total):
        out = pages_dir / f"page_{i + 1:04d}.png"
        if not out.exists():
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = doc[i].get_pixmap(matrix=mat)
            pix.save(str(out))
            print(f"  Rendered {i + 1}/{total} → {out.name}")
        paths.append(out)
    doc.close()
    return paths


# ── Progress ──────────────────────────────────────────────────────────────────

def load_progress() -> set[int]:
    if not PROGRESS_FILE.exists():
        return set()
    return {int(l) for l in PROGRESS_FILE.read_text().splitlines() if l.strip().isdigit()}


def save_progress(page_num: int) -> None:
    with open(PROGRESS_FILE, "a") as f:
        f.write(f"{page_num}\n")


# ── JSON helpers ──────────────────────────────────────────────────────────────

def extract_json(text: str) -> dict | None:
    """Try several extraction strategies, most strict first."""
    # 1. Explicit markers
    m = re.search(r"BEGIN_JSON\s*(\{.*?\})\s*END_JSON", text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            pass

    # 2. Markdown code block: ```json { ... } ```
    m = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            pass

    # 3. Bare JSON object anywhere in the text
    m = re.search(r"(\{[^{}]*\"translation_ru\".*?\})", text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            pass

    return None


def validate(data: dict, page_num: int) -> bool:
    return (
        isinstance(data, dict)
        and isinstance(data.get("translation_ru"), str)
        and len(data["translation_ru"].strip()) > 10
    )


def save_json(data: dict) -> Path:
    JSON_DIR.mkdir(exist_ok=True)
    out = JSON_DIR / f"page_{data['page']:04d}.json"
    out.write_text(json.dumps(data, ensure_ascii=False, indent=2))
    return out


# ── ChatGPT via Playwright ────────────────────────────────────────────────────

def wait_for_response(page: Page, timeout_ms: int = 150_000) -> str:
    """Poll until ChatGPT response stabilizes (text stops changing)."""
    deadline = time.time() + timeout_ms / 1000

    def last_text() -> str:
        msgs = page.query_selector_all('[data-message-author-role="assistant"]')
        if not msgs:
            return ""
        return msgs[-1].inner_text()

    # Phase 1: wait until any response appears (up to 30s)
    print("  Waiting for response…", end="", flush=True)
    for _ in range(60):
        if time.time() > deadline:
            break
        txt = last_text()
        if len(txt) > 10:
            break
        time.sleep(0.5)
    print()

    # Phase 2: poll until text stops changing for 3 consecutive checks
    stable_count = 0
    prev = ""
    while time.time() < deadline:
        txt = last_text()
        if txt == prev and len(txt) > 10:
            stable_count += 1
            if stable_count >= 3:
                break
        else:
            stable_count = 0
        prev = txt
        time.sleep(1.0)

    result = last_text()
    print(f"  Response length: {len(result)} chars")
    return result


def send_page(page: Page, image_path: Path, prompt: str) -> None:
    """Attach image and send prompt as ONE message."""

    # 1. Find the hidden file input and set the image
    file_input = page.query_selector('input[type="file"]')
    if file_input is None:
        try:
            page.click('button[aria-label*="ttach"]', timeout=5_000)
            time.sleep(0.8)
        except Exception:
            pass
        file_input = page.query_selector('input[type="file"]')

    if file_input:
        file_input.set_input_files(str(image_path))
        try:
            page.wait_for_selector('img[alt*="Uploaded"]', timeout=20_000)
        except PWTimeout:
            time.sleep(3)
    else:
        print("  [warn] Could not find file input — sending without image")

    # 2. Type prompt into textarea
    textarea = page.query_selector('div#prompt-textarea')
    if textarea is None:
        textarea = page.query_selector('textarea')
    textarea.click()
    time.sleep(0.3)
    textarea.fill(prompt)
    time.sleep(0.3)

    # 3. Send
    page.keyboard.press("Enter")


def send_text(page: Page, prompt: str) -> None:
    """Send a text-only message (repair prompt)."""
    textarea = page.query_selector('div#prompt-textarea')
    if textarea is None:
        textarea = page.query_selector('textarea')
    textarea.click()
    time.sleep(0.2)
    textarea.fill(prompt)
    time.sleep(0.2)
    page.keyboard.press("Enter")


def open_new_chat(page: Page) -> None:
    page.goto(CHAT_URL, wait_until="domcontentloaded", timeout=30_000)
    time.sleep(2.5)


# ── Translation loop ──────────────────────────────────────────────────────────

def translate_pages(image_paths: list[Path], start: int, end: int,
                    max_retries: int, delay_s: float, force: bool) -> None:
    JSON_DIR.mkdir(exist_ok=True)
    RAW_LOG_DIR.mkdir(exist_ok=True)
    done = load_progress()

    subset = [p for p in image_paths if start <= _page_num(p) <= end]

    with sync_playwright() as pw:
        ctx = pw.chromium.launch_persistent_context(
            str(BOT_PROFILE),
            headless=False,
            args=["--start-maximized"],
            no_viewport=True,
        )
        page = ctx.pages[0] if ctx.pages else ctx.new_page()
        _stealth.apply_stealth_sync(page)

        open_new_chat(page)
        input("\n>> Press ENTER when ChatGPT is ready (logged in, new chat open)…\n")

        for img_path in subset:
            pnum = _page_num(img_path)

            out_json = JSON_DIR / f"page_{pnum:04d}.json"
            if out_json.exists() and not force:
                print(f"[skip] page {pnum} (JSON exists)")
                save_progress(pnum)
                done.add(pnum)
                continue

            print(f"\n── Page {pnum} ──────────────────────────────")
            prompt = PROMPT_TEMPLATE.format(page=pnum)
            result = None

            for attempt in range(1, max_retries + 1):
                try:
                    open_new_chat(page)
                    time.sleep(1)

                    print(f"  Sending (attempt {attempt})…")
                    send_page(page, img_path, prompt)
                    raw = wait_for_response(page)

                    (RAW_LOG_DIR / f"page_{pnum:04d}_a{attempt}.txt").write_text(raw)

                    data = extract_json(raw)

                    if data is None:
                        print("  [warn] No JSON markers — repair prompt…")
                        send_text(page, REPAIR_PROMPT)
                        raw2 = wait_for_response(page, timeout_ms=60_000)
                        (RAW_LOG_DIR / f"page_{pnum:04d}_a{attempt}_repair.txt").write_text(raw2)
                        data = extract_json(raw2)

                    if data is not None and validate(data, pnum):
                        data["page"] = pnum
                        data["translation_ru"] = clean_translation(data["translation_ru"])
                        flags = quality_check(data["translation_ru"], pnum)
                        if flags:
                            data["quality_flags"] = flags
                            print(f"  [quality] flags: {flags}")
                        result = data
                        break

                    print("  [warn] Invalid or empty — retrying…")

                except Exception as e:
                    print(f"  [error] attempt {attempt}: {e}")

                if attempt < max_retries:
                    time.sleep(delay_s)

            if result:
                path = save_json(result)
                save_progress(pnum)
                done.add(pnum)
                preview = result["translation_ru"][:80].replace("\n", " ")
                print(f"  ✓ {path.name}  |  «{preview}…»")
            else:
                print(f"  ✗ FAILED page {pnum} — skipping")

            time.sleep(delay_s)

        ctx.close()

    print(f"\nDone. Total translated: {len(done)}")


def _page_num(p: Path) -> int:
    return int(p.stem.split("_")[1])


# ── Audit ─────────────────────────────────────────────────────────────────────

def audit_json(start: int, end: int) -> None:
    """Print a report of pages with quality flags."""
    pages = sorted(JSON_DIR.glob("page_*.json"), key=lambda p: _page_num(p))
    pages = [p for p in pages if start <= _page_num(p) <= end]

    flagged = []
    for json_path in pages:
        data = json.loads(json_path.read_text())
        pnum = data["page"]
        text = data.get("translation_ru", "")

        # Re-run quality check (catches pages translated before flags were added)
        existing = data.get("quality_flags", [])
        fresh = quality_check(text, pnum)
        flags = sorted(set(existing) | set(fresh))

        if flags:
            flagged.append((pnum, flags, text[:60].replace("\n", " ")))

    if not flagged:
        print("Все страницы выглядят нормально.")
        return

    print(f"\n{'Стр.':>5}  {'Флаги':<30}  Превью")
    print("─" * 80)
    for pnum, flags, preview in flagged:
        print(f"{pnum:>5}  {', '.join(flags):<30}  {preview!r}")
    print(f"\nВсего с флагами: {len(flagged)} из {len(pages)}")


# ── DOCX helpers ──────────────────────────────────────────────────────────────

# Matches "Хинду:" or "Сикх:" at the start of a line (with optional space after)
_DIALOGUE_RE = re.compile(r"^(Хинду:|Сикх:)\s*")

# Matches citation references like (Бхай Гурдас, Вар 1) or (Шри Раг М 1, с. 15)
_CITATION_RE = re.compile(r"^\(.*\)\s*$")

# Footnote lines starting with a digit/superscript + dot or close-paren
_FOOTNOTE_RE = re.compile(r"^[\d¹²³⁴⁵⁶⁷⁸⁹]+[\.\)]\s+\S")

# Boundary marker [→] anywhere in text
_MARKER_RE = re.compile(r"\[→\]")


def _add_run(para, text: str, size: int = 11, bold: bool = False,
             italic: bool = False, color: RGBColor | None = None) -> None:
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_translation_block(doc: Document, translation: str) -> None:
    """
    Add translation to the document with rich formatting:
      - Хинду: / Сикх: labels in bold + colour
      - Standalone citation lines in italic
      - Footnote lines in small grey
      - [→] markers in light grey
      - Everything else in plain black
    """
    lines = translation.split("\n")

    for line in lines:
        raw = line.rstrip()

        # Blank line → spacer paragraph
        if not raw.strip():
            doc.add_paragraph()
            continue

        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)

        # ── Footnote line ──────────────────────────────────────────────────
        if _FOOTNOTE_RE.match(raw):
            _add_run(para, raw, size=9, color=COLOR_FOOTNOTE)
            continue

        # ── Citation-only line ─────────────────────────────────────────────
        if _CITATION_RE.match(raw.strip()):
            _add_run(para, raw, size=10, italic=True, color=COLOR_FOOTNOTE)
            para.paragraph_format.left_indent = Pt(24)
            continue

        # ── Build runs for a normal line, handling [→] and dialogue labels ──

        # Strip and handle leading [→]
        remainder = raw
        if remainder.startswith("[→]"):
            _add_run(para, "[→] ", size=10, color=COLOR_MARKER)
            remainder = remainder[4:].lstrip()

        # Check for dialogue label at start
        m = _DIALOGUE_RE.match(remainder)
        if m:
            label = m.group(1)
            color = COLOR_HINDU if label == "Хинду:" else COLOR_SIKH
            _add_run(para, label + " ", bold=True, color=color)
            remainder = remainder[m.end():]

        # Now emit the rest, splitting on any embedded [→]
        parts = _MARKER_RE.split(remainder)
        for i, part in enumerate(parts):
            if part:
                _add_run(para, part, color=COLOR_RUSSIAN)
            if i < len(parts) - 1:
                _add_run(para, " [→] ", size=10, color=COLOR_MARKER)


# ── DOCX builder ──────────────────────────────────────────────────────────────

def build_docx(start: int, end: int, output: Path) -> None:
    doc = Document()

    title = doc.add_heading("Мы — не индуисты", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Кан Сингх Набха (Kahn Singh Nabha)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = COLOR_PAGE
    sub.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    pages = sorted(JSON_DIR.glob("page_*.json"), key=lambda p: _page_num(p))
    pages = [p for p in pages if start <= _page_num(p) <= end]

    if not pages:
        print("No JSON files found — nothing to build.")
        return

    for json_path in pages:
        data = json.loads(json_path.read_text())
        pnum = data["page"]
        translation = data.get("translation_ru", "").strip()
        flags = data.get("quality_flags", [])

        # Page label
        label = doc.add_paragraph(f"— стр. {pnum} —")
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label.runs[0].font.color.rgb = COLOR_PAGE
        label.runs[0].font.size = Pt(9)

        # Quality warning (if any)
        if flags:
            warn_para = doc.add_paragraph()
            _add_run(warn_para, f"⚠ {', '.join(flags)}", size=9, color=COLOR_FLAG)

        if translation:
            add_translation_block(doc, translation)

        doc.add_paragraph()

    # Backup existing file before overwriting
    if output.exists():
        backup = output.with_suffix(".bak.docx")
        output.rename(backup)
        print(f"  Backup: {backup.name}")

    doc.save(str(output))
    print(f"DOCX saved: {output}  ({len(pages)} pages)")


# ── CLI ───────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description="we_are_not_hindu PDF → Russian DOCX via ChatGPT")
    parser.add_argument("--start", type=int, default=1)
    parser.add_argument("--end", type=int, default=9999)
    parser.add_argument("--dpi", type=int, default=200)
    parser.add_argument("--delay", type=float, default=4.0)
    parser.add_argument("--max-retries", type=int, default=3)
    parser.add_argument("--output", type=Path, default=OUTPUT_DOCX)
    parser.add_argument("--docx-start", type=int, default=1,
                        help="First page to include in DOCX (default: 1)")
    parser.add_argument("--docx-end", type=int, default=9999,
                        help="Last page to include in DOCX (default: all)")
    parser.add_argument("--rebuild-docx", action="store_true",
                        help="Rebuild DOCX from existing JSON, skip ChatGPT")
    parser.add_argument("--render-only", action="store_true",
                        help="Only render PDF to images")
    parser.add_argument("--force", action="store_true",
                        help="Re-translate already-done pages")
    parser.add_argument("--audit", action="store_true",
                        help="Print quality report for existing JSON files")
    args = parser.parse_args()

    print(f"Rendering PDF: {PDF_PATH.name}")
    image_paths = pdf_to_images(PDF_PATH, PAGES_DIR, dpi=args.dpi)
    total = len(image_paths)
    end = min(args.end, total)
    print(f"Total pages: {total}  |  Processing: {args.start}–{end}")

    if args.render_only:
        return

    if args.audit:
        audit_json(args.docx_start, args.docx_end)
        return

    if args.rebuild_docx:
        build_docx(args.docx_start, args.docx_end, args.output)
        return

    translate_pages(
        image_paths=image_paths,
        start=args.start,
        end=end,
        max_retries=args.max_retries,
        delay_s=args.delay,
        force=args.force,
    )
    # Always build DOCX from ALL available pages (not just the translated range)
    build_docx(args.docx_start, args.docx_end, args.output)


if __name__ == "__main__":
    main()
