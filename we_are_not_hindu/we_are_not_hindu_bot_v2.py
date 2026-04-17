#!/usr/bin/env python3
# we_are_not_hindu_bot_v2.py  — книжное оформление DOCX
#
# Идентичен v1 по логике перевода.
# Отличия только в сборке DOCX:
#   - страничные метки: маленькая правосторонняя сноска вместо центрального разделителя
#   - «цитаты» → курсив
#   - заголовки (ALL CAPS / короткие отдельные строки) → жирный
#   - Хальса → Кхальса
#   - inline (p. X) ссылки → серые

from __future__ import annotations

import argparse
import json
import re
import time
from pathlib import Path

import fitz  # pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from playwright.sync_api import Page, TimeoutError as PWTimeout, sync_playwright
from playwright_stealth import Stealth

_stealth = Stealth()

# ── Paths ─────────────────────────────────────────────────────────────────────

HERE = Path(__file__).parent
PDF_PATH = HERE / "Sikhs_We_Are_Not_Hindus_By_Kahan_Singh_Nabha.pdf"
PAGES_DIR = HERE / "pages"
JSON_DIR = HERE / "page_json"
RAW_LOG_DIR = HERE / "raw_logs"
OUTPUT_DOCX = HERE / "we_are_not_hindu_ru_v2.docx"
PROGRESS_FILE = HERE / ".progress_we_are_not_hindu.txt"

BOT_PROFILE = HERE.parent / "custom_khoj_sahib_singh" / "bot_profile"

CHAT_URL = "https://chatgpt.com/"

# ── Colours for DOCX ──────────────────────────────────────────────────────────

COLOR_RUSSIAN   = RGBColor(0x00, 0x00, 0x00)
COLOR_PAGE      = RGBColor(0xBB, 0xBB, 0xBB)   # очень светлый — метки страниц
COLOR_HINDU     = RGBColor(0x8B, 0x00, 0x00)   # тёмно-красный — Хинду
COLOR_SIKH      = RGBColor(0x00, 0x40, 0x80)   # тёмно-синий  — Сикх
COLOR_FOOTNOTE  = RGBColor(0x77, 0x77, 0x77)   # серый — сноски
COLOR_MARKER    = RGBColor(0xAA, 0xAA, 0xAA)   # светло-серый — [→]
COLOR_FLAG      = RGBColor(0xCC, 0x66, 0x00)   # оранжевый — ⚠ метка качества

# ── Noise patterns to strip from translation ──────────────────────────────────

NOISE_PATTERNS = [
    r"sikhbookclub\.com",
    r"www\.sikhbookclub\.com",
    r"SIKHBOOKCLUB\.COM",
    r"WWW\.SIKHBOOKCLUB\.COM",
]
_NOISE_RE = re.compile("|".join(NOISE_PATTERNS), re.IGNORECASE)

# ── Quality-flag patterns ─────────────────────────────────────────────────────

_GPT_META_PATTERNS = re.compile(
    r"анализ изображени|i cannot|i'm unable|sorry|as an ai|"
    r"на изображении (видно|представлен|показан)|это изображение",
    re.IGNORECASE,
)

_PAGE_MARKER_RE = re.compile(
    r"^[\s\d\.\-—–]*"
    r"(страница\s*\d+|page\s*\d+|www\.|http)?"
    r"[\s\d\.\-—–]*$",
    re.IGNORECASE,
)


def quality_check(text: str, page_num: int) -> list[str]:
    """Return list of quality flag strings (empty = looks fine)."""
    flags: list[str] = []
    stripped = text.strip()

    if not stripped:
        return ["empty"]

    if _PAGE_MARKER_RE.match(stripped) or len(stripped) < 40:
        flags.append("short")

    if _NOISE_RE.search(stripped):
        flags.append("noise")

    if _GPT_META_PATTERNS.search(stripped):
        flags.append("gpt_meta")

    if page_num >= 11:
        latin = sum(1 for c in stripped if "a" <= c.lower() <= "z")
        cyrillic = sum(1 for c in stripped if "\u0400" <= c <= "\u04ff")
        total = latin + cyrillic
        if total > 0 and latin / total > 0.35:
            flags.append("english_leak")

    return flags


def clean_translation(text: str) -> str:
    text = _NOISE_RE.sub("", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── Prompt ────────────────────────────────────────────────────────────────────

PROMPT_TEMPLATE = """\
Это страница {page} из книги «Ham Hindu Nahin» («Мы — не индуисты») Кан Сингх Набха. Книга на английском. Это диалогическая книга: в ней ведётся диалог Сикха и Хинду.

Переведи текст на этой странице на русский язык. Строго по содержанию, без добавлений.

Правила:
- Переводи только смысловой текст. Не включай в перевод адреса сайтов, издательские пометки, колонтитулы, номера страниц.
- Диалог: «Hindu:» → «Хинду:», «Sikh:» → «Сикх:».
- Цитаты из СГГС или других источников сохраняй как цитаты.
- Если текст начинается с середины предложения (продолжение предыдущей страницы) — начни перевод с [→].
- Если текст обрывается на середине предложения — заверши перевод на [→].
- Переводи понятно, как книжный текст, не дословно.
- Если страница пустая (нет текста, только фон, номер страницы или водяной знак) — верни "empty_page": true и "translation_ru": "".

Верни ответ строго между BEGIN_JSON и END_JSON, без markdown, без пояснений:

BEGIN_JSON
{{
  "page": {page},
  "empty_page": false,
  "translation_ru": "...русский перевод..."
}}
END_JSON
"""

REPAIR_PROMPT = """\
Перепиши свой последний ответ в правильном JSON-формате.

Нужен только валидный JSON между BEGIN_JSON и END_JSON.
Без markdown, без code fences, без пояснений.

Обязательные поля: page (integer), translation_ru (string).

BEGIN_JSON
{
  "page": ...,
  "translation_ru": "..."
}
END_JSON
"""

# ── PDF → images ──────────────────────────────────────────────────────────────

def pdf_to_images(pdf_path: Path, pages_dir: Path, dpi: int = 200) -> list[Path]:
    pages_dir.mkdir(exist_ok=True)
    doc = fitz.open(str(pdf_path))
    total = len(doc)
    paths = []
    for i in range(total):
        out = pages_dir / f"page_{i + 1:04d}.png"
        if not out.exists():
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = doc[i].get_pixmap(matrix=mat)
            pix.save(str(out))
            print(f"  Rendered {i + 1}/{total} → {out.name}")
        paths.append(out)
    doc.close()
    return paths


# ── Progress ──────────────────────────────────────────────────────────────────

def load_progress() -> set[int]:
    if not PROGRESS_FILE.exists():
        return set()
    return {int(l) for l in PROGRESS_FILE.read_text().splitlines() if l.strip().isdigit()}


def save_progress(page_num: int) -> None:
    with open(PROGRESS_FILE, "a") as f:
        f.write(f"{page_num}\n")


# ── JSON helpers ──────────────────────────────────────────────────────────────

def extract_json(text: str) -> dict | None:
    m = re.search(r"BEGIN_JSON\s*(\{.*?\})\s*END_JSON", text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            pass

    m = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            pass

    m = re.search(r"(\{[^{}]*\"translation_ru\".*?\})", text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            pass

    return None


def validate(data: dict, page_num: int) -> bool:
    if not isinstance(data, dict):
        return False
    if data.get("empty_page"):
        return True
    return (
        isinstance(data.get("translation_ru"), str)
        and len(data["translation_ru"].strip()) > 10
    )


def save_json(data: dict) -> Path:
    JSON_DIR.mkdir(exist_ok=True)
    out = JSON_DIR / f"page_{data['page']:04d}.json"
    out.write_text(json.dumps(data, ensure_ascii=False, indent=2))
    return out


# ── ChatGPT via Playwright ────────────────────────────────────────────────────

def wait_for_response(page: Page, timeout_ms: int = 150_000) -> str:
    deadline = time.time() + timeout_ms / 1000

    def last_text() -> str:
        msgs = page.query_selector_all('[data-message-author-role="assistant"]')
        if not msgs:
            return ""
        return msgs[-1].inner_text()

    print("  Waiting for response…", end="", flush=True)
    for _ in range(60):
        if time.time() > deadline:
            break
        txt = last_text()
        if len(txt) > 10:
            break
        time.sleep(0.5)
    print()

    stable_count = 0
    prev = ""
    while time.time() < deadline:
        txt = last_text()
        if txt == prev and len(txt) > 10:
            stable_count += 1
            if stable_count >= 3:
                break
        else:
            stable_count = 0
        prev = txt
        time.sleep(1.0)

    result = last_text()
    print(f"  Response length: {len(result)} chars")
    return result


def send_page(page: Page, image_path: Path, prompt: str) -> None:
    file_input = page.query_selector('input[type="file"]')
    if file_input is None:
        try:
            page.click('button[aria-label*="ttach"]', timeout=5_000)
            time.sleep(0.8)
        except Exception:
            pass
        file_input = page.query_selector('input[type="file"]')

    if file_input:
        file_input.set_input_files(str(image_path))
        try:
            page.wait_for_selector('img[alt*="Uploaded"]', timeout=20_000)
        except PWTimeout:
            time.sleep(3)
    else:
        print("  [warn] Could not find file input — sending without image")

    textarea = page.query_selector('div#prompt-textarea')
    if textarea is None:
        textarea = page.query_selector('textarea')
    textarea.click()
    time.sleep(0.3)
    textarea.fill(prompt)
    time.sleep(0.3)

    page.keyboard.press("Enter")


def send_text(page: Page, prompt: str) -> None:
    textarea = page.query_selector('div#prompt-textarea')
    if textarea is None:
        textarea = page.query_selector('textarea')
    textarea.click()
    time.sleep(0.2)
    textarea.fill(prompt)
    time.sleep(0.2)
    page.keyboard.press("Enter")


def open_new_chat(page: Page) -> None:
    page.goto(CHAT_URL, wait_until="domcontentloaded", timeout=30_000)
    time.sleep(2.5)


# ── Translation loop ──────────────────────────────────────────────────────────

def translate_pages(image_paths: list[Path], start: int, end: int,
                    max_retries: int, delay_s: float, force: bool,
                    explicit_pages: list[Path] | None = None) -> None:
    JSON_DIR.mkdir(exist_ok=True)
    RAW_LOG_DIR.mkdir(exist_ok=True)
    done = load_progress()

    if explicit_pages is not None:
        subset = sorted(explicit_pages, key=_page_num)
    else:
        subset = [p for p in image_paths if start <= _page_num(p) <= end]

    with sync_playwright() as pw:
        ctx = pw.chromium.launch_persistent_context(
            str(BOT_PROFILE),
            headless=False,
            args=["--start-maximized"],
            no_viewport=True,
        )
        page = ctx.pages[0] if ctx.pages else ctx.new_page()
        _stealth.apply_stealth_sync(page)

        open_new_chat(page)
        input("\n>> Press ENTER when ChatGPT is ready (logged in, new chat open)…\n")

        for img_path in subset:
            pnum = _page_num(img_path)

            out_json = JSON_DIR / f"page_{pnum:04d}.json"
            if out_json.exists() and not force:
                print(f"[skip] page {pnum} (JSON exists)")
                save_progress(pnum)
                done.add(pnum)
                continue

            print(f"\n── Page {pnum} ──────────────────────────────")
            prompt = PROMPT_TEMPLATE.format(page=pnum)
            result = None

            for attempt in range(1, max_retries + 1):
                try:
                    open_new_chat(page)
                    time.sleep(1)

                    print(f"  Sending (attempt {attempt})…")
                    send_page(page, img_path, prompt)
                    raw = wait_for_response(page)

                    (RAW_LOG_DIR / f"page_{pnum:04d}_a{attempt}.txt").write_text(raw)

                    data = extract_json(raw)

                    if data is None:
                        print("  [warn] No JSON markers — repair prompt…")
                        send_text(page, REPAIR_PROMPT)
                        raw2 = wait_for_response(page, timeout_ms=60_000)
                        (RAW_LOG_DIR / f"page_{pnum:04d}_a{attempt}_repair.txt").write_text(raw2)
                        data = extract_json(raw2)

                    if data is not None and validate(data, pnum):
                        data["page"] = pnum
                        if data.get("empty_page"):
                            data["translation_ru"] = ""
                            print(f"  [empty page] стр. {pnum} — пустая страница")
                        else:
                            data["translation_ru"] = clean_translation(data["translation_ru"])
                            flags = quality_check(data["translation_ru"], pnum)
                            if flags:
                                data["quality_flags"] = flags
                                print(f"  [quality] flags: {flags}")
                        result = data
                        break

                    print("  [warn] Invalid or empty — retrying…")

                except Exception as e:
                    print(f"  [error] attempt {attempt}: {e}")

                if attempt < max_retries:
                    time.sleep(delay_s)

            if result:
                path = save_json(result)
                save_progress(pnum)
                done.add(pnum)
                preview = result["translation_ru"][:80].replace("\n", " ")
                print(f"  ✓ {path.name}  |  «{preview}…»")
            else:
                print(f"  ✗ FAILED page {pnum} — skipping")

            time.sleep(delay_s)

        ctx.close()

    print(f"\nDone. Total translated: {len(done)}")


def _page_num(p: Path) -> int:
    return int(p.stem.split("_")[1])


# ── Audit ─────────────────────────────────────────────────────────────────────

def audit_json(start: int, end: int) -> None:
    pages = sorted(JSON_DIR.glob("page_*.json"), key=lambda p: _page_num(p))
    pages = [p for p in pages if start <= _page_num(p) <= end]

    flagged = []
    for json_path in pages:
        data = json.loads(json_path.read_text())
        pnum = data["page"]
        text = data.get("translation_ru", "")

        existing = data.get("quality_flags", [])
        fresh = quality_check(text, pnum)
        flags = sorted(set(existing) | set(fresh))

        if flags:
            flagged.append((pnum, flags, text[:60].replace("\n", " ")))

    if not flagged:
        print("Все страницы выглядят нормально.")
        return

    print(f"\n{'Стр.':>5}  {'Флаги':<30}  Превью")
    print("─" * 80)
    for pnum, flags, preview in flagged:
        print(f"{pnum:>5}  {', '.join(flags):<30}  {preview!r}")
    print(f"\nВсего с флагами: {len(flagged)} из {len(pages)}")


# ── Missing-pages helpers ─────────────────────────────────────────────────────

def find_missing_pages(image_paths: list[Path], start: int, end: int) -> list[Path]:
    """Return page images in [start, end] without a JSON file, sorted by page number."""
    return sorted(
        [p for p in image_paths
         if start <= _page_num(p) <= end
         and not (JSON_DIR / f"page_{_page_num(p):04d}.json").exists()],
        key=_page_num,
    )


def scan_missing(image_paths: list[Path], start: int, end: int) -> None:
    missing = find_missing_pages(image_paths, start, end)
    total_in_range = sum(1 for p in image_paths if start <= _page_num(p) <= end)
    if not missing:
        print(f"Нет пропущенных страниц в диапазоне {start}–{end}.")
        return
    nums = [_page_num(p) for p in missing]
    print(f"\nПропущено {len(nums)} из {total_in_range} страниц (диапазон {start}–{end}):")
    # Выводим диапазонами для компактности
    ranges: list[tuple[int, int]] = []
    rs = prev = nums[0]
    for n in nums[1:]:
        if n == prev + 1:
            prev = n
        else:
            ranges.append((rs, prev))
            rs = prev = n
    ranges.append((rs, prev))
    for a, b in ranges:
        if a == b:
            print(f"  стр. {a}")
        else:
            print(f"  стр. {a}–{b}  ({b - a + 1} шт.)")


# ── DOCX helpers (v2 — книжное оформление) ────────────────────────────────────

_DIALOGUE_RE = re.compile(r"^(Хинду:|Сикх:)\s*")
_CITATION_RE = re.compile(r"^\(.*\)\s*$")
_FOOTNOTE_RE = re.compile(r"^[\d¹²³⁴⁵⁶⁷⁸⁹]+[\.\)]\s+\S")
_MARKER_RE   = re.compile(r"\[→\]")

# «текст в кавычках» — курсив
_QUOTE_RE = re.compile(r'(«[^»]*»)')

# inline (p. X) или (с. X) в конце строки — серый
_INLINE_PAGEREF_RE = re.compile(r'\s*\((?:p|с)\.\s*\d+\)\s*$')

# заголовки: строка целиком из заглавных (рус/лат) букв + пробелы/цифры/знаки
_ALLCAPS_RE = re.compile(r'^[А-ЯЁA-Z\d\s\.\-–—,:;!?/()«»"\']+$')

# замены терминов (без GPT)
_TERM_FIXES = [
    (re.compile(r'Хальса', re.IGNORECASE), 'Кхальса'),
    (re.compile(r'ХАЛЬСА', re.IGNORECASE), 'КХАЛЬСА'),
]


def _apply_term_fixes(text: str) -> str:
    for pattern, replacement in _TERM_FIXES:
        text = pattern.sub(replacement, text)
    return text


def _is_header(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 80:
        return False
    if _FOOTNOTE_RE.match(stripped):
        return False
    if _DIALOGUE_RE.match(stripped):
        return False
    alpha_chars = [c for c in stripped if c.isalpha()]
    if not alpha_chars:
        return False
    # Строка целиком заглавными буквами
    return all(c == c.upper() for c in alpha_chars)


def _add_run(para, text: str, size: int = 11, bold: bool = False,
             italic: bool = False, color: RGBColor | None = None) -> None:
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _emit_with_quotes(para, text: str, size: int = 11, bold: bool = False,
                      base_color: RGBColor | None = None) -> None:
    """Emit text, rendering «...» fragments as italic."""
    segments = _QUOTE_RE.split(text)
    for seg in segments:
        if not seg:
            continue
        is_quote = seg.startswith('«') and seg.endswith('»')
        _add_run(para, seg, size=size, bold=bold, italic=is_quote, color=base_color)


def add_translation_block(doc: Document, translation: str) -> None:
    translation = _apply_term_fixes(translation)
    lines = translation.split("\n")

    for line in lines:
        raw = line.rstrip()

        if not raw.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            continue

        # ── Footnote ──────────────────────────────────────────────────────
        if _FOOTNOTE_RE.match(raw):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.left_indent = Pt(12)
            _add_run(para, raw, size=9, color=COLOR_FOOTNOTE)
            continue

        # ── Citation-only line  (Бхай Гурдас, Вар 1) ──────────────────────
        if _CITATION_RE.match(raw.strip()):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.left_indent = Pt(24)
            _add_run(para, raw.strip(), size=10, italic=True, color=COLOR_FOOTNOTE)
            continue

        # ── Header (ALL CAPS) ──────────────────────────────────────────────
        if _is_header(raw):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(6)
            _add_run(para, raw.strip(), size=12, bold=True, color=COLOR_RUSSIAN)
            continue

        # ── Normal / dialogue line ─────────────────────────────────────────
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(5)

        # Split off trailing inline page ref: (p. 144)
        pageref = ""
        m_ref = _INLINE_PAGEREF_RE.search(raw)
        if m_ref:
            pageref = raw[m_ref.start():].strip()
            raw = raw[:m_ref.start()]

        remainder = raw

        # Leading [→]
        if remainder.startswith("[→]"):
            _add_run(para, "[→] ", size=10, color=COLOR_MARKER)
            remainder = remainder[4:].lstrip()

        # Dialogue label
        m_dial = _DIALOGUE_RE.match(remainder)
        if m_dial:
            label = m_dial.group(1)
            color = COLOR_HINDU if label == "Хинду:" else COLOR_SIKH
            _add_run(para, label + " ", bold=True, color=color)
            remainder = remainder[m_dial.end():]

        # Split on [→] markers, emit inner quotes as italic
        parts = _MARKER_RE.split(remainder)
        for i, part in enumerate(parts):
            if part:
                _emit_with_quotes(para, part, base_color=COLOR_RUSSIAN)
            if i < len(parts) - 1:
                _add_run(para, " [→] ", size=10, color=COLOR_MARKER)

        # Trailing page ref
        if pageref:
            _add_run(para, "  " + pageref, size=9, color=COLOR_FOOTNOTE)


# ── DOCX builder ──────────────────────────────────────────────────────────────

def build_docx(start: int, end: int, output: Path) -> None:
    doc = Document()

    # Поля чуть меньше для книжного вида
    for section in doc.sections:
        section.top_margin    = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin   = Pt(80)
        section.right_margin  = Pt(72)

    title = doc.add_heading("Мы — не индуисты", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Кан Сингх Набха (Kahn Singh Nabha)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = COLOR_FOOTNOTE
    sub.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    pages = sorted(JSON_DIR.glob("page_*.json"), key=lambda p: _page_num(p))
    pages = [p for p in pages if start <= _page_num(p) <= end]

    if not pages:
        print("No JSON files found — nothing to build.")
        return

    for json_path in pages:
        data = json.loads(json_path.read_text())
        pnum = data["page"]
        translation = data.get("translation_ru", "").strip()
        flags = data.get("quality_flags", [])

        # Пустая страница — только метка, без контента
        if data.get("empty_page"):
            label_para = doc.add_paragraph()
            label_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            label_para.paragraph_format.space_before = Pt(0)
            label_para.paragraph_format.space_after  = Pt(2)
            run = label_para.add_run(f"стр. {pnum}  [пустая]")
            run.font.size = Pt(8)
            run.font.color.rgb = COLOR_PAGE
            continue

        # Метка страницы — маленький правосторонний маркер, не разрывает поток
        label_para = doc.add_paragraph()
        label_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        label_para.paragraph_format.space_before = Pt(0)
        label_para.paragraph_format.space_after  = Pt(2)
        run = label_para.add_run(f"стр. {pnum}")
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_PAGE

        # Флаги качества (если есть)
        if flags:
            warn = doc.add_paragraph()
            warn.paragraph_format.space_after = Pt(2)
            _add_run(warn, f"⚠ {', '.join(flags)}", size=9, color=COLOR_FLAG)

        if translation:
            add_translation_block(doc, translation)

    if output.exists():
        backup = output.with_suffix(".bak.docx")
        output.rename(backup)
        print(f"  Backup: {backup.name}")

    doc.save(str(output))
    print(f"DOCX saved: {output}  ({len(pages)} pages)")


# ── CLI ───────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description="we_are_not_hindu PDF → Russian DOCX via ChatGPT (v2)")
    parser.add_argument("--start", type=int, default=1)
    parser.add_argument("--end", type=int, default=9999)
    parser.add_argument("--dpi", type=int, default=200)
    parser.add_argument("--delay", type=float, default=4.0)
    parser.add_argument("--max-retries", type=int, default=3)
    parser.add_argument("--output", type=Path, default=OUTPUT_DOCX)
    parser.add_argument("--docx-start", type=int, default=1)
    parser.add_argument("--docx-end", type=int, default=9999)
    parser.add_argument("--rebuild-docx", action="store_true",
                        help="Rebuild DOCX from existing JSON, skip ChatGPT")
    parser.add_argument("--render-only", action="store_true")
    parser.add_argument("--force", action="store_true")
    parser.add_argument("--audit", action="store_true")
    parser.add_argument("--scan-missing", action="store_true",
                        help="Print list of missing pages and exit")
    parser.add_argument("--fix-missing", action="store_true",
                        help="Translate only missing pages (in order), then rebuild DOCX")
    args = parser.parse_args()

    print(f"Rendering PDF: {PDF_PATH.name}")
    image_paths = pdf_to_images(PDF_PATH, PAGES_DIR, dpi=args.dpi)
    total = len(image_paths)
    end = min(args.end, total)
    print(f"Total pages: {total}  |  Processing: {args.start}–{end}")

    if args.render_only:
        return

    if args.audit:
        audit_json(args.docx_start, args.docx_end)
        return

    if args.scan_missing:
        scan_missing(image_paths, args.start, end)
        return

    if args.fix_missing:
        missing = find_missing_pages(image_paths, args.start, end)
        if not missing:
            print("Нет пропущенных страниц — всё готово.")
            return
        scan_missing(image_paths, args.start, end)
        translate_pages(
            image_paths=image_paths,
            start=args.start,
            end=end,
            max_retries=args.max_retries,
            delay_s=args.delay,
            force=False,
            explicit_pages=missing,
        )
        build_docx(args.docx_start, args.docx_end, args.output)
        return

    if args.rebuild_docx:
        build_docx(args.docx_start, args.docx_end, args.output)
        return

    translate_pages(
        image_paths=image_paths,
        start=args.start,
        end=end,
        max_retries=args.max_retries,
        delay_s=args.delay,
        force=args.force,
    )
    build_docx(args.docx_start, args.docx_end, args.output)


if __name__ == "__main__":
    main()
